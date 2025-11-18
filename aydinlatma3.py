import re
import os
import json
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
import pandas as pd
from dataclasses import dataclass, asdict
import logging

try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class LightingCriteria:
    """Aydınlatma Ölçüm Raporu kriterleri veri sınıfı"""
    genel_rapor_bilgileri: Dict[str, Any]
    olcum_metodu_standart: Dict[str, Any]
    olcum_sonuc_tablosu: Dict[str, Any]
    uygunluk_degerlendirmesi: Dict[str, Any]
    gorsel_teknik_dokumantasyon: Dict[str, Any]
    olcum_cihazi_bilgileri: Dict[str, Any]
    sonuc_oneriler: Dict[str, Any]

@dataclass
class LightingAnalysisResult:
    """Aydınlatma Ölçüm Raporu analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

class LightingReportAnalyzer:
    """Aydınlatma Ölçüm Raporu analiz sınıfı"""
    
    def __init__(self):
        logger.info("Aydınlatma Ölçüm Raporu analiz sistemi başlatılıyor...")
        
        self.criteria_weights = {
            "Genel Rapor Bilgileri": 10,
            "Ölçüm Metodu ve Standart Referansları": 15,
            "Ölçüm Sonuç Tablosu": 25,
            "Uygunluk Değerlendirmesi": 20,
            "Görsel ve Teknik Dokümantasyon": 5,
            "Ölçüm Cihazı Bilgileri": 10,
            "Sonuç ve Öneriler": 15
        }
        
        self.criteria_details = {
            "Genel Rapor Bilgileri": {
                "proje_adi_numarasi": {"pattern": r"(?:Proje\s*Ad[ıi]|Project\s*Name|Proje\s*No|Project\s*Number|Proje\s*Kodu|Project\s*Code|İş\s*No|Job\s*Number|Sipariş\s*No|Order\s*Number)", "weight": 2},
                "olcum_rapor_tarihleri": {"pattern": r"(?:Ölçüm\s*Tarih|Measurement\s*Date|Rapor\s*Tarih|Report\s*Date|Test\s*Tarih|Test\s*Date|Analiz\s*Tarih|Analysis\s*Date|Değerlendirme\s*Tarih|\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4}|\d{2,4}[\/\.\-]\d{1,2}[\/\.\-]\d{1,2})", "weight": 2},
                "tesis_bolge_alan": {"pattern": r"(?:Tesis|Facility|Fabrika|Factory|Ofis|Office|Bina|Building|İş\s*yeri|Workplace|Alan|Area|Bölge|Zone|Konum|Location|Adres|Address|Mekân|Space)", "weight": 1},
                "rapor_no_revizyon": {"pattern": r"(?:Rapor\s*No|Report\s*No|Rapor\s*Numaras[ıi]|Report\s*Number|Rev|Revizyon|Revision|Ver|Version|Sürüm|Döküman\s*No|Document\s*No)", "weight": 2},
                "olcumu_yapan_firma": {"pattern": r"(?:Ölçümü\s*Yapan|Measured\s*By|Test\s*Yapan|Tested\s*By|Firma|Company|Şirket|Corporation|Kurum|Institution|Organizasyon|Organization|Muayene\s*Kuruluş)", "weight": 1},
                "onay_imza": {"pattern": r"(?:Onay|Approval|İmza|Signature|Onaylayan|Approved\s*By|İmzalayan|Signed\s*By|Sorumlu|Responsible|Yetkili|Authorized|Mühür|Stamp|Seal)", "weight": 2}
            },
            "Ölçüm Metodu ve Standart Referansları": {
                "olcum_cihazi": {"pattern": r"(?:Lüksmetre|Luxmeter|Lux\s*Meter|Işık\s*Ölçer|Light\s*Meter|Işıklılık\s*Ölçer|Luminance\s*Meter|Fotometre|Photometer|Cihaz\s*Marka|Device\s*Brand|Model|Seri\s*No|Serial\s*Number)", "weight": 4},
                "kalibrasyon_bilgi": {"pattern": r"(?:Kalibrasyon|Calibration|Sertifika|Certificate|Akreditasyon|Accreditation|Geçerlilik|Validity|Son\s*Kalibrasyon|Last\s*Calibration|Kalibre|Calibrated)", "weight": 3},
                "olcum_yontemi": {"pattern": r"(?:Ölçüm\s*Yöntem|Measurement\s*Method|Test\s*Yöntem|Test\s*Method|Prosedür|Procedure|Metodoloji|Methodology|Ortlama|Average|Dört\s*Nokta|Four\s*Point|Grid|Izgara)", "weight": 4},
                "standartlar": {"pattern": r"(?:TS\s*EN\s*12464|ISO\s*8995|İş\s*Sağlığ[ıi]|Work\s*Safety|Occupational\s*Health|Standart|Standard|Norm|Specification|Tüzük|Regulation|Yönetmelik|Directive)", "weight": 4}
            },
            "Ölçüm Sonuç Tablosu": {
                "tablo_yapisi": {"pattern": r"(?:Tablo|Table|Liste|List|Sıra\s*No|Row\s*No|S[ıi]ra|Order|Numara|Number|Index|Çizelge|Chart|Matris|Matrix)", "weight": 5},
                "calisma_alani": {"pattern": r"(?:Çalışma\s*Alan[ıi]|Work\s*Area|İş\s*Alan[ıi]|Work\s*Zone|Bölge\s*Ad[ıi]|Area\s*Name|Konum|Location|Nokta|Point|Pozisyon|Position)", "weight": 4},
                "olculen_degerler": {"pattern": r"(?:Lüks|Lux|lx|Aydınlatma\s*Şiddet|Illumination|Light\s*Level|Işık\s*Seviye|Light\s*Intensity|Ölçülen|Measured|Mevcut|Current|Actual)", "weight": 8},
                "hedeflenen_degerler": {"pattern": r"(?:Hedeflenen|Target|İstenen|Desired|Gerekli|Required|Minimum|Standart|Standard|Önerilen|Recommended|Limit|Thresh)", "weight": 4},
                "uygunluk_durumu": {"pattern": r"(?:Uygun|Suitable|Conform|Uygunsuz|Not\s*Suitable|Non[\\-\\s]*Conform|Geçerli|Valid|Geçersiz|Invalid|PASS|FAIL|OK|NOK)", "weight": 4}
            },
            "Uygunluk Değerlendirmesi": {
                "toplu_degerlendirme": {"pattern": r"(?:Genel\s*Değerlendirme|General\s*Assessment|Toplu\s*Değerlendirme|Overall\s*Evaluation|Özet|Summary|Sonuç|Result|Analiz|Analysis)", "weight": 5},
                "limit_disi_degerler": {"pattern": r"(?:Limit\s*Dış[ıi]|Out\s*of\s*Limit|Standart\s*Dış[ıi]|Non[\\-\\s]*Standard|Uygunsuz|Non[\\-\\s]*Compliant|Eksik|Insufficient|Fazla|Excessive|Aş[ıi]r[ıi]|Over)", "weight": 5},
                "risk_belirtme": {"pattern": r"(?:Risk|Tehlike|Hazard|Göz\s*Yorgunluk|Eye\s*Fatigue|Verimlilik|Productivity|Güvenlik|Safety|Dikkat\s*Dağ[ıi]n[ıi]kl[ıi]|Distraction)", "weight": 5},
                "duzeltici_faaliyet": {"pattern": r"(?:Düzeltici\s*Faaliyet|Corrective\s*Action|İyileştirme|Improvement|Öneri|Recommendation|Çözüm|Solution|Aksiyon|Action|Tedbir|Measure)", "weight": 5}
            },
            "Görsel ve Teknik Dokümantasyon": {
                "alan_fotograflari": {"pattern": r"(?:Fotoğraf|Photo|Görsel|Visual|Resim|Picture|Image|Şekil|Figure|Çekim|Shot)", "weight": 1},
                "cihaz_fotograflari": {"pattern": r"(?:Cihaz\s*Fotoğraf|Device\s*Photo|Alet\s*Fotoğraf|Equipment\s*Photo|Ölçüm\s*Cihaz|Measurement\s*Device)", "weight": 1},
                "kroki_sema": {"pattern": r"(?:Kroki|Sketch|Şema|Schema|Plan|Layout|Çizim|Drawing|Diyagram|Diagram|Harita|Map|Yerleşim|Placement)", "weight": 1},
                "armatur_teknik": {"pattern": r"(?:Armatür|Fixture|Lamba|Lamp|LED|Fotometrik|Photometric|Lümen|Lumen|lm|Wat|Watt|W|Işık\s*Ak[ıi]s[ıi]|Luminous\s*Flux|Verim|Efficacy)", "weight": 2}
            },
            "Ölçüm Cihazı Bilgileri": {
                "cihaz_detay": {"pattern": r"(?:Marka|Brand|Model|Tip|Type|Seri|Serial|SN|S/N|Üretici|Manufacturer|Kalibrasyon|Calibration)", "weight": 5},
                "cihaz_ozellikleri": {"pattern": r"(?:Hassasiyet|Accuracy|Precision|Doğruluk|Range|Aralık|Ölçüm\s*Aral[ıi]ğ[ıi]|Measurement\s*Range|Çözünürlük|Resolution)", "weight": 3},
                "cihaz_durumu": {"pattern": r"(?:Durum|Status|Çalışır|Working|Aktif|Active|Geçerli|Valid|Uygun|Suitable|Kullan[ıi]labilir|Usable)", "weight": 2}
            },
            "Sonuç ve Öneriler": {
                "genel_uygunluk": {"pattern": r"(?:Genel\s*Sonuç|Overall\s*Result|Uygun|Suitable|Uygunsuz|Non[\\-\\s]*Suitable|Geçerli|Valid|Geçersiz|Invalid|PASS|FAIL|Başar[ıi]l[ıi]|Successful)", "weight": 4},
                "standart_atif": {"pattern": r"(?:Standart|Standard|TS\s*EN|ISO|Referans|Reference|Atıf|Citation|Uygunluk|Compliance|Conformity)", "weight": 3},
                "iyilestirme_onerileri": {"pattern": r"(?:İyileştirme|Improvement|Öneri|Recommendation|Aksiyon|Action|Tedbir|Measure|Çözüm|Solution|Gelişt|Develop)", "weight": 4},
                "tekrar_olcum": {"pattern": r"(?:Tekrar\s*Ölçüm|Re[\\-\\s]*Measurement|Periyot|Period|S[ıi]kl[ıi]k|Frequency|Süre|Duration|Kontrol|Control|İzleme|Monitoring)", "weight": 4}
            }
        }
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni Türkçe'ye çevir - şimdilik devre dışı"""
        if source_lang != 'tr':
            logger.info(f"Tespit edilen dil: {source_lang.upper()} - Çeviri yapılmıyor, orijinal metin kullanılıyor")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosya formatına göre metin çıkarma"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.doc':
            logger.warning("DOC formatı için DOCX'e dönüştürme gerekiyor veya OCR kullanılacak")
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatı: {file_ext}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den sadece PyPDF2 ile metin çıkarma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF analizi başlatılıyor - Toplam sayfa: {total_pages}")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = page_text.replace('|', ' ')
                    page_text = page_text.strip()
                    
                    all_text += page_text + "\n"
                
                # Metni temizle
                all_text = all_text.replace('—', '-')
                all_text = all_text.replace('"', '"').replace('"', '"')
                all_text = all_text.replace('´', "'")
                all_text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]+', ' ', all_text)
                all_text = all_text.strip()
                
                logger.info(f"✅ PDF analizi tamamlandı:")
                logger.info(f"   📊 Toplam metin uzunluğu: {len(all_text):,} karakter")
                logger.info(f"   📄 Toplam sayfa: {total_pages}")
                
                return all_text
                
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Paragrafları oku
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Tablolardan metin çıkar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            logger.info(f"DOCX'den {len(text)} karakter metin çıkarıldı")
            return text
            
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasından metin çıkarma"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text:
                # UTF-8 başarısız olursa diğer encoding'leri dene
                encodings = ['cp1254', 'iso-8859-9', 'latin1']
                for encoding in encodings:
                    try:
                        with open(txt_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        if text:
                            break
                    except:
                        continue
            
            logger.info(f"TXT'den {len(text)} karakter metin çıkarıldı")
            return text.strip()
            
        except Exception as e:
            logger.error(f"TXT metin çıkarma hatası: {e}")
            return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, LightingAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                # Scoring algoritması
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
                    
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = LightingAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={
                    "pattern_used": pattern,
                    "matches_count": len(matches) if matches else 0
                }
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, LightingAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """Aydınlatma raporuna özgü değerleri çıkar"""
        values = {
            "rapor_numarasi": "Bulunamadı",
            "proje_adi": "Bulunamadı",
            "olcum_tarihi": "Bulunamadı", 
            "rapor_tarihi": "Bulunamadı",
            "olcum_cihazi": "Bulunamadı",
            "tesis_adi": "Bulunamadı",
            "genel_uygunluk": "Bulunamadı"
        }
        
            # RAPOR NUMARASI - Bu bölümü ekleyin (ölçüm tarihi pattern'larından önce)
        report_no_patterns = [
            # Tablo yapısında rapor no
            r"(?i)(?:RAPOR\s*NO|REPORT\s*NO|RAPOR\s*NUMARAS[II])\s*[|\s]*\s*([A-Z0-9\-\/]{3,20})",
            
            # Klasik format'lar
            r"(?i)(?:RAPOR\s*NO|REPORT\s*NO|RAPOR\s*NUMARAS[II]|REPORT\s*NUMBER|DÖKÜMAN\s*NO|DOCUMENT\s*NO)\s*[:=]?\s*([A-Z0-9\-\/]{3,20})",
            
            r"(?i)(?:TEST\s*NO|BELGE\s*NO|REFERANS\s*NO|REF\s*NO)\s*[:=]?\s*([A-Z0-9\-\/]{3,20})"
        ]
        
        for pattern in report_no_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 3 <= len(result) <= 20:
                    values["rapor_numarasi"] = result
                    break

        # PROJE ADI
        project_patterns = [
            # Başlık varsa ve başında gereksiz karakterler olabilir
            r"(?i)(?:PROJE\s*ADI|PROJECT\s*NAME|PROJE\s*TANIM|PROJECT\s*TITLE)?\s*[:=-]?\s*(?:\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\s*REV\s*-?\d*\s*\d*\.*\s*BÖLÜM\s*)?[-\s]*(.*?(?:ÖLÇÜM|TEST|RAPOR|REPORT))",
            
            # Başta tarih/REV yok, sadece ÖLÇÜM/TEST/RAPOR/REPORT ile bitiyor
            r"(?i)[-\s]*([A-ZÇĞİÖŞÜ0-9\s]{5,100}.*?(?:ÖLÇÜM|TEST|RAPOR|REPORT))",

            # Başlık + proje adı, tek satır
            r"(?i)(?:PROJE\s*ADI|PROJECT\s*NAME|PROJE\s*TANIM|PROJECT\s*TITLE)\s*[:=-]?\s*[-\s]*(.*?(?:ÖLÇÜM|TEST|RAPOR|REPORT))",

            # Satır başında tarih veya REV var
            r"(?i)(?:\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\s*REV\s*-?\d*\s*\d*\.*\s*BÖLÜM\s*)[-\s]*(.*?(?:ÖLÇÜM|TEST|RAPOR|REPORT))"
        ]

        
        for pattern in project_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                result = re.sub(r'\s+', ' ', result)  # fazla boşlukları temizle
                if 3 <= len(result) <= 100:
                    values["proje_adi"] = result
                    break
        
        # ÖLÇÜM TARİHİ
        # ÖLÇÜM TARİHİ - Geliştirilmiş pattern'lar
        measurement_date_patterns = [
            # Tablo yapısı için: "Ölçüm Tarihi / Saati" sonrası gelen tarih
            r"(?i)(?:ÖLÇÜM\s*TARİH[İI]?\s*\/?\s*SAAT[İI]?)\s*[|\s]*\s*(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
            
            # Saat aralığı ile birlikte: 28.11.2023 / 10:58-11:07
            r"(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})\s*\/?\s*\d{1,2}:\d{1,2}[\-:]\d{1,2}:\d{1,2}",
            
            # Klasik format'lar
            r"(?i)(?:ÖLÇÜM\s*TARİH[İI]?|MEASUREMENT\s*DATE|TEST\s*TARİH[İI]?|TEST\s*DATE)\s*[:=]\s*(\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})",
            
            # Tablo hücresi içinde sadece tarih
            r"(?i)(?:ÖLÇÜM.*?(?:TARİH|DATE).*?)\s*[|\s]*\s*(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
            
            # Alternatif tablo yapıları
            r"(?i)(?:ÖLÇÜM\s*YAPILDI[ĞG]I|MEASURED\s*ON)\s*[:=]?\s*(\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})",
            
            # PDF'den çıkan metin bozuklukları için
            r"(?:TARİH[İI]?\s*\/?\s*SAAT[İI]?).*?(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
            
            # Genel tarih pattern'ı (en son denenir)
            r"\b(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{4})\b"
        ]
        
        for pattern in measurement_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["olcum_tarihi"] = match.group(1)
                break
        
        # RAPOR TARİHİ
        report_date_patterns = [
            # Rapor tarihi tablo yapısı
            r"(?i)(?:RAPOR\s*TARİH[İI]?\s*\/?\s*SAAT[İI]?)\s*[|\s]*\s*(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
            
            # Klasik format'lar
            r"(?i)(?:RAPOR\s*TARİH[İI]?|REPORT\s*DATE|HAZIRLANMA\s*TARİH[İI]?|PREPARED\s*ON|DÜZENLEME\s*TARİH[İI]?)\s*[:=]\s*(\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})",
            
            # Belge başlığında tarih
            r"(?i)(?:BELGE|DÖKÜMAN|DOCUMENT).*?TARİH[İI]?.*?(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})"
        ]
        
        for pattern in report_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["rapor_tarihi"] = match.group(1)
                break
    
        # Eğer rapor tarihi bulunamazsa, ölçüm tarihi ile aynı olabilir
        if values["rapor_tarihi"] == "Bulunamadı" and values["olcum_tarihi"] != "Bulunamadı":
            values["rapor_tarihi"] = "Rapor tarihi ayrı belirtilmemiş"

        # ÖLÇÜM CİHAZI
        device_patterns = [
        r"(?i)(?:LÜKSMETRE|LUXMETER|LUX\s*METER)\s*[:=]?\s*([A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜ0-9\s\.\-]{2,50}(?=\s*(?:\d{4,}|Seri|No|Nolu|$)))",
        r"(?i)(?:MARKA|BRAND)\s*[:=]?\s*([A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜ0-9\s\.\-]{2,30})(?:.*?(?:MODEL|TİP|TYPE)\s*[:=]?\s*([A-ZÇĞİÖŞÜ0-9\-]{1,30}))?",
        r"(?i)(?:CİHAZ|DEVICE|ALET|INSTRUMENT)\s*[:=]?\s*([A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜ0-9\s\.\-]{2,50}(?=\s*(?:\d{4,}|Seri|No|Nolu|$)))",
        r"(?i)([A-ZÇĞİÖŞÜ0-9]+(?:\s+[A-ZÇĞİÖŞÜ0-9]+){0,3}\s+(?:LUXMETER|LUX\s*METER|LÜKSMETRE)(?=\s*(?:\d{4,}|Seri|No|Nolu|$)))",
        r"(?i)(?:IŞIK\s*ŞİDDETİ\s*ÖLÇÜM\s*CİHAZI)\s*([A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜ0-9\s\.\-]{2,50}(?=\s*(?:\d{4,}|Seri|No|Nolu|$)))"
        ]

        
        for pattern in device_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 2 <= len(result) <= 35:
                    values["olcum_cihazi"] = result
                    break
        
        # TESİS ADI
        facility_patterns = [
            r"(?<!\w)([A-ZÇĞİÖŞÜ][A-Za-zÇĞİÖŞÜçğıöşü\s\.\&]{2,50}?(?:A\.Ş|LTD\.?\s*ŞTİ|SANAYİ|TİCARET|İÇECEK|GIDA|TEKSTİL|OTOMOTİV|İNŞAAT|MAKİNA|COMPANY|CORP|CORPORATION|INC|INCORPORATED)\.?)(?!\w)",
            r"(?i)(?:TESİS\s*ADI|FACILITY\s*NAME|FABRİKA|FACTORY)\s*[:\-]?\s*([A-ZÇĞİÖŞÜ0-9][A-ZÇĞİÖŞÜ0-9\.\&\süçğıöş\-]{3,60})",
            r"(?i)(?:İŞ\s*YERİ|WORKPLACE|KURULUŞ|COMPANY)\s*[:\-]?\s*([A-ZÇĞİÖŞÜ0-9][A-ZÇĞİÖŞÜ0-9\.\&\süçğıöş\-]{3,60})",
            r"(?i)([A-ZÇĞİÖŞÜ0-9][A-ZÇĞİÖŞÜ0-9\.\&\süçğıöş\-]{3,60})\s+(?:TESİS|FABRİKA|FACTORY|PLANT)",
            r"(?i)([A-ZÇĞİÖŞÜ0-9][A-ZÇĞİÖŞÜ0-9\-\s\.&]*?(?:A\.Ş|LTD\.?\s*ŞTİ|SANAYİ|TİCARET)[A-ZÇĞİÖŞÜ0-9\-\s\.&]*)",
        ]

        for pattern in facility_patterns:
            matches = re.findall(pattern, text)
            for result in matches:
                result = result.strip()
                if 3 <= len(result) <= 60: 
                    values["tesis_adi"] = result
                    break
            if "tesis_adi" in values:
                break

        
        # GENEL UYGUNLUK
        compliance_patterns = [
            r"(?i)\b(UYGUN|SUITABLE|CONFORM|GEÇERLİ|VALID|PASS)\b",
            r"(?i)\b(UYGUNSUZ|NOT\s*SUITABLE|NON[\\-\\s]*CONFORM|GEÇERSİZ|INVALID|FAIL)\b",
            r"(?i)(?:GENEL\s*SONUÇ|OVERALL\s*RESULT|SONUÇ|RESULT)\s*[:=]?\s*(UYGUN|UYGUNSUZ|SUITABLE|NOT\s*SUITABLE|PASS|FAIL|GEÇERLİ|GEÇERSİZ)"
        ]
        
        for pattern in compliance_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip().upper()
                if result in ["UYGUN", "SUITABLE", "CONFORM", "GEÇERLİ", "VALID", "PASS"]:
                    values["genel_uygunluk"] = "UYGUN"
                elif result in ["UYGUNSUZ", "NOT SUITABLE", "NON-CONFORM", "GEÇERSİZ", "INVALID", "FAIL"]:
                    values["genel_uygunluk"] = "UYGUNSUZ"
                break
        
        return values

    def validate_lighting_document(self, text: str) -> bool:
        """Dokümanın aydınlatma ölçüm raporu olup olmadığını kontrol et"""
        
        lighting_keywords = [
            # Temel aydınlatma terimleri
            "aydınlatma", "lighting", "lüks", "lux", "ışık", "light", "luminans", "luminance",
            
            # Ölçüm terimleri  
            "ölçüm", "measurement", "test", "analiz", "analysis", "değerlendirme", "assessment",
            
            # Cihaz terimleri
            "lüksmetre", "luxmeter", "fotometre", "photometer", "ışıkölçer", "light meter",
            
            # Rapor terimleri
            "rapor", "report", "değerlendirme", "evaluation", "sonuç", "result",
            
            # Standart terimleri
            "standart", "standard", "ts en", "iso", "norm", "specification",
            
            # Tesis terimleri
            "işyeri", "workplace", "ofis", "office", "fabrika", "factory", "tesis", "facility",
            
            # Uygunluk terimleri
            "uygun", "suitable", "uygunsuz", "unsuitable", "geçerli", "valid"
        ]
        
        # En az 2 anahtar kelime bulunmalı
        found_keywords = 0
        found_words = []
        
        for keyword in lighting_keywords:
            if re.search(rf"\b{keyword}\b", text, re.IGNORECASE):
                found_keywords += 1
                found_words.append(keyword)
                
        logger.info(f"Doküman validasyonu: {found_keywords} anahtar kelime bulundu: {found_words[:10]}")
        
        return found_keywords >= 2

    def check_date_validity(self, measurement_date: str, report_date: str) -> Tuple[bool, str]:
        """Ölçüm ve rapor tarihlerini bugünkü tarih ile kontrol et (1 yıl kuralı)"""
        
        # Hiçbir tarih bulunamadıysa
        if measurement_date == "Bulunamadı" and report_date == "Bulunamadı":
            return False, "Ne ölçüm ne de rapor tarihi bulunamadı"
        
        # Kontrol edilecek tarihleri topla
        dates_to_check = []
        
        if measurement_date != "Bulunamadı":
            dates_to_check.append(("Ölçüm", measurement_date))
        
        if report_date != "Bulunamadı" and report_date != "Rapor tarihi ayrı belirtilmemiş":
            dates_to_check.append(("Rapor", report_date))
        
        # Hiç kontrol edilecek tarih yoksa
        if not dates_to_check:
            return False, "Geçerli tarih bulunamadı"
        
        try:
            date_formats = ['%d/%m/%Y', '%d.%m.%Y', '%d-%m-%Y', '%d/%m/%y', '%d.%m.%y', '%d-%m-%y']
            today = datetime.now()
            
            for date_type, date_str in dates_to_check:
                # Tarihi parse et
                parsed_date = None
                for fmt in date_formats:
                    try:
                        parsed_date = datetime.strptime(date_str, fmt)
                        break
                    except:
                        continue
                
                if not parsed_date:
                    return False, f"{date_type} tarihi formatı tanınmadı ({date_str})"
                
                # 2 haneli yıl düzeltmesi
                if parsed_date.year < 1950:
                    parsed_date = parsed_date.replace(year=parsed_date.year + 2000)
                
                # Bugünden farkı hesapla
                diff = abs((today - parsed_date).days)
                
                # 1 yıldan (365 gün) eskiyse geçersiz
                if diff > 365:
                    return False, f"{date_type} tarihi 1 yıldan eski ({diff} gün önce - {date_str})"
            
            # Tüm tarihler geçerliyse
            checked_dates = [f"{dt[0]}: {dt[1]}" for dt in dates_to_check]
            return True, f"Tüm tarihler geçerli - {', '.join(checked_dates)}"
            
        except Exception as e:
            return False, f"Tarih kontrolü yapılamadı - {e}"
        
    def generate_improvement_actions(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Dinamik iyileştirme önerileri oluştur"""
        actions = []
        
        # Kategorileri puana göre sırala (düşükten yükseğe)
        sorted_categories = sorted(
            scores["category_scores"].items(), 
            key=lambda x: x[1]["percentage"]
        )
        
        category_actions = {
            "Genel Rapor Bilgileri": [
                "Proje adı ve numarasını netleştiriniz",
                "Ölçüm ve rapor tarihlerini açıkça belirtiniz",
                "Tesis/bölge/alan bilgilerini detaylandırınız",
                "Rapor numarası ve revizyon bilgisini ekleyiniz",
                "Ölçümü yapan firma ve personel bilgilerini yazınız",
                "Raporu hazırlayanın onay/imzasını alınız"
            ],
            "Ölçüm Metodu ve Standart Referansları": [
                "Ölçüm cihazının marka, model ve seri numarasını belirtiniz",
                "Cihaz kalibrasyon bilgilerini ve sertifikalarını ekleyiniz",
                "Ölçüm yöntemini detaylandırınız (dört nokta ortalaması vb.)",
                "TS EN 12464-1, ISO 8995 gibi standart referanslarını ekleyiniz"
            ],
            "Ölçüm Sonuç Tablosu": [
                "Ölçüm sonuç tablosunu sıra no ile düzenleyiniz",
                "Çalışma alanı/bölge adlarını açıkça tanımlayınız",
                "Ölçülen aydınlatma şiddeti değerlerini (lüks) eksiksiz yazınız",
                "Hedeflenen aydınlatma şiddeti değerlerini belirtiniz",
                "Her nokta için uygunluk durumunu (UYGUN/UYGUN DEĞİL) belirtiniz"
            ],
            "Uygunluk Değerlendirmesi": [
                "Tüm ölçüm noktalarının genel değerlendirmesini yapınız",
                "Limit dışı değerlerin listesini çıkarınız",
                "Yetersiz/aşırı aydınlatmanın risklerini belirtiniz",
                "Somut düzeltici faaliyet önerileri sununuz"
            ],
            "Görsel ve Teknik Dokümantasyon": [
                "Ölçüm yapılan alan fotoğraflarını ekleyiniz",
                "Ölçüm cihazı fotoğraflarını çekiniz",
                "Ölçüm noktaları kroki veya şemasını hazırlayınız",
                "Armatür teknik belgelerini (fotometrik raporlar) ekleyiniz"
            ],
            "Ölçüm Cihazı Bilgileri": [
                "Cihaz marka, model, seri no detaylarını tam olarak yazınız",
                "Cihaz hassasiyet, doğruluk, ölçüm aralığı özelliklerini belirtiniz",
                "Cihazın çalışır durumda olduğunu teyit ediniz"
            ],
            "Sonuç ve Öneriler": [
                "Genel uygunluk sonucunu (UYGUN/UYGUNSUZ) açıkça belirtiniz",
                "İlgili standartlara atıf yapınız",
                "Somut iyileştirme önerilerini listeyiniz",
                "Tekrar ölçüm periyodu önerisinde bulununuz"
            ]
        }
        
        # En düşük puanlı 5 kategori için öneriler ekle
        for category, score_data in sorted_categories[:5]:
            if score_data["percentage"] < 70:
                actions.extend(category_actions.get(category, []))
        
        # Genel öneriler
        if scores["percentage"] < 50:
            actions.insert(0, "ÖNCELİK: Rapor yapısını ve içeriğini kapsamlı olarak yeniden düzenleyiniz")
        
        return actions

    def generate_recommendations(self, analysis_results: Dict, scores: Dict, date_valid: bool = True, date_message: str = "") -> List[str]:
        """Öneriler oluştur - Güncellenmiş versiyon"""
        recommendations = []
        
        total_percentage = scores["percentage"]
        
        # Tarih geçerlilik kontrolü
        if not date_valid:
            # date_message'dan hangi tarih sorunu olduğunu belirle
            if "Ne ölçüm ne de rapor tarihi bulunamadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ne ölçüm ne de rapor tarihi bulunamadı")
            elif "Ölçüm tarihi formatı tanınmadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ölçüm tarihi formatı tanınmadı")
            elif "Rapor tarihi formatı tanınmadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Rapor tarihi formatı tanınmadı")
            elif "Ölçüm tarihi 1 yıldan eski" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ölçüm tarihi 1 yıldan eski")
            elif "Rapor tarihi 1 yıldan eski" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Rapor tarihi 1 yıldan eski")
            elif "Geçerli tarih bulunamadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Geçerli tarih bulunamadı")
            else:
                recommendations.append(f"❌ RAPOR GEÇERSİZ: {date_message}")
            return recommendations
        
        # Tarih geçerliyse normal değerlendirme
        if total_percentage >= 70:
            recommendations.append(f"✅ Aydınlatma Ölçüm Raporu GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ Aydınlatma Ölçüm Raporu GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 40:
                recommendations.append(f"🔴 {category} bölümü yetersiz (%{category_score:.1f})")
                missing_items = [name for name, result in results.items() if not result.found]
                if missing_items:
                    recommendations.append(f"   Eksik: {', '.join(missing_items[:3])}")
            elif category_score < 70:
                recommendations.append(f"🟡 {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"🟢 {category} bölümü yeterli (%{category_score:.1f})")
        
        return recommendations

    def analyze_lighting_report(self, file_path: str) -> Dict[str, Any]:
        """Ana Aydınlatma Ölçüm Raporu analiz fonksiyonu - Güncellenmiş"""
        logger.info("Aydınlatma Ölçüm Raporu analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        if detected_lang != 'tr':
            logger.info(f"{detected_lang.upper()} dilinden Türkçe'ye çeviriliyor...")
            text = self.translate_to_turkish(text, detected_lang)
        
        # Doküman türü kontrolü
        if not self.validate_lighting_document(text):
            return {
                "error": "YANLIŞ DOKÜMAN: Bu dosya aydınlatma ölçüm raporu değil!",
                "document_type": "UNKNOWN",
                "suggestion": "Lütfen geçerli bir aydınlatma ölçüm raporu yükleyiniz."
            }
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        
        # Tarih geçerlilik kontrolü
        date_valid, date_message = self.check_date_validity(
            extracted_values.get("olcum_tarihi", "Bulunamadı"),
            extracted_values.get("rapor_tarihi", "Bulunamadı")
        )
        
        # Önerileri date_message ile birlikte oluştur
        recommendations = self.generate_recommendations(analysis_results, scores, date_valid, date_message)
        improvement_actions = self.generate_improvement_actions(analysis_results, scores)
        
        # Tarih geçersizse rapor otomatik başarısız
        if not date_valid:
            final_status = "FAIL"
            final_percentage = 0
        else:
            final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
            final_percentage = scores["percentage"]
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "tarih_gecerlilik": {
                "valid": date_valid,
                "message": date_message
            },
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "iyilestirme_eylemleri": improvement_actions,
            "ozet": {
                "toplam_puan": scores["total_score"] if date_valid else 0,
                "yuzde": final_percentage,
                "durum": final_status,
                "rapor_tipi": "AYDINLATMA_OLCUM_RAPORU",
                "tarih_gecersiz": not date_valid
            }
        }
        
        return report

def main():
    """Ana fonksiyon - Güncellenmiş"""
    analyzer = LightingReportAnalyzer()

    # Test için örnek dosya yolları
    file_path = r"C:\Users\nuvo_teknik_2\Desktop\PILZ DOCUMENTS\2.4 Aydınlatma\Aydınlatma Ölçüm Raporu.pdf"

    if not os.path.exists(file_path):
        print(f"❌ Dosya bulunamadı: {file_path}")
        return
    
    print("💡 Aydınlatma Ölçüm Raporu Analizi Başlatılıyor...")
    print("=" * 60)
    
    report = analyzer.analyze_lighting_report(file_path)
    
    if "error" in report:
        print(f"❌ Hata: {report['error']}")
        return
    
    print("\n📊 ANALİZ SONUÇLARI")
    print("=" * 60)
    
    print(f"📅 Analiz Tarihi: {report['analiz_tarihi']}")
    print(f"📄 Dosya Tipi: {report['dosya_bilgisi']['file_type'].upper()}")
    print(f"🔍 Tespit Edilen Dil: {report['dosya_bilgisi']['detected_language'].upper()}")
    
    # Tarih geçerlilik kontrolü (güncellenmiş anahtar)
    if report['ozet']['tarih_gecersiz']:
        print(f"⚠️ TARİH GEÇERSİZLİĞİ: {report['tarih_gecerlilik']['message']}")
        print(f"📋 Toplam Puan: 0/100 (Otomatik başarısız)")
        print(f"📈 Yüzde: %0")
        print(f"🎯 Durum: FAIL")
    else:
        print(f"📋 Toplam Puan: {report['ozet']['toplam_puan']}/100")
        print(f"📈 Yüzde: %{report['ozet']['yuzde']}")
        print(f"🎯 Durum: {report['ozet']['durum']}")
        print(f"📅 Tarih Geçerliliği: {report['tarih_gecerlilik']['message']}")
    
    print(f"📄 Rapor Tipi: {report['ozet']['rapor_tipi']}")
    
    print("\n📋 ÖNEMLİ ÇIKARILAN DEĞERLER")
    print("-" * 40)
    for key, value in report['cikarilan_degerler'].items():
        display_name = {
            "rapor_numarasi": "Rapor Numarası",
            "proje_adi": "Proje Adı",
            "olcum_tarihi": "Ölçüm Tarihi",
            "rapor_tarihi": "Rapor Tarihi",
            "olcum_cihazi": "Ölçüm Cihazı",
            "tesis_adi": "Tesis Adı",
            "genel_uygunluk": "Genel Uygunluk"
        }.get(key, key.replace('_', ' ').title())
        print(f"{display_name}: {value}")
    
    # ESKİ: if not report['ozet']['sifir_tolerans_uyguland']:
    # YENİ: if not report['ozet']['tarih_gecersiz']:
    if not report['ozet']['tarih_gecersiz']:
        print("\n📊 KATEGORİ PUANLARI")
        print("-" * 40)
        for category, score_data in report['puanlama']['category_scores'].items():
            print(f"{category}: {score_data['normalized']}/{score_data['max_weight']} (%{score_data['percentage']:.1f})")
    
    print("\n💡 ÖNERİLER VE DEĞERLENDİRME")
    print("-" * 40)
    for recommendation in report['oneriler']:
        print(recommendation)
    
    print("\n📋 GENEL DEĞERLENDİRME")
    print("=" * 60)
    
    # ESKİ: if report['ozet']['sifir_tolerans_uyguland']:
    # YENİ: if report['ozet']['tarih_gecersiz']:
    if report['ozet']['tarih_gecersiz']:
        print("❌ SONUÇ: GEÇERSİZ (TARİH SORUNU)")
        print(f"📝 Değerlendirme: {report['tarih_gecerlilik']['message']}")
        print("\n📌 YAPILMASI GEREKENLER:")
        print("1. Geçerli tarihli ölçümleri yapınız")
        print("2. Rapor tarihini güncelleyiniz")  
        print("3. Tarihlerin 1 yıldan eski olmadığından emin olunuz")
    elif report['ozet']['yuzde'] >= 70:
        print("✅ SONUÇ: GEÇERLİ")
        print(f"🌟 Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Aydınlatma ölçüm raporu genel olarak yeterli kriterleri sağlamaktadır.")
    else:
        print("❌ SONUÇ: GEÇERSİZ")
        print(f"⚠️ Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Aydınlatma ölçüm raporu minimum gereklilikleri sağlamamaktadır.")
        
        print("\n⚠️ EKSİK GEREKLİLİKLER:")
        for category, results in report['kategori_analizleri'].items():
            missing_items = []
            for criterion, result in results.items():
                if not result.found:
                    missing_items.append(criterion)
            
            if missing_items:
                print(f"\n🔍 {category}:")
                for item in missing_items:
                    readable_name = {
                        "proje_adi_numarasi": "Proje Adı/Numarası",
                        "olcum_rapor_tarihleri": "Ölçüm/Rapor Tarihleri",
                        "tesis_bolge_alan": "Tesis/Bölge/Alan Bilgisi",
                        "rapor_no_revizyon": "Rapor No/Revizyon",
                        "proje_adi_numarasi": "Proje Adı/Numarası",
                        "olcumu_yapan_firma": "Ölçümü Yapan Firma",
                        "onay_imza": "Onay/İmza",
                        "olcum_cihazi": "Ölçüm Cihazı",
                        "kalibrasyon_bilgi": "Kalibrasyon Bilgileri",
                        "olcum_yontemi": "Ölçüm Yöntemi",
                        "standartlar": "Standart Referansları",
                        "tablo_yapisi": "Tablo Yapısı",
                        "calisma_alani": "Çalışma Alanı",
                        "olculen_degerler": "Ölçülen Değerler",
                        "hedeflenen_degerler": "Hedeflenen Değerler",
                        "uygunluk_durumu": "Uygunluk Durumu",
                        "toplu_degerlendirme": "Toplu Değerlendirme",
                        "limit_disi_degerler": "Limit Dışı Değerler",
                        "risk_belirtme": "Risk Belirtme",
                        "duzeltici_faaliyet": "Düzeltici Faaliyet",
                        "alan_fotograflari": "Alan Fotoğrafları",
                        "cihaz_fotograflari": "Cihaz Fotoğrafları",
                        "kroki_sema": "Kroki/Şema",
                        "armatur_teknik": "Armatür Teknik Belgeleri",
                        "cihaz_detay": "Cihaz Detayları",
                        "cihaz_ozellikleri": "Cihaz Özellikleri",
                        "cihaz_durumu": "Cihaz Durumu",
                        "genel_uygunluk": "Genel Uygunluk",
                        "standart_atif": "Standart Atıf",
                        "iyilestirme_onerileri": "İyileştirme Önerileri",
                        "tekrar_olcum": "Tekrar Ölçüm"
                    }.get(item, item.replace('_', ' ').title())
                    print(f"   ❌ {readable_name}")
        
        print("\n📌 YAPILMASI GEREKENLER:")
        if "iyilestirme_eylemleri" in report:
            for i, action in enumerate(report['iyilestirme_eylemleri'], 1):
                print(f"{i}. {action}")

if __name__ == "__main__":
    main()
