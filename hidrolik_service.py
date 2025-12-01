# ============================================
# HİDROLİK DEVRE ŞEMASI ANALİZ SERVİSİ
# Standalone Service - Azure App Service Ready
# Database-driven configuration ile dinamik pattern yönetimi
# Port: 8011
# ============================================

# ============================================
# IMPORTS
# ============================================
import os
import json
import io
from datetime import datetime
import re
from typing import Dict, List, Tuple, Any, Optional
import PyPDF2
import pytesseract
from PIL import Image
from dataclasses import dataclass, asdict
import logging
import math
import cv2
import numpy as np
import fitz  # PyMuPDF
from docx import Document

from flask import Flask, request, jsonify, current_app
from werkzeug.utils import secure_filename
import pdf2image

# ============================================
# DATABASE IMPORTS (YENİ)
# ============================================
from database import db, init_db
from db_loader import load_service_config

# ============================================
# CONFIG IMPORT (YENİ)
# ============================================
from config import Config

# ============================================
# LOGGING CONFIGURATION
# ============================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================
# TESSERACT CHECK
# ============================================
def check_tesseract_installation():
    """Check if Tesseract is properly installed"""
    try:
        version = pytesseract.get_tesseract_version()
        logger.info(f"Tesseract OCR kurulu - Sürüm: {version}")
        return True, f"Tesseract v{version}"
    except Exception as e:
        logger.error(f"Tesseract OCR kurulu değil: {e}")
        return False, str(e)

tesseract_available, tesseract_info = check_tesseract_installation()

# ============================================
# ANALİZ SINIFI - DATA CLASSES
# ============================================
@dataclass
class ComponentDetection:
    """Detected component information"""
    component_type: str
    label: str
    position: Tuple[int, int]
    confidence: float
    bounding_box: Tuple[int, int, int, int]

@dataclass
class CircuitAnalysisResult:
    """Analysis result for each criterion"""
    criteria_name: str
    found: bool
    content: str
    score: float
    max_score: float
    details: Dict[str, Any]
    visual_evidence: List[ComponentDetection]

# ============================================
# ANALİZ SINIFI - MAIN ANALYZER
# ============================================
class AdvancedCircuitAnalyzer:
    """Advanced circuit diagram analyzer"""
    
    def __init__(self, app=None):
        logger.info("Hidrolik Devre Şeması analiz sistemi başlatılıyor...")
        
        # Flask app context varsa DB'den yükle, yoksa boş başlat
        if app:
            with app.app_context():
                try:
                    config = load_service_config('hydraulic_circuit')
                    
                    # DB'den yüklenen veriler
                    self.hydraulic_criteria_weights = config.get('criteria_weights', {})
                    self.hydraulic_criteria_details = config.get('criteria_details', {})
                    self.pattern_definitions = config.get('pattern_definitions', {})
                    self.validation_keywords = config.get('validation_keywords', {})
                    self.category_actions = config.get('category_actions', {})
                    
                    # component_templates pattern_definitions içinde
                    self.component_templates = self.pattern_definitions.get('hydraulic', {})
                    
                    logger.info(f"✅ Veritabanından yüklendi: {len(self.hydraulic_criteria_weights)} kategori")
                    
                except Exception as e:
                    logger.error(f"⚠️ Veritabanından yükleme başarısız: {e}")
                    logger.warning("⚠️ Fallback: Boş config kullanılıyor")
                    self.hydraulic_criteria_weights = {}
                    self.hydraulic_criteria_details = {}
                    self.pattern_definitions = {}
                    self.validation_keywords = {}
                    self.category_actions = {}
                    self.component_templates = {}
        else:
            # Flask app yoksa boş başlat
            logger.warning("⚠️ Flask app context yok, boş config kullanılıyor")
            self.hydraulic_criteria_weights = {}
            self.hydraulic_criteria_details = {}
            self.pattern_definitions = {}
            self.validation_keywords = {}
            self.category_actions = {}
            self.component_templates = {}

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using PyPDF2 and OCR fallback"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = page_text.replace('|', ' ')
                    text += page_text + "\n"
                
                text = text.replace('—', '-')
                text = text.replace('"', '"').replace('"', '"')
                text = text.replace('´', "'")
                text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]+', ' ', text)
                text = text.strip()
                
                if len(text.strip()) < 50:
                    logger.info("PyPDF2 extracted minimal text, trying OCR...")
                    ocr_text = self.extract_text_with_ocr(pdf_path)
                    if len(ocr_text) > len(text):
                        text = ocr_text
                
                return text
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            return self.extract_text_with_ocr(pdf_path)

    def extract_text_with_ocr(self, pdf_path: str) -> str:
        """Extract text from PDF using OCR with enhanced settings"""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            
            for page_num in range(min(5, len(doc))):
                page = doc[page_num]
                mat = fitz.Matrix(3.0, 3.0)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.pil_tobytes(format="PNG")
                img = Image.open(io.BytesIO(img_data))
                img_array = np.array(img)
                
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                
                temp_path = f"temp_page_{page_num}.png"
                cv2.imwrite(temp_path, thresh)
                
                ocr_configs = [
                    '--psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,:-/+()[]{}ÇĞİÖŞÜçğıöşü ',
                    '--psm 4',
                    '--psm 11',
                    '--psm 8',
                    '--psm 3'
                ]
                
                page_text = ""
                for config in ocr_configs:
                    try:
                        ocr_result = pytesseract.image_to_string(temp_path, lang='tur+eng+deu', config=config)
                        if len(ocr_result.strip()) > len(page_text.strip()):
                            page_text = ocr_result
                    except:
                        continue
                
                text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"
                
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"OCR extraction error: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""

    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1254', 'iso-8859-9']
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""

    def extract_text_from_image(self, image_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            text = pytesseract.image_to_string(image_path, lang='tur+eng', config='--psm 6')
            return text.strip()
        except Exception as e:
            logger.error(f"Image OCR error: {e}")
            return ""

    def extract_images_from_pdf(self, pdf_path: str) -> List[Any]:
        """Extract images from PDF"""
        logger.info("Image extraction is temporarily disabled")
        return []

    def perform_ocr_on_images(self, images: List[Any]) -> List[str]:
        """Perform OCR on extracted images"""
        logger.info("OCR functionality is temporarily disabled")
        return []

    def detect_components_in_images(self, images: List[Any], circuit_type: str) -> List[ComponentDetection]:
        """Detect components in images"""
        logger.info("Component detection is temporarily disabled")
        return []

    def determine_circuit_type(self, text: str, images: List[Any]) -> Tuple[str, float]:
        """Always return hydraulic"""
        return "hydraulic", 1.0

    def analyze_text_quality(self, text: str) -> str:
        """Analyze OCR text quality"""
        if len(text) < 100:
            return "poor"
        
        technical_terms = len(re.findall(r'(?i)\b(?:hydraulic|pressure|valve|pump|cylinder|motor|system|control|bar|psi)\b', text))
        garbled_patterns = len(re.findall(r'[^a-zA-Z0-9\s]{3,}', text))
        readable_words = len(re.findall(r'\b[a-zA-Z]{3,}\b', text))
        text_density = readable_words / max(1, len(text.split()))
        
        if technical_terms >= 10 and text_density > 0.3:
            return "excellent"
        elif technical_terms >= 3 and text_density > 0.2:
            return "good"  
        elif garbled_patterns > 5 or text_density < 0.1:
            return "poor"
        else:
            return "normal"

    def analyze_criteria(self, text: str, images: List[Any], category: str, circuit_type: str) -> Dict[str, CircuitAnalysisResult]:
        """Analyze criteria"""
        results = {}
        criteria = self.hydraulic_criteria_details.get(category, {})
        
        combined_text = text
        if images:
            ocr_results = self.perform_ocr_on_images(images)
            combined_text += " " + " ".join(ocr_results)
        
        detected_components = self.detect_components_in_images(images, circuit_type)
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            text_matches = re.findall(pattern, combined_text, re.IGNORECASE | re.MULTILINE)
            relevant_components = [comp for comp in detected_components if self._is_relevant_component(comp, criterion_name)]
            
            if text_matches or relevant_components:
                content_parts = []
                if text_matches:
                    content_parts.append(f"Text: {str(text_matches[:3])}")
                if relevant_components:
                    comp_labels = [comp.label for comp in relevant_components[:5]]
                    content_parts.append(f"Components: {comp_labels}")
                
                content = " | ".join(content_parts)
                found = True
                
                text_score = min(weight * 0.8, len(text_matches) * (weight * 0.2))
                component_score = min(weight * 0.2, len(relevant_components) * (weight * 0.1))
                score = text_score + component_score
                score = min(score, weight)
            else:
                content = "Not found"
                found = False
                score = 0
            
            results[criterion_name] = CircuitAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={
                    "pattern_used": pattern,
                    "text_matches": len(text_matches) if text_matches else 0,
                    "visual_matches": len(relevant_components)
                },
                visual_evidence=relevant_components
            )
        
        return results

    def _is_relevant_component(self, component: ComponentDetection, criterion_name: str) -> bool:
        """Check if component is relevant - DB'den"""
        # DB'den relevance_map al
        relevance_map = self.pattern_definitions.get('relevance_map', {})
        relevant_types = relevance_map.get(criterion_name, [])
        return component.component_type in relevant_types

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, CircuitAnalysisResult]], circuit_type: str) -> Dict[str, Any]:
        """Calculate scores"""
        category_scores = {}
        total_score = 0
        
        sample_text = getattr(self, '_last_extracted_text', '')
        text_quality = self.analyze_text_quality(sample_text)
        
        logger.info(f"Detected text quality: {text_quality}")

        for category, results in analysis_results.items():
            category_max = self.hydraulic_criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())

            if category_possible > 0:
                raw_percentage = category_earned / category_possible
                
                if text_quality == "excellent":
                    if raw_percentage > 0.7:
                        adjusted_percentage = raw_percentage
                    elif raw_percentage > 0.3:
                        adjusted_percentage = raw_percentage * 1.1
                    else:
                        adjusted_percentage = raw_percentage * 0.8
                        
                elif text_quality == "poor":
                    if raw_percentage > 0:
                        adjusted_percentage = max(0.6, math.pow(raw_percentage, 0.3))
                    else:
                        adjusted_percentage = 0.3
                        
                else:
                    if raw_percentage > 0:
                        adjusted_percentage = max(0.4, math.pow(raw_percentage, 0.6))
                    else:
                        adjusted_percentage = 0.1
                    
                normalized_score = min(category_max, adjusted_percentage * category_max)
            else:
                normalized_score = 0

            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round((normalized_score / category_max * 100), 2)
            }

            total_score += normalized_score

        total_found_criteria = sum(
            sum(1 for result in results.values() if result.found) 
            for results in analysis_results.values()
        )
        total_possible_criteria = sum(
            len(results) for results in analysis_results.values()
        )
        
        hydraulic_validity_percentage = (total_found_criteria / total_possible_criteria * 100) if total_possible_criteria > 0 else 0
        
        if hydraulic_validity_percentage >= 25:
            logger.info(f"Hidrolik geçerlilik %{hydraulic_validity_percentage:.1f} - Otomatik geçer puan veriliyor")
            total_score = max(total_score, 75.0)
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "total_max_score": 100,
            "overall_percentage": round((total_score / 100 * 100), 2),
            "text_quality": text_quality
        }

    def extract_specific_values(self, text: str, circuit_type: str) -> Dict[str, Any]:
        """Extract specific values - Enhanced for OCR - DB'den"""
        values = {
            "proje_no": "Bulunamadı",
            "sistem_tipi": "Bulunamadı",
            "tarih": "Bulunamadı",
            "hidrolik_unite": "Bulunamadı",
            "tank_hacmi": "Bulunamadı",
            "motor_gucu": "Bulunamadı",
            "devir": "Bulunamadı",
            "debi": "Bulunamadı",
            "tambur": "Bulunamadı"
        }
        
        # DB'den extract_values pattern'lerini al
        extract_values = self.pattern_definitions.get('extract_values', {})
        
        # Her field için pattern listesini çek ve uygula
        for field_name in values.keys():
            patterns = extract_values.get(field_name, [])
            for pattern in patterns:
                match = re.search(pattern, text)
                if match:
                    if match.groups():
                        values[field_name] = next((m for m in match.groups() if m), match.group())
                    else:
                        values[field_name] = match.group()
                    break
        
        return values

    def generate_recommendations(self, analysis_results: Dict, scores: Dict, circuit_type: str) -> List[str]:
        """Generate recommendations"""
        recommendations = []
        
        valid_criteria_count = sum(1 for category, results in analysis_results.items() 
                                 for result in results.values() if result.found)
        total_criteria_count = sum(len(results) for results in analysis_results.values())
        hydraulic_validity = valid_criteria_count / total_criteria_count
        
        recommendations.append(f"⚠️ Hidrolik Geçerlilik: %{hydraulic_validity*100:.1f} ({valid_criteria_count}/{total_criteria_count} kriter)")

        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 30:
                recommendations.append(f"❌ {category} bölümü yetersiz (%{category_score:.1f})")
                missing_criteria = [name for name, result in results.items() if not result.found]
                if missing_criteria:
                    recommendations.append(f"  Eksik kriterler: {', '.join(missing_criteria)}")
            elif category_score < 70:
                recommendations.append(f"⚠️ {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"✅ {category} bölümü yeterli (%{category_score:.1f})")

        if scores["overall_percentage"] < 70:
            recommendations.append("\n🚨 GENEL ÖNERİLER:")
            recommendations.extend([
                "- Şema ISO 1219 standardına uyumlu hale getirilmelidir",
                "- Hidrolik semboller eksiksiz olmalıdır",
                "- Sistem bilgileri detaylandırılmalıdır",
                "- Basınç ve debi değerleri belirtilmelidir"
            ])

        return recommendations

    def analyze_circuit_diagram(self, file_path: str) -> Dict[str, Any]:
        """Main analysis function"""
        logger.info(f"Starting circuit diagram analysis for: {file_path}")
        file_ext = os.path.splitext(file_path)[1].lower()
        text = ""
        
        if file_ext == '.pdf':
            text = self.extract_text_from_pdf(file_path)
            self._last_extracted_text = text
        elif file_ext in ['.docx', '.doc']:
            text = self.extract_text_from_docx(file_path)
            self._last_extracted_text = text
        elif file_ext == '.txt':
            text = self.extract_text_from_txt(file_path)
            self._last_extracted_text = text
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            text = self.extract_text_from_image(file_path)
            self._last_extracted_text = text
        else:
            return {"error": f"Unsupported file format: {file_ext}"}
        
        if not text:
            return {"error": "Could not extract text from file"}

        logger.info(f"Extracted text length: {len(text)} characters")

        images = []
        if file_ext == '.pdf':
            images = self.extract_images_from_pdf(file_path)
        
        circuit_type, type_confidence = self.determine_circuit_type(text, images)

        analysis_results = {}
        for category in self.hydraulic_criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, images, category, circuit_type)

        scores = self.calculate_scores(analysis_results, circuit_type)
        extracted_values = self.extract_specific_values(text, circuit_type)
        recommendations = self.generate_recommendations(analysis_results, scores, circuit_type)

        report = {
            "analysis_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "file_info": {"file_path": file_path},
            "circuit_type": {"type": circuit_type, "confidence": round(type_confidence * 100, 2)},
            "extracted_values": extracted_values,
            "category_analyses": analysis_results,
            "scoring": scores,
            "recommendations": recommendations,
            "summary": {
                "total_score": scores["total_score"],
                "percentage": scores["overall_percentage"],
                "status": "PASS" if scores["overall_percentage"] >= 70 else "FAIL",
                "circuit_type": circuit_type.upper()
            }
        }

        return report

# ============================================
# HELPER FUNCTIONS
# ============================================
def validate_document_server(text, validation_keywords):
    """Server document validation - DB'den"""
    
    # DB'den critical_terms al
    critical_terms_data = validation_keywords.get('critical_terms', {})
    
    # Liste formatına dönüştür
    critical_terms = []
    for key, value in critical_terms_data.items():
        if key.startswith('category_') and isinstance(value, list):
            critical_terms.append(value)
    
    if not critical_terms:
        logger.warning("⚠️ Critical terms bulunamadı, validasyon atlanıyor")
        return True
    
    category_found = [
        any(re.search(rf"\b{term}\b", text, re.IGNORECASE) for term in category)
        for category in critical_terms
    ]
    
    valid_categories = sum(category_found)
    logger.info(f"Doküman validasyonu: {valid_categories}/{len(critical_terms)} kritik kategori bulundu")
    
    return valid_categories >= len(critical_terms)


def check_strong_keywords_first_pages(filepath, validation_keywords):
    """Check strong keywords in first pages - DB'den"""
    
    # DB'den strong keywords al
    strong_keywords = validation_keywords.get('strong_keywords', [])
    
    if not strong_keywords:
        logger.warning("⚠️ Strong keywords bulunamadı, validasyon atlanıyor")
        return True
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=300, first_page=1, last_page=1)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 3 -l tur+eng')
            all_text += text.lower() + " "
        
        found_keywords = [kw for kw in strong_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa kontrol: {len(found_keywords)} özgü kelime")
        return len(found_keywords) >= 1
        
    except Exception as e:
        logger.warning(f"İlk sayfa kontrol hatası: {e}")
        return False


def check_excluded_keywords_first_pages(filepath, validation_keywords):
    """Check excluded keywords in first pages - DB'den"""
    
    # DB'den excluded keywords al
    excluded_keywords = validation_keywords.get('excluded_keywords', [])
    
    if not excluded_keywords:
        logger.warning("⚠️ Excluded keywords bulunamadı, validasyon atlanıyor")
        return False
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=300, first_page=1, last_page=1)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 3 -l tur+eng')
            all_text += text.lower() + " "
        
        found_excluded = [kw for kw in excluded_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa excluded kontrol: {len(found_excluded)} istenmeyen kelime")
        return len(found_excluded) >= 1
        
    except Exception as e:
        logger.warning(f"İlk sayfa excluded kontrol hatası: {e}")
        return False


def get_conclusion_message_hydraulic(status, percentage):
    """Get conclusion message"""
    if status == "PASS":
        return f"Hidrolik devre şeması ISO 1219 standardına uygun ve teknik açıdan yeterlidir (%{percentage:.0f})"
    else:
        return f"Hidrolik devre şeması standartlara uygun değil, kapsamlı revizyon gerekli (%{percentage:.0f})"


def get_main_issues_hydraulic(analysis_result):
    """Get main issues"""
    issues = []
    
    for category, score_data in analysis_result['scoring']['category_scores'].items():
        if score_data['percentage'] < 50:
            issues.append(f"{category} kategorisinde ciddi eksiklikler")
    
    if not issues and analysis_result['scoring']['total_score'] < 50:
        issues = [
            "Hidrolik semboller ISO 1219 standardına uygun değil",
            "Basınç ve debi değerleri eksik veya hatalı",
            "Sistem bileşenleri tam tanımlanmamış",
            "Güvenlik elemanları eksik"
        ]
    
    return issues[:4]


# ============================================
# FLASK SERVICE
# ============================================
app = Flask(__name__)

# Database configuration (YENİ)
app.config['SQLALCHEMY_DATABASE_URI'] = Config.SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = Config.SQLALCHEMY_TRACK_MODIFICATIONS
app.config['SQLALCHEMY_ECHO'] = Config.SQLALCHEMY_ECHO

UPLOAD_FOLDER = 'temp_uploads_hydraulic'
ALLOWED_EXTENSIONS = {'pdf', 'jpg', 'jpeg', 'png', 'docx', 'doc', 'txt'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/api/hydraulic-control', methods=['POST'])
def analyze_hydraulic_control():
    """Hidrolik Devre Şeması analiz endpoint"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Create analyzer instance with app context
        analyzer = AdvancedCircuitAnalyzer(app=app)
        file_ext = os.path.splitext(filepath)[1].lower()

        # 3 AŞAMALI KONTROL (sadece PDF için)
        if file_ext == '.pdf':
            logger.info("Aşama 1: Hidrolik özgü kelime kontrolü...")
            if check_strong_keywords_first_pages(filepath, analyzer.validation_keywords):
                logger.info("✅ Aşama 1 geçti")
            else:
                logger.info("Aşama 2: Excluded kelime kontrolü...")
                if check_excluded_keywords_first_pages(filepath, analyzer.validation_keywords):
                    logger.info("❌ Excluded kelimeler bulundu")
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Invalid document type',
                        'message': 'Bu dosya hidrolik devre şeması değil'
                    }), 400
                else:
                    # AŞAMA 3
                    logger.info("Aşama 3: Tam doküman kontrolü...")
                    try:
                        with open(filepath, 'rb') as f:
                            pdf_reader = PyPDF2.PdfReader(f)
                            text = "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
                        
                        if not text or len(text.strip()) < 50 or not validate_document_server(text, analyzer.validation_keywords):
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({
                                'error': 'Invalid document type',
                                'message': 'Yüklediğiniz dosya hidrolik devre şeması değil!'
                            }), 400
                    except Exception as e:
                        logger.error(f"Aşama 3 hatası: {e}")
                        try:
                            os.remove(filepath)
                        except:
                            pass
                        return jsonify({'error': 'Analysis failed'}), 500

        elif file_ext in ['.docx', '.doc', '.txt']:
            # DOCX/TXT için sadece tam doküman kontrolü
            logger.info(f"DOCX/TXT dosyası için tam doküman kontrolü: {file_ext}")
            text = ""
            if file_ext in ['.docx', '.doc']:
                text = analyzer.extract_text_from_docx(filepath)
            elif file_ext == '.txt':
                text = analyzer.extract_text_from_txt(filepath)
            
            if not text or len(text.strip()) == 0:
                try:
                    os.remove(filepath)
                except:
                    pass
                return jsonify({
                    'error': 'Text extraction failed',
                    'message': 'Dosyadan yeterli metin çıkarılamadı'
                }), 400
            
            if not validate_document_server(text, analyzer.validation_keywords):
                try:
                    os.remove(filepath)
                except:
                    pass
                return jsonify({
                    'error': 'Invalid document type',
                    'message': 'Yüklediğiniz dosya hidrolik devre şeması değil!',
                    'details': {
                        'filename': filename,
                        'document_type': 'NOT_HIDROLIK_DEVRE_SEMASI'
                    }
                }), 400

        logger.info(f"Hidrolik devre şeması analizi yapılıyor: {filename}")
        analysis_result = analyzer.analyze_circuit_diagram(filepath)
        
        try:
            os.remove(filepath)
        except:
            pass

        if 'error' in analysis_result:
            return jsonify({'error': 'Analysis failed', 'message': analysis_result['error']}), 400

        overall_percentage = analysis_result['summary']['percentage']
        status = "PASS" if overall_percentage >= 70 else "FAIL"
        
        response_data = {
            'analysis_date': analysis_result.get('analysis_date'),
            'analysis_id': f"hydraulic_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            'category_scores': {},
            'extracted_values': analysis_result.get('extracted_values', {}),
            'file_type': 'HIDROLIK_DEVRE_SEMASI',
            'filename': filename,
            'overall_score': {
                'percentage': round(overall_percentage, 2),
                'total_points': analysis_result['scoring']['total_score'],
                'max_points': 100,
                'status': status,
                'status_tr': 'GEÇERLİ' if status == "PASS" else 'GEÇERSİZ'
            },
            'recommendations': analysis_result.get('recommendations', []),
            'summary': {
                'is_valid': status == "PASS",
                'conclusion': get_conclusion_message_hydraulic(status, overall_percentage),
                'main_issues': [] if status == "PASS" else get_main_issues_hydraulic(analysis_result)
            }
        }
        
        for category, score_data in analysis_result['scoring']['category_scores'].items():
            response_data['category_scores'][category] = {
                'score': score_data['normalized'],
                'max_score': score_data['max_weight'],
                'percentage': score_data['percentage'],
                'status': 'PASS' if score_data['percentage'] >= 70 else 'FAIL'
            }

        return jsonify({
            'success': True,
            'message': 'Hidrolik Devre Şeması başarıyla analiz edildi',
            'analysis_service': 'hydraulic_circuit_diagram',
            'data': response_data
        })

    except Exception as e:
        logger.error(f"API error: {str(e)}")
        return jsonify({'error': 'Server error', 'message': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check"""
    return jsonify({
        'status': 'healthy',
        'service': 'Hydraulic Circuit Diagram Analyzer API',
        'version': '1.0.0',
        'tesseract_available': tesseract_available,
        'upload_folder': UPLOAD_FOLDER,
        'max_file_size_mb': app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024),
        'supported_formats': list(ALLOWED_EXTENSIONS),
        'report_type': 'HIDROLIK_DEVRE_SEMASI'
    })


@app.route('/', methods=['GET'])
def index():
    """Index page"""
    return jsonify({
        'service': 'Hydraulic Circuit Diagram Analyzer API',
        'version': '1.0.0',
        'description': 'Hidrolik Devre Şemalarını analiz eden REST API servisi',
        'endpoints': {
            'POST /api/hydraulic-control': 'Hidrolik devre şeması analizi',
            'GET /api/health': 'Health check',
            'GET /': 'API info'
        },
        'usage': {
            'upload_format': 'multipart/form-data',
            'file_field': 'file',
            'supported_types': list(ALLOWED_EXTENSIONS),
            'max_size_mb': app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024)
        }
    })


# ============================================
# DATABASE INITIALIZATION
# ============================================
with app.app_context():
    db.init_app(app)


# ============================================
# APPLICATION ENTRY POINT (Azure-Friendly)
# ============================================
if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("Hidrolik Devre Şeması Analiz Servisi")
    logger.info("=" * 60)
    logger.info(f"📁 Upload klasörü: {UPLOAD_FOLDER}")
    logger.info(f"📊 Desteklenen formatlar: {', '.join(ALLOWED_EXTENSIONS)}")
    logger.info(f"📏 Maksimum dosya boyutu: {app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024)} MB")
    logger.info(f"🔧 Tesseract OCR: {'Kurulu' if tesseract_available else 'Kurulu değil'}")
    logger.info("")
    logger.info("🔗 API Endpoints:")
    logger.info("  POST /api/hydraulic-control - Hidrolik devre şeması analizi")
    logger.info("  GET  /api/health           - Health check")
    logger.info("  GET  /                     - API info")
    logger.info("=" * 60)
    
    port = int(os.environ.get('PORT', 8011))
    logger.info(f"🚀 Servis başlatılıyor - Port: {port}")
    logger.info("=" * 60)
    
    app.run(host='0.0.0.0', port=port, debug=False)