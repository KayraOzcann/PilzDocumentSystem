import re
import os
import json
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
import pandas as pd
from dataclasses import dataclass, asdict
import logging

# OCR için import'lar
import pytesseract
import cv2
import numpy as np
from PIL import Image
import pdf2image

try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class HRCCriteria:
    """HRC Kuvvet-Basınç Ölçüm Raporu kriterleri veri sınıfı"""
    genel_bilgiler: Dict[str, Any]
    test_kosullari_senaryo: Dict[str, Any]
    olcum_noktalari_metodoloji: Dict[str, Any]
    kuvvet_basinc_sonuclari: Dict[str, Any]
    sinir_degerleri_karsilastirma: Dict[str, Any]
    risk_degerlendirmesi_sonuc: Dict[str, Any]
    oneriler_onlemler: Dict[str, Any]
    ekler_kalibrasyon_belgeleri: Dict[str, Any]

@dataclass
class HRCAnalysisResult:
    """HRC analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

class HRCReportAnalyzer:
    """HRC Kuvvet-Basınç Ölçüm Raporu analiz sınıfı"""
    
    def __init__(self):
        logger.info("HRC Kuvvet-Basınç Ölçüm Raporu analiz sistemi başlatılıyor...")
        logger.info("✅ OCR sistemi aktif - Tüm dosyalar OCR ile işlenmektedir")
        
        self.criteria_weights = {
            "Genel Bilgiler": 10,
            "Test Koşulları ve Senaryo Tanımı": 10,
            "Ölçüm Noktaları ve Metodoloji": 15,
            "Kuvvet ve Basınç Ölçüm Sonuçları": 25,
            "Sınır Değerlerle Karşılaştırma": 20,
            "Risk Değerlendirmesi ve Sonuç": 10,
            "Öneriler ve Önlemler": 5,
            "Ekler ve Kalibrasyon Belgeleri": 5
        }
        
        self.criteria_details = {
            "Genel Bilgiler": {
                "test_tarihi": {"pattern": r"(?i)(?:test|ölçüm|analiz|measurement|analysis|tarih|date)[\s\W]*[:=]?\s*(\d{1,2}[\./\-]\d{1,2}[\./\-]\d{2,4})|(\d{1,2}[\./\-]\d{1,2}[\./\-]\d{4})(?:\s*(?:tarih|date))?", "weight": 2},
                "test_yapan_kurum": {"pattern": r"(?i)(?:test\s*yapan|tested\s*by|ölçüm\s*yapan|measured\s*by|kurum|institution|organization|şirket|company|firma|sorumlu|responsible)[\s\W]*[:=]?\s*([A-ZÇĞİÖŞÜa-züçğıöşü][A-Za-züçğıöşüÇĞİÖŞÜ\s\.\&\-]{3,50})", "weight": 2},
                "robot_modeli_seri": {"pattern": r"(?i)(?:robot\s*model|robot\s*tip|model|seri\s*no|serial\s*number|s\/n|robot\s*type|tip|robot\s*id)[\s\W]*[:=]?\s*([A-Z0-9\-]+[A-Z0-9\-\s]*)|([A-Z]{2,4}[\-\s]*[0-9]{1,5}[A-Z0-9]*)", "weight": 3},
                "uygulama_istasyon": {"pattern": r"(?i)(?:uygulama|application|istasyon|station|workstation|iş\s*istasyon|work\s*station|test\s*setup|test\s*düzen|çalışma\s*alan|work\s*area)[\s\W]*[:=]?\s*([A-ZÇĞİÖŞÜa-züçğıöşü][A-Za-züçğıöşüÇĞİÖŞÜ0-9\s\.\-]{3,50})", "weight": 3}
            },
            "Test Koşulları ve Senaryo Tanımı": {
                "robot_hiz_durum": {"pattern": r"(?i)(?:hız|speed|velocity|hızlı|durum|position|pose|duruş|stance|hareket|motion)[\s\W]*[:=]?\s*([0-9.,]+\s*(?:mm\/s|m\/s|%|derece|degree|°|rpm)|[A-ZÇĞİÖŞÜa-züçğıöşü][A-Za-züçğıöşüÇĞİÖŞÜ\s]{3,30})", "weight": 3},
                "calisma_modu": {"pattern": r"(?i)(?:çalışma\s*mod|work\s*mode|operation\s*mode|manuel|manual|otomatik|automatic|işbirlik|collaborative|cobot|hrc|interaction|mod|mode)", "weight": 2},
                "temas_bolgeleri": {"pattern": r"(?i)(?:temas\s*bölge|contact\s*area|contact\s*region|vücut\s*bölge|body\s*region|bölge|region|area|kol|arm|el|hand|gövde|body|torso|baş|head|bacak|leg|finger|parmak)", "weight": 3},
                "cevresel_sartlar": {"pattern": r"(?i)(?:sıcaklık|temperature|ortam|environment|çevre|ambient|celsius|nem|humidity|koşul|condition|şart)[\s\W]*[:=]?\s*([0-9.,]+\s*(?:°c|c|%|derece)|[A-ZÇĞİÖŞÜa-züçğıöşü][A-Za-züçğıöşüÇĞİÖŞÜ\s]{3,30})", "weight": 2}
            },
            "Ölçüm Noktaları ve Metodoloji": {
                "vucut_bolgeleri_iso15066": {"pattern": r"(?i)(?:iso\/ts\s*15066|iso\s*15066|ts\s*15066|vücut\s*bölge|body\s*region|tablo|table|kafa|head|yüz|face|boyun|neck|sırt|back|göğüs|chest|karın|abdomen|pelvis|kol|arm|dirsek|elbow|ön\s*kol|forearm|el|hand|parmak|finger|15066|skull|forehead|temple|masticatory|muscle)", "weight": 5},
                "olcum_yontemi": {"pattern": r"(?i)(?:ölçüm\s*yöntem|measurement\s*method|test\s*method|yöntem|method|statik|static|dinamik|dynamic|temas|contact|serbest\s*hareket|free\s*motion|engel|obstacle|barrier|prosedür|procedure|ölçüm\s*nokta|measurement\s*point|nokta\s*\d+|point\s*\d+|film|llw|max\s*pressure|mpa)", "weight": 5},
                "tekrar_sayisi_tutarlilik": {"pattern": r"(?i)(?:tekrar|repeat|iteration|test\s*sayı|test\s*count|n\s*=|n=|örnek\s*sayı|sample\s*count)[\s\W]*[:=]?\s*(\d{1,3})|(?:tutarlılık|consistency|repeatability|reproducibility|tekrarlama)|(?:a1|a2|b1|b2|c1|c2)[\s\W]*[=:]\s*([0-9.,]+)|(?:kuvvet\s*ölçüm|force\s*measurement)[\s\S]{0,50}(?:gelişim|development|progress)|(\d{1,2})\s*(?:nokta|point|ölçüm|measurement)", "weight": 5}
            },
            "Kuvvet ve Basınç Ölçüm Sonuçları": {
                "maksimum_temas_kuvveti": {"pattern": r"(?i)(?:maksimum\s*kuvvet|maximum\s*force|max\s*force|fmax|f\s*max|kuvvet\s*değer|force\s*value|en\s*büyük\s*kuvvet)[\s\W]*[:=]?\s*([0-9.,]+)\s*(?:n|newton|kn)|(?:f[\s\W]*max|fr[\s\W]*max|fs[\s\W]*max)[\s\W]*([0-9.,]+)|(\d{1,3})\s*(?=\s*\-?\s*\d{1,2}\s)|(?:maximum\s*permissible\s*kuvvet|maximum\s*permissible\s*force|izin\s*verilen\s*kuvvet)[\s\S]{0,50}([0-9.,]+)|(?:quasi[\-\s]*static|transient)[\s\S]{0,100}(?:kuvvet|force)[\s\S]{0,50}([0-9.,]+)", "weight": 7},
                "temas_suresi": {"pattern": r"(?i)(?:temas\s*süre|contact\s*time|contact\s*duration|süre|duration|time|t\s*=|t=|zaman|contact)[\s\W]*[:=]?\s*([0-9.,]+)\s*(?:ms|millisecond|s|second|saniye|msn|sn)|(\d{1,3})\s*(?:ms|s|saniye|sn)\b|(?:süre|duration|time)[\s\W]*(\d{1,3})|(?:transient|quasi[\-\s]*static)[\s\S]{0,50}(?:süre|time|duration)[\s\S]{0,30}([0-9.,]+)", "weight": 6},
                "basinc_degeri": {"pattern": r"(?i)(?:basınç\s*ts|pressure\s*ts|basınç|pressure|press|baskı)[\s\W]*[:=]?\s*([0-9.,]+)\s*(?:n\/cm²|n\/cm2|pa|kpa|mpa|bar|pascal)|(?:ps[\s\W]*max|p[\s\W]*max|basınç[\s\W]*max)[\s\W]*([0-9.,]+)|(\d{1,3})\s*(?:n\/cm²|n\/cm2|pa|bar)|(?:ts\s*15066)[\s\S]{0,50}(\d{1,3})\s*(?:n\/cm²|n\/cm2)|(?:maximum\s*permissible\s*basınç|maximum\s*permissible\s*pressure|izin\s*verilen\s*basınç)[\s\S]{0,50}([0-9.,]+)|(?:quasi[\-\s]*static|transient)[\s\S]{0,100}(?:basınç|pressure)[\s\S]{0,50}([0-9.,]+)", "weight": 6},
                "grafiksel_gosterim": {"pattern": r"(?i)(?:grafik|graph|chart|eğri|curve|kuvvet[\-\s]*zaman|force[\-\s]*time|plot|çizim|drawing|şekil\s*\d+|figure\s*\d+|diyagram|diagram|görsel|visual|resim\s*\d+|image\s*\d+|fig\s*\d+|şek\s*\d+|tablo\s*\d+|table\s*\d+|sonuç|result|ölçüm\s*sonuç|measurement\s*result)", "weight": 6}
            },
            "Sınır Değerlerle Karşılaştırma": {
                "iso15066_sinir_karsilastirma": {"pattern": r"(?i)(?:iso\/ts\s*15066|iso\s*15066|ts\s*15066|15066)[\s\S]*?(?:sınır|limit|threshold|eşik|karşılaştırma|comparison|compare|kıyas|standart|uygun|compliant)", "weight": 8},
                "vucut_bolgesi_limitleri": {"pattern": r"(?i)(?:izin\s*verilen|allowed|permitted|limit|sınır|maksimum\s*izin|maximum\s*allowed|kabul\s*edilebilir|acceptable)[\s\S]*?(?:kuvvet|force|basınç|pressure)", "weight": 6},
                "asim_risk_isaret": {"pattern": r"(?i)(?:aşım|exceed|over|fazla|limit\s*aş|limit\s*over|risk|tehlike|hazard|warning|uyarı|alert|güvenli\s*değil|not\s*safe|tehlikeli|dangerous)", "weight": 6}
            },
            "Risk Değerlendirmesi ve Sonuç": {
                "risk_seviye_analizi": {"pattern": r"(?i)(?:risk\s*analiz|risk\s*analysis|risk\s*assessment|risk\s*değerlendirme|risk|seviye|level|kategori|category|düşük|low|orta|medium|yüksek|high|değerlendirme|assessment)", "weight": 4},
                "risk_kabul_edilebilir": {"pattern": r"(?i)(?:kabul\s*edilebilir|acceptable|accept|uygun|suitable|güvenli|safe|güvenlik|safety|onay|approve|red|reject|kabul|uygunluk|compliance)", "weight": 3},
                "gereken_onlemler": {"pattern": r"(?i)(?:önlem|measure|action|tedbir|hız\s*sınır|speed\s*limit|güvenlik\s*sensör|safety\s*sensor|uç\s*efektör|end\s*effector|koruma|protection|emniyet|security)", "weight": 3}
            },
            "Öneriler ve Önlemler": {
                "emniyet_stratejisi": {"pattern": r"(?i)(?:emniyet\s*stratejisi|safety\s*strategy|strateji|güvenlik\s*stratejisi|hız\s*sınır|speed\s*limit|kuvvet\s*sınır|force\s*limit|yastıklama|padding|koruyucu|protective|eşik\s*değer|threshold\s*value|limit\s*değer|uygun\s*değer|yeşil|green|renk|color|renklendir|colored)", "weight": 2},
                "operatör_egitimi": {"pattern": r"(?i)(?:operatör|operator|eğitim|training|bilgilendirme|information|uyarı|warning|not|note|kullanıcı|user|personel|personnel|kabul\s*edilebilir|acceptable|uygun|appropriate|yeterli|sufficient)", "weight": 2},
                "periyodik_test": {"pattern": r"(?i)(?:periyodik|periodic|tekrarlayan|repeated|test\s*tekrar|test\s*repeat|düzenli|regular|planlı|scheduled|rutin|routine|bakım|maintenance|ölçüm\s*tablo|measurement\s*table|tablo|table|değerlendirme|evaluation)", "weight": 1}
            },
            "Ekler ve Kalibrasyon Belgeleri": {
                "kalibrasyon_sertifika": {"pattern": r"(?i)(?:kalibrasyon|calibration|sertifika|certificate|cert|belge|document|iso\s*17025|akreditasyon|accreditation|onay|approval|doğrulama|verification)", "weight": 2},
                "fotograf_video": {"pattern": r"(?i)(?:fotoğraf|photo|photograph|video|resim|image|picture|görsel|visual|kayıt|record|çekim|shot|dokümantasyon|documentation)", "weight": 2},
                "test_prosedur_referans": {"pattern": r"(?i)(?:prosedür|procedure|referans|reference|standart|standard|kılavuz|guide|manual|dokümantasyon|documentation|kaynak|source|metod|method)", "weight": 1}
            }
        }
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni Türkçe'ye çevir - şimdilik devre dışı"""
        if source_lang != 'tr':
            logger.info(f"Tespit edilen dil: {source_lang.upper()} - Çeviri yapılmıyor, orijinal metin kullanılıyor")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosya formatına göre OCR ile metin çıkarma"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Tüm dosya türleri için OCR kullan
        if file_ext == '.pdf':
            return self.extract_text_with_ocr(file_path)
        elif file_ext in ['.docx', '.doc']:
            # DOCX'i önce standart yöntemle dene, başarısızsa OCR'a geç
            try:
                text = self.extract_text_from_docx(file_path)
                if len(text.strip()) > 50:
                    return text
                else:
                    logger.info("DOCX standart yöntem yetersiz, OCR deneniyor...")
                    return self.extract_text_with_ocr(file_path)
            except:
                logger.info("DOCX okuma başarısız, OCR deneniyor...")
                return self.extract_text_with_ocr(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self.extract_text_with_ocr(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatı: {file_ext}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den standart PyPDF2 ile metin çıkarma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF analizi başlatılıyor - Toplam sayfa: {total_pages}")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = page_text.replace('|', ' ')
                    page_text = page_text.strip()
                    all_text += page_text + "\n"
                
                # Metni temizle
                all_text = all_text.replace('—', '-')
                all_text = all_text.replace('"', '"').replace('"', '"')
                all_text = all_text.replace('´', "'")
                all_text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]+', ' ', all_text)
                all_text = all_text.strip()
                
                logger.info(f"✅ PDF analizi tamamlandı: {len(all_text):,} karakter")
                return all_text
                
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_with_ocr(self, file_path: str) -> str:
        """OCR ile metin çıkarma"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                # PDF'i görüntülere dönüştür
                pages = pdf2image.convert_from_path(file_path, dpi=200)
                all_text = ""
                
                for i, page in enumerate(pages):
                    # PIL'den OpenCV formatına
                    opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
                    
                    # Görüntü ön işleme
                    gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
                    processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
                    
                    # OCR uygula
                    text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
                    all_text += text + "\n"
                    
                    logger.info(f"PDF sayfa {i+1}/{len(pages)} OCR tamamlandı")
                
                logger.info(f"✅ PDF OCR tamamlandı: {len(all_text):,} karakter")
                return all_text.strip()
                
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                # Görüntü dosyası
                image = cv2.imread(file_path)
                if image is None:
                    logger.error(f"Görüntü yüklenemedi: {file_path}")
                    return ""
                
                # Görüntü ön işleme
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                _, processed = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)
                
                # OCR uygula
                text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
                
                logger.info(f"✅ Görüntü OCR tamamlandı: {len(text):,} karakter")
                return text.strip()
            
            elif file_ext in ['.docx', '.doc']:
                # DOCX/DOC dosyalarını görüntüye dönüştürüp OCR uygula
                logger.warning(f"DOCX/DOC dosyası OCR ile işlenemez: {file_ext}")
                return ""
            
            else:
                logger.warning(f"OCR için desteklenmeyen format: {file_ext}")
                return ""
                
        except Exception as e:
            logger.error(f"OCR hatası: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            logger.info(f"DOCX'den {len(text)} karakter metin çıkarıldı")
            return text
            
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasından metin çıkarma"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text:
                encodings = ['cp1254', 'iso-8859-9', 'latin1']
                for encoding in encodings:
                    try:
                        with open(txt_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        if text:
                            break
                    except:
                        continue
            
            logger.info(f"TXT'den {len(text)} karakter metin çıkarıldı")
            return text.strip()
            
        except Exception as e:
            logger.error(f"TXT metin çıkarma hatası: {e}")
            return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, HRCAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                # Weight'e göre scoring algoritması
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
                    
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = HRCAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={
                    "pattern_used": pattern,
                    "matches_count": len(matches) if matches else 0
                }
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, HRCAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """HRC raporuna özgü spesifik değerleri çıkar"""
        values = {
            "robot_modeli": "Bulunamadı",
            "test_tarihi": "Bulunamadı",
            "olcum_cihazi": "Bulunamadı"
        }
        
        # Robot Modeli
        robot_patterns = [
            r"(?i)(?:Robot\s*tip[i]?|Robot\s*type|Robot\s*model)\s*[|\s]\s*([A-Z0-9\-]+(?:[A-Z0-9\-]*))(?=\s|$|\n)",
            r"(?i)(?:Robot\s*Model|Model)\s*[:=]\s*([A-Z0-9]+\-[A-Z0-9]+[A-Z]*)",
            r"([A-Z]{2,4}\-[0-9]{1,3}[A-Z]*)\b",
            r"(?i)(?:model|tip|type).*?[|\s]{2,}([A-Z0-9]+\-[A-Z0-9]+[A-Z]*)"
        ]
        
        for pattern in robot_patterns:
            match = re.search(pattern, text)
            if match:
                values["robot_modeli"] = match.group(1).strip()
                break
        
        # Test Tarihi
        date_patterns = [
            r"(?i)(?:Test\s*Tarih|Test\s*Date|Ölçüm\s*Tarih)\s*[:=]\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})",
            r"(?i)(?:Analiz\s*Tarih|Analysis\s*Date)\s*[:=]\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})",
            r"(\d{1,2}[./]\d{1,2}[./]\d{4})\s*(?:tarih|date)"
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                values["test_tarihi"] = match.group(1).strip()
                break

        olcum_cihazi_patterns = [
            r"(?i)(?:ölçüm\s*cihaz|measurement\s*device)[\s\W]*[:=]?\s*([A-Z][A-Za-z0-9\s\-]{4,20})(?=\s*\n|\s*$|\s*[|])",
            # Generic format: Büyük harf kelime + küçük harf kelime + büyük harf kelime + kod, sonra sayı dizisinden önce dur
            r"([A-Z]{3,8}\s+[a-z]{2,6}\s+[A-Z][a-z]+\s+[A-Z0-9]+)(?=\s+\d{4}|\s*\n|\s*$)",
            # Daha basit: Harfle başlayıp kodla biten, sonra versiyon numarasından önce dur
            r"([A-Z][A-Za-z\s]+[A-Z0-9]+)(?=\s+[\d\.]+\s+\d{4})"
        ]

        for pattern in olcum_cihazi_patterns:
            match = re.search(pattern, text)
            if match:
                found_value = match.group(1).strip()
                # Makul uzunlukta ve Set kelimesi içeren cihaz isimleri
                if 10 <= len(found_value) <= 25:
                    values["olcum_cihazi"] = found_value
                    break
                
        return values

    def validate_hrc_document(self, text: str) -> bool:
        """Dokümanın HRC kuvvet-basınç ölçüm raporu olup olmadığını kontrol et"""
        
        # HRC raporlarında MUTLAKA olması gereken kritik kelimeler
        critical_hrc_terms = [
            # HRC/Cobot temel terimleri (en az 1 tane olmalı)
            ["hrc", "collaborative", "işbirlik", "cobot", "kolaboratif", "human robot collaboration"],
            
            # Kuvvet/Basınç ölçüm terimleri (en az 1 tane olmalı)  
            ["kuvvet", "force", "basınç", "pressure", "temas", "contact", "newton"],
            
            # ISO 15066 standardı (mutlaka olmalı)
            ["iso 15066", "iso/ts 15066", "ts 15066", "15066"],
            
            # Vücut bölgeleri (HRC raporlarına özgü, en az 1 tane olmalı)
            ["vücut", "body", "kol", "arm", "el", "hand", "baş", "head", "gövde", "torso", "boyun", "neck"]
        ]
        
        # Her kategori için kontrol
        category_found = []
        
        for i, category in enumerate(critical_hrc_terms):
            found_in_category = False
            for term in category:
                if re.search(rf"\b{term}\b", text, re.IGNORECASE):
                    found_in_category = True
                    logger.info(f"HRC Kategori {i+1} bulundu: '{term}'")
                    break
            category_found.append(found_in_category)
        
        # Tüm kategorilerden en az bir terim bulunmalı
        valid_categories = sum(category_found)
        
        logger.info(f"HRC doküman validasyonu: {valid_categories}/4 kritik kategori bulundu")
        
        # 4 kategorinin tamamında terim bulunmalı (daha sıkı kontrol)
        return valid_categories >= 4

    def generate_improvement_actions(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Dinamik iyileştirme önerileri oluştur"""
        actions = []
        
        # Kategorileri puana göre sırala (düşükten yükseğe)
        sorted_categories = sorted(
            scores["category_scores"].items(), 
            key=lambda x: x[1]["percentage"]
        )
        
        category_actions = {
            "Genel Bilgiler": [
                "Test tarihini, test yapan kurum/kişi bilgilerini netleştiriniz",
                "Test edilen robot modeli ve seri numarasını belirtiniz",
                "Uygulama/istasyon tanımını detaylandırınız",
                "Kullanılan test ekipmanı ve kalibrasyon durumunu ekleyiniz"
            ],
            "Test Koşulları ve Senaryo Tanımı": [
                "Robot hızı, duruşu, uç efektör tipini belirtiniz",
                "Çalışma modunu (manuel, otomatik, işbirlikçi) tanımlayınız",
                "Test edilen temas bölgelerini (kol, el, gövde) listeleyiniz",
                "Çevresel şartları (sıcaklık, ortam) kaydediniz"
            ],
            "Ölçüm Noktaları ve Metodoloji": [
                "ISO/TS 15066 tablosuna göre vücut bölgelerini belirtiniz",
                "Ölçüm yöntemini (statik/dinamik temas) açıklayınız",
                "Tekrar sayısı ve yöntem tutarlılığını ekleyiniz",
                "Test metodolojisini detaylandırınız"
            ],
            "Kuvvet ve Basınç Ölçüm Sonuçları": [
                "Maksimum temas kuvvetini [N] cinsinden kaydediniz",
                "Temas süresini [ms] cinsinden belirtiniz",
                "Basınç değerlerini [N/cm²] cinsinden ekleyiniz",
                "Grafiksel gösterimleri (kuvvet-zaman eğrisi) oluşturunuz"
            ],
            "Sınır Değerlerle Karşılaştırma": [
                "ISO/TS 15066 sınır değerleri ile karşılaştırma yapınız",
                "Vücut bölgeleri için izin verilen limitleri belirtiniz",
                "Aşımlar varsa işaretleyiniz ve risk değerlendirmesi ekleyiniz",
                "Standart uygunluk durumunu netleştiriniz"
            ],
            "Risk Değerlendirmesi ve Sonuç": [
                "Ölçümlere dayalı risk seviyesi analizini yapınız",
                "Risk kabul edilebilirlik durumunu belirtiniz",
                "Gereken önlemleri (hız sınırı, güvenlik sensörü) listeleyiniz",
                "Sonuç değerlendirmesini detaylandırınız"
            ],
            "Öneriler ve Önlemler": [
                "Emniyet stratejisini (hız & kuvvet sınırlaması) tanımlayınız",
                "Operatör bilgilendirmesi ve eğitim notlarını ekleyiniz",
                "Periyodik test önerilerini belirtiniz",
                "Ek güvenlik önlemlerini listeleyiniz"
            ],
            "Ekler ve Kalibrasyon Belgeleri": [
                "Kalibrasyon sertifikalarını ekleyiniz",
                "Fotoğraf ve video kayıtlarını dahil ediniz",
                "Test prosedürü referanslarını belirtiniz",
                "Destekleyici dokümanları iliştirin"
            ]
        }
        
        # En düşük puanlı 5 kategori için öneriler ekle
        for category, score_data in sorted_categories[:5]:
            if score_data["percentage"] < 70:
                actions.extend(category_actions.get(category, []))
        
        # Genel öneriler
        if scores["percentage"] < 50:
            actions.insert(0, "ÖNCE: HRC rapor formatını ve yapısını yeniden gözden geçiriniz")
        
        return actions
    
    def generate_recommendations(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Öneriler oluştur"""
        recommendations = []
        
        total_percentage = scores["percentage"]
        
        if total_percentage >= 70:
            recommendations.append(f"✅ HRC Kuvvet-Basınç Raporu GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ HRC Kuvvet-Basınç Raporu GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 40:
                recommendations.append(f"🔴 {category} bölümü yetersiz (%{category_score:.1f})")
                missing_items = [name for name, result in results.items() if not result.found]
                if missing_items:
                    recommendations.append(f"   Eksik: {', '.join(missing_items[:3])}")
            elif category_score < 70:
                recommendations.append(f"🟡 {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"🟢 {category} bölümü yeterli (%{category_score:.1f})")
        
        return recommendations

    def analyze_hrc_report(self, file_path: str) -> Dict[str, Any]:
        """Ana HRC Kuvvet-Basınç Ölçüm Raporu analiz fonksiyonu"""
        logger.info("HRC Kuvvet-Basınç Ölçüm Raporu analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        if detected_lang != 'tr':
            logger.info(f"{detected_lang.upper()} dilinden Türkçe'ye çeviriliyor...")
            text = self.translate_to_turkish(text, detected_lang)
        
        # Doküman türü kontrolü
        if not self.validate_hrc_document(text):
            return {
                "error": "YANLIŞ DOKÜMAN: Bu dosya HRC kuvvet-basınç ölçüm raporu değil!",
                "document_type": "UNKNOWN",
                "suggestion": "Lütfen geçerli bir HRC kuvvet-basınç ölçüm raporu yükleyiniz."
            }
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        recommendations = self.generate_recommendations(analysis_results, scores)
        improvement_actions = self.generate_improvement_actions(analysis_results, scores)
        
        final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "iyilestirme_eylemleri": improvement_actions,
            "ozet": {
                "toplam_puan": scores["total_score"],
                "yuzde": scores["percentage"],
                "durum": final_status,
                "rapor_tipi": "HRC_KUVVET_BASINC_RAPORU"
            }
        }
        
        return report

def main():
    """Ana fonksiyon"""
    analyzer = HRCReportAnalyzer()

    # Test için örnek dosya yolu
    file_path = r"C:\Users\nuvo_teknik_2\Desktop\PILZ DOCUMENTS\2.7 HRC\C23.234 - SD Conta Robotu HRC Ölçümü Raporu .pdf"


    if not os.path.exists(file_path):
        print(f"❌ Dosya bulunamadı: {file_path}")
        return
    
    print("🤖 HRC Kuvvet-Basınç Ölçüm Raporu Analizi Başlatılıyor...")
    print("=" * 60)
    print("✅ OCR sistemi aktif - Tüm dosyalar OCR ile işlenmektedir")
    
    report = analyzer.analyze_hrc_report(file_path)
    
    if "error" in report:
        print(f"❌ Hata: {report['error']}")
        return
    
    print("\n📊 ANALİZ SONUÇLARI")
    print("=" * 60)
    
    print(f"📅 Analiz Tarihi: {report['analiz_tarihi']}")
    print(f"📄 Dosya Tipi: {report['dosya_bilgisi']['file_type'].upper()}")
    print(f"🔍 Tespit Edilen Dil: {report['dosya_bilgisi']['detected_language'].upper()}")
    print(f"📋 Toplam Puan: {report['ozet']['toplam_puan']}/100")
    print(f"📈 Yüzde: %{report['ozet']['yuzde']}")
    print(f"🎯 Durum: {report['ozet']['durum']}")
    print(f"📄 Rapor Tipi: {report['ozet']['rapor_tipi']}")
    
    print("\n🤖 ÖNEMLİ ÇIKARILAN DEĞERLER")
    print("-" * 40)
    for key, value in report['cikarilan_degerler'].items():
        display_name = {
            "robot_modeli": "Robot Modeli",
            "test_tarihi": "Test Tarihi",
            "olcum_cihazi": "Ölçüm Cihazı"
        }.get(key, key.replace('_', ' ').title())
        print(f"{display_name}: {value}")
    
    print("\n📊 KATEGORİ PUANLARI")
    print("-" * 40)
    for category, score_data in report['puanlama']['category_scores'].items():
        print(f"{category}: {score_data['normalized']}/{score_data['max_weight']} (%{score_data['percentage']:.1f})")
    
    print("\n💡 ÖNERİLER VE DEĞERLENDİRME")
    print("-" * 40)
    for recommendation in report['oneriler']:
        print(recommendation)
    
    print("\n📋 GENEL DEĞERLENDİRME")
    print("=" * 60)
    
    if report['ozet']['yuzde'] >= 70:
        print("✅ SONUÇ: GEÇERLİ")
        print(f"🌟 Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: HRC kuvvet-basınç ölçüm raporu EN ISO 10218 ve ISO/TS 15066 standartlarına uygun.")
    else:
        print("❌ SONUÇ: GEÇERSİZ")
        print(f"⚠️ Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: HRC raporu minimum gereklilikleri sağlamamaktadır.")
        
        print("\n⚠️ EKSİK GEREKLİLİKLER:")
        for category, results in report['kategori_analizleri'].items():
            missing_items = []
            for criterion, result in results.items():
                if not result.found:
                    missing_items.append(criterion)
            
            if missing_items:
                print(f"\n🔍 {category}:")
                for item in missing_items:
                    readable_name = {
                        "test_tarihi": "Test Tarihi",
                        "test_yapan_kurum": "Test Yapan Kurum",
                        "robot_modeli_seri": "Robot Modeli/Seri",
                        "uygulama_istasyon": "Uygulama/İstasyon Tanımı",
                        "robot_hiz_durum": "Robot Hızı/Duruşu",
                        "calisma_modu": "Çalışma Modu",
                        "temas_bolgeleri": "Temas Bölgeleri",
                        "cevresel_sartlar": "Çevresel Şartlar",
                        "vucut_bolgeleri_iso15066": "Vücut Bölgeleri (ISO 15066)",
                        "olcum_yontemi": "Ölçüm Yöntemi",
                        "tekrar_sayisi_tutarlilik": "Tekrar Sayısı/Tutarlılık",
                        "maksimum_temas_kuvveti": "Maksimum Temas Kuvveti",
                        "temas_suresi": "Temas Süresi",
                        "basinc_degeri": "Basınç Değeri",
                        "grafiksel_gosterim": "Grafiksel Gösterim",
                        "iso15066_sinir_karsilastirma": "ISO 15066 Sınır Karşılaştırması",
                        "vucut_bolgesi_limitleri": "Vücut Bölgesi Limitleri",
                        "asim_risk_isaret": "Aşım/Risk İşareti",
                        "risk_seviye_analizi": "Risk Seviye Analizi",
                        "risk_kabul_edilebilir": "Risk Kabul Edilebilirlik",
                        "gereken_onlemler": "Gereken Önlemler",
                        "emniyet_stratejisi": "Emniyet Stratejisi",
                        "operatör_egitimi": "Operatör Eğitimi",
                        "periyodik_test": "Periyodik Test Önerisi",
                        "kalibrasyon_sertifika": "Kalibrasyon Sertifikaları",
                        "fotograf_video": "Fotoğraf/Video Kayıtları",
                        "test_prosedur_referans": "Test Prosedürü Referansları"
                    }.get(item, item.replace('_', ' ').title())
                    print(f"   ❌ {readable_name}")
        
        print("\n📌 YAPILMASI GEREKENLER:")
        if "iyilestirme_eylemleri" in report:
            for i, action in enumerate(report['iyilestirme_eylemleri'], 1):
                print(f"{i}. {action}")

if __name__ == "__main__":
    main()
