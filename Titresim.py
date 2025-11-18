import os
import json
from datetime import datetime, timedelta
import re
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
import pandas as pd
from dataclasses import dataclass, asdict
import logging

try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class VibrationCriteria:
    """Mekanik Titreşim Ölçüm Raporu kriterleri veri sınıfı"""
    rapor_kimlik_bilgileri: Dict[str, Any]
    olcum_ortam_makine_calisan: Dict[str, Any]
    olcum_cihazi_kalibrasyon: Dict[str, Any]
    titresim_turleri_yasal: Dict[str, Any]
    olcum_metodolojisi_standartlar: Dict[str, Any]
    olcum_sonuclari_analizler: Dict[str, Any]
    degerlendirme_yorum_onlemler: Dict[str, Any]
    ekler_gorseller: Dict[str, Any]

@dataclass
class VibrationAnalysisResult:
    """Mekanik Titreşim Ölçüm Raporu analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

class VibrationReportAnalyzer:
    """Mekanik Titreşim Ölçüm Raporu analiz sınıfı"""
    
    def __init__(self):
        logger.info("Mekanik Titreşim Ölçüm Raporu analiz sistemi başlatılıyor...")
        
        self.criteria_weights = {
            "Rapor Kimlik Bilgileri": 10,
            "Ölçüm Ortam, Makine ve Çalışan Bilgileri": 10,
            "Ölçüm Cihazı ve Kalibrasyon Bilgileri": 10,
            "Titreşim Türleri ve Yasal Değerler": 15,
            "Ölçüm Metodolojisi ve Uygulanan Standartlar": 10,
            "Ölçüm Sonuçları ve Analizler": 20,
            "Değerlendirme, Yorum ve Önlemler": 20,
            "Ekler ve Görseller": 5
        }
        
        self.criteria_details = {
            "Rapor Kimlik Bilgileri": {
                "rapor_numarasi": {"pattern": r"(?:Rapor\s*No|Report\s*No|Rapor\s*Numaras[ıi]|Report\s*Number|Döküman\s*No|Document\s*No|Belge\s*No|Test\s*No)", "weight": 2},
                "rapor_tarihi": {"pattern": r"(?:Rapor\s*Tarih|Report\s*Date|Düzenleme\s*Tarih|Hazırlanma\s*Tarih|Prepared\s*Date|Test\s*Tarih|\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})", "weight": 2},
                "olcum_tarihi_saati": {"pattern": r"(?:Ölçüm\s*Tarih|Measurement\s*Date|Test\s*Tarih|Test\s*Date|Ölçüm\s*Saat|Measurement\s*Time|Test\s*Time|\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})", "weight": 2},
                "hazırlayan_kurulus": {"pattern": r"(?:Hazırlayan\s*Kuruluş|Prepared\s*By|Ölçümü\s*Yapan|Measured\s*By|Test\s*Yapan|Tested\s*By|Firma|Company|Şirket|Corporation|Kurum|Institution|Akredite|Accredited)", "weight": 2},
                "olcum_uzman": {"pattern": r"(?:Ölçüm\s*Yapan\s*Uzman|Measurement\s*Expert|Test\s*Uzman|Test\s*Expert|Uzman|Expert|Yetkili|Authorized|İmza|Signature|Onay|Approval|Sorumlu|Responsible)", "weight": 2}
            },
            "Ölçüm Ortam, Makine ve Çalışan Bilgileri": {
                "firma_adi_adres": {"pattern": r"(?:Firma\s*Ad[ıi]|Company\s*Name|Şirket|Corporation|İşyeri|Workplace|Kuruluş|Organization|Adres|Address|Konum|Location)", "weight": 2},
                "ortam_tanimi": {"pattern": r"(?:Ortam|Environment|Fabrika|Factory|Atölye|Workshop|Şantiye|Construction|Ofis|Office|Tesis|Facility|Alan|Area|Bölge|Zone)", "weight": 2},
                "makine_ekipman": {"pattern": r"(?:Makine|Machine|Ekipman|Equipment|Alet|Tool|Cihaz|Device|Model|Seri\s*No|Serial\s*Number|Konum|Position|Yerleşim|Layout)", "weight": 2},
                "calisan_bilgileri": {"pattern": r"(?:Çalışan|Employee|Worker|Personel|Staff|Sicil\s*No|Personnel\s*No|ID\s*No|Görev|Task|Job|Position|Unvan|Title)", "weight": 2},
                "maruziyet_suresi": {"pattern": r"(?:Maruziyet\s*Süre|Exposure\s*Time|Çalışma\s*Süre|Working\s*Time|Saat|Hour|Dakika|Minute|Günlük|Daily|Haftalık|Weekly)", "weight": 2}
            },
            "Ölçüm Cihazı ve Kalibrasyon Bilgileri": {
                "cihaz_marka_model": {"pattern": r"(?:Titreşim\s*Ölçer|Vibration\s*Meter|İvmeölçer|Accelerometer|Sensör|Sensor|Marka|Brand|Model|Tip|Type|Cihaz|Device|Instrument)", "weight": 4},
                "seri_numarasi": {"pattern": r"(?:Seri\s*No|Serial\s*Number|S/N|SN|Seri|Serial|Numara|Number|ID)", "weight": 2},
                "kalibrasyon_durumu": {"pattern": r"(?:Kalibrasyon|Calibration|Sertifika|Certificate|Akreditasyon|Accreditation|Geçerlilik|Validity|Son\s*Kalibrasyon|Last\s*Calibration|İzlenebilir|Traceable)", "weight": 4}
            },
            "Titreşim Türleri ve Yasal Değerler": {
                "yasal_dayanak": {"pattern": r"(?:6331\s*sayılı|İş\s*Sağlığ[ıi]|Occupational\s*Health|Work\s*Safety|Yönetmelik|Regulation|Directive|Kanun|Law|Mevzuat|Legislation)", "weight": 3},
                "titresim_turleri": {"pattern": r"(?:El[\\-\\s]*Kol\s*Titreşim|Hand[\\-\\s]*Arm\s*Vibration|HAV|Bütün\s*Vücut\s*Titreşim|Whole\s*Body\s*Vibration|WBV|Titreşim\s*Tür|Vibration\s*Type)", "weight": 4},
                "maruziyet_degerler": {"pattern": r"(?:Maruziyet\s*Değer|Exposure\s*Value|Eylem\s*Değer|Action\s*Value|Sınır\s*Değer|Limit\s*Value|A\\(8\\)|m/s2|2[\\.,]5|5[\\.,]0|0[\\.,]5|1[\\.,]15)", "weight": 4},
                "deger_tablolari": {"pattern": r"(?:Tablo|Table|Liste|List|Değer|Value|El[\\-\\s]*Kol|Hand[\\-\\s]*Arm|Bütün\s*Vücut|Whole\s*Body|2[\\.,]5|5[\\.,]0|0[\\.,]5|1[\\.,]15)", "weight": 4}
            },
            "Ölçüm Metodolojisi ve Uygulanan Standartlar": {
                "uygulanan_standartlar": {"pattern": r"(?:TS\s*EN\s*ISO\s*5349|TS\s*ISO\s*2631|ISO\s*5349|ISO\s*2631|Standart|Standard|Norm|Specification)", "weight": 4},
                "olcum_sekli": {"pattern": r"(?:Tri[\\-\\s]*aksiyel|Tri[\\-\\s]*axial|İvmeölçer|Accelerometer|Sensör|Sensor|Ölçüm\s*Şekl|Measurement\s*Method|Test\s*Method)", "weight": 3},
                "a8_hesaplaması": {"pattern": r"(?:A\\(8\\)|Günlük\s*Maruziyet|Daily\s*Exposure|8\s*saat|8\s*hour|Referans|Reference|Normalleştir|Normalize|Hesaplama|Calculation)", "weight": 3}
            },
            "Ölçüm Sonuçları ve Analizler": {
                "olcum_verileri": {"pattern": r"(?:İvme|Acceleration|Hız|Velocity|Yer\s*değiştirme|Displacement|m/s2|mm/s|µm|Tablo|Table|Veri|Data|Sonuç|Result)", "weight": 6},
                "a8_degerleri": {"pattern": r"(?:A\\(8\\)|Günlük\s*Maruziyet\s*Değer|Daily\s*Exposure\s*Value|Hesaplanan|Calculated|m/s2)", "weight": 6},
                "makine_sagligi": {"pattern": r"(?:Makine\s*Sağlığ[ıi]|Machine\s*Health|Arıza\s*Teşhis|Fault\s*Diagnosis|FFT|Spektral\s*Analiz|Spectral\s*Analysis|Fourier|Dalga\s*Form|Waveform)", "weight": 4},
                "sonuc_tablolari": {"pattern": r"(?:Sonuç\s*Tablosu|Result\s*Table|Değerlendirme\s*Tablosu|Evaluation\s*Table|Yasal\s*Değerlendirme|Legal\s*Assessment|Uygun|Suitable|Eylem\s*Değeri\s*Aşıl|Action\s*Value\s*Exceeded)", "weight": 4}
            },
            "Değerlendirme, Yorum ve Önlemler": {
                "yasal_uygunluk": {"pattern": r"(?:Yasal\s*Uygunluk|Legal\s*Compliance|Uygun|Suitable|Uygunsuz|Non[\\-\\s]*Suitable|Eylem\s*Değer|Action\s*Value|Sınır\s*Değer|Limit\s*Value|Aşıl|Exceed)", "weight": 5},
                "uzman_yorumu": {"pattern": r"(?:Uzman\s*Yorumu|Expert\s*Opinion|Değerlendirme|Evaluation|Analiz|Analysis|Yorum|Comment|Görüş|Opinion|Çevresel\s*Koşul|Environmental\s*Condition)", "weight": 5},
                "iyilestirici_onlemler": {"pattern": r"(?:İyileştirici\s*Önlem|Improvement\s*Measure|Mühendislik\s*Kontrol|Engineering\s*Control|İdari\s*Kontrol|Administrative\s*Control|KKD|PPE|Öneri|Recommendation|Tedbir|Measure)", "weight": 5},
                "onlem_detaylari": {"pattern": r"(?:Titreşim\s*Sönümleyici|Vibration\s*Damper|Anti[\\-\\s]*titreşim|Anti[\\-\\s]*vibration|Eldiven|Glove|İzolasyon|Isolation|Rotasyon|Rotation|Eğitim|Training|Bakım|Maintenance)", "weight": 5}
            },
            "Ekler ve Görseller": {
                "kroki_plan": {"pattern": r"(?:Kroki|Sketch|Plan|Layout|Yerleşim\s*Plan|Site\s*Plan|Ölçüm\s*Nokta|Measurement\s*Point|Konum|Location)", "weight": 1},
                "fotograflar": {"pattern": r"(?:Fotoğraf|Photo|Görsel|Visual|Resim|Picture|Image|Şekil|Figure)", "weight": 1},
                "kalibrasyon_sertifikalari": {"pattern": r"(?:Kalibrasyon\s*Sertifika|Calibration\s*Certificate|Sertifika\s*Kopya|Certificate\s*Copy)", "weight": 1},
                "mevzuat_standart_listesi": {"pattern": r"(?:Mevzuat\s*Liste|Legislation\s*List|Standart\s*Liste|Standard\s*List|Referans|Reference|Atıf|Citation)", "weight": 1},
                "hesaplama_detaylari": {"pattern": r"(?:A\\(8\\)\s*Hesaplama|A\\(8\\)\s*Calculation|Hesaplama\s*Detay|Calculation\s*Detail|Formül|Formula|Denklem|Equation)", "weight": 1}
            }
        }
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni Türkçe'ye çevir - şimdilik devre dışı"""
        if source_lang != 'tr':
            logger.info(f"Tespit edilen dil: {source_lang.upper()} - Çeviri yapılmıyor, orijinal metin kullanılıyor")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosya formatına göre metin çıkarma"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.doc':
            logger.warning("DOC formatı için DOCX'e dönüştürme gerekiyor veya OCR kullanılacak")
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatı: {file_ext}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den sadece PyPDF2 ile metin çıkarma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF analizi başlatılıyor - Toplam sayfa: {total_pages}")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = page_text.replace('|', ' ')
                    page_text = page_text.strip()
                    
                    all_text += page_text + "\n"
                
                # Metni temizle
                all_text = all_text.replace('—', '-')
                all_text = all_text.replace('"', '"').replace('"', '"')
                all_text = all_text.replace('´', "'")
                all_text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]+', ' ', all_text)
                all_text = all_text.strip()
                
                logger.info(f"✅ PDF analizi tamamlandı:")
                logger.info(f"   📊 Toplam metin uzunluğu: {len(all_text):,} karakter")
                logger.info(f"   📄 Toplam sayfa: {total_pages}")
                
                return all_text
                
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Paragrafları oku
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Tablolardan metin çıkar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            logger.info(f"DOCX'den {len(text)} karakter metin çıkarıldı")
            return text
            
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasından metin çıkarma"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text:
                # UTF-8 başarısız olursa diğer encoding'leri dene
                encodings = ['cp1254', 'iso-8859-9', 'latin1']
                for encoding in encodings:
                    try:
                        with open(txt_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        if text:
                            break
                    except:
                        continue
            
            logger.info(f"TXT'den {len(text)} karakter metin çıkarıldı")
            return text.strip()
            
        except Exception as e:
            logger.error(f"TXT metin çıkarma hatası: {e}")
            return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, VibrationAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                # Scoring algoritması - aydınlatma kodundaki mantığı koruyoruz
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
                    
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = VibrationAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={
                    "pattern_used": pattern,
                    "matches_count": len(matches) if matches else 0
                }
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, VibrationAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """Titreşim raporuna özgü değerleri çıkar"""
        values = {
            "rapor_numarasi": "Bulunamadı",
            "olcum_tarihi": "Bulunamadı",
            "rapor_tarihi": "Bulunamadı",
            "olcum_cihazi": "Bulunamadı",
            "firma_adi": "Bulunamadı",
            "titresim_turu": "Bulunamadı",
            "a8_degeri": "Bulunamadı",
            "yasal_uygunluk": "Bulunamadı"
        }
        
        # RAPOR NUMARASI
        report_no_patterns = [
            # Tablo yapısında rapor no
            r"(?i)(?:RAPOR\s*NO|REPORT\s*NO|RAPOR\s*NUMARAS[II])\s*[|\s]*\s*([A-Z0-9\-\/]{3,20})",
            
            # Klasik format'lar
            r"(?i)(?:RAPOR\s*NO|REPORT\s*NO|RAPOR\s*NUMARAS[II]|REPORT\s*NUMBER|DÖKÜMAN\s*NO|DOCUMENT\s*NO)\s*[:=]?\s*([A-Z0-9\-\/]{3,20})",
            
            r"(?i)(?:TEST\s*NO|BELGE\s*NO|REFERANS\s*NO|REF\s*NO)\s*[:=]?\s*([A-Z0-9\-\/]{3,20})"
        ]
        
        for pattern in report_no_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 3 <= len(result) <= 20:
                    values["rapor_numarasi"] = result
                    break
        
        # ÖLÇÜM TARİHİ
        measurement_date_patterns = [
        # Tablo yapısı için: "Ölçüm Tarihi / Saati" sonrası gelen tarih
        r"(?i)(?:ÖLÇÜM\s*TARİH[İI]?\s*\/?\s*SAAT[İI]?)\s*[|\s]*\s*(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
        
        # Saat aralığı ile birlikte: 28.11.2023 / 10:58-11:07
        r"(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})\s*\/?\s*\d{1,2}:\d{1,2}[\-:]\d{1,2}:\d{1,2}",
        
        # Klasik format'lar
        r"(?i)(?:ÖLÇÜM\s*TARİH[İI]?|MEASUREMENT\s*DATE|TEST\s*TARİH[İI]?|TEST\s*DATE)\s*[:=]\s*(\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})",
        
        # Tablo hücresi içinde sadece tarih
        r"(?i)(?:ÖLÇÜM.*?(?:TARİH|DATE).*?)\s*[|\s]*\s*(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
        
        # Alternatif tablo yapıları
        r"(?i)(?:ÖLÇÜM\s*YAPILDI[ĞG]I|MEASURED\s*ON)\s*[:=]?\s*(\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})",
        
        # PDF'den çıkan metin bozuklukları için
        r"(?:TARİH[İI]?\s*\/?\s*SAAT[İI]?).*?(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
        
        # Genel tarih pattern'ı (en son denenir)
        r"\b(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{4})\b"
    ]
        
        for pattern in measurement_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["olcum_tarihi"] = match.group(1)
                break
        
        # RAPOR TARİHİ
        report_date_patterns = [
             # Rapor tarihi tablo yapısı
            r"(?i)(?:RAPOR\s*TARİH[İI]?\s*\/?\s*SAAT[İI]?)\s*[|\s]*\s*(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})",
            
            # Klasik format'lar
            r"(?i)(?:RAPOR\s*TARİH[İI]?|REPORT\s*DATE|HAZIRLANMA\s*TARİH[İI]?|PREPARED\s*ON|DÜZENLEME\s*TARİH[İI]?)\s*[:=]\s*(\d{1,2}[\/\.\-]\d{1,2}[\/\.\-]\d{2,4})",
            
            # Belge başlığında tarih
            r"(?i)(?:BELGE|DÖKÜMAN|DOCUMENT).*?TARİH[İI]?.*?(\d{1,2}[\.\/\-]\d{1,2}[\.\/\-]\d{2,4})"
        ]
        
        for pattern in report_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["rapor_tarihi"] = match.group(1)
                break
            
            # Eğer rapor tarihi bulunamazsa, ölçüm tarihi ile aynı olabilir
        if values["rapor_tarihi"] == "Bulunamadı" and values["olcum_tarihi"] != "Bulunamadı":
            values["rapor_tarihi"] = "Rapor tarihi ayrı belirtilmemiş"
        
        # ÖLÇÜM CİHAZI
        device_patterns = [
            r"(?i)(?:TİTREŞİM\s*ÖLÇER|VIBRATION\s*METER)\s*[:=]?\s*([A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜ0-9\s\.\-]{2,50}(?=\s*(?:\d{4,}|Seri|No|$)))",
            r"(?i)(?:İVMEÖLÇER|ACCELEROMETER)\s*[:=]?\s*([A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜ0-9\s\.\-]{2,50}(?=\s*(?:\d{4,}|Seri|No|$)))",
            r"(?i)(?:MARKA|BRAND)\s*[:=]?\s*([A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜ0-9\s\.\-]{2,30})(?:.*?(?:MODEL|TİP|TYPE)\s*[:=]?\s*([A-ZÇĞİÖŞÜ0-9\-]{1,30}))?"
        ]
        
        for pattern in device_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 2 <= len(result) <= 35:
                    values["olcum_cihazi"] = result
                    break
        
        # FİRMA ADI
        company_patterns = [
            r"(?i)(?:FİRMA\s*ADI|COMPANY\s*NAME|ŞİRKET|İŞYERİ|WORKPLACE)\s*[:\-]?\s*([A-ZÇĞİÖŞÜ0-9][A-ZÇĞİÖŞÜ0-9\.\&\süçğıöş\-]{3,60})",
            r"(?<!\w)([A-ZÇĞİÖŞÜ][A-Za-zÇĞİÖŞÜçğıöşü\s\.\&]{2,50}?(?:A\.Ş|LTD\.?\s*ŞTİ|SANAYİ|TİCARET|İÇECEK|GIDA|TEKSTİL|OTOMOTİV|İNŞAAT|MAKİNA|COMPANY|CORP|CORPORATION|INC|INCORPORATED)\.?)(?!\w)"
        ]
        
        for pattern in company_patterns:
            matches = re.findall(pattern, text)
            for result in matches:
                result = result.strip()
                if 3 <= len(result) <= 60:
                    values["firma_adi"] = result
                    break
            if values["firma_adi"] != "Bulunamadı":
                break
        
        # TİTREŞİM TÜRÜ
        vibration_type_patterns = [
            r"(?i)(EL[\\-\\s]*KOL\s*TİTREŞİM|HAND[\\-\\s]*ARM\s*VIBRATION|HAV)",
            r"(?i)(BÜTÜN\s*VÜCUT\s*TİTREŞİM|WHOLE\s*BODY\s*VIBRATION|WBV)"
        ]
        
        for pattern in vibration_type_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip().upper()
                if "EL" in result or "HAND" in result or "HAV" in result:
                    values["titresim_turu"] = "El-Kol Titreşimi"
                elif "VÜCUT" in result or "BODY" in result or "WBV" in result:
                    values["titresim_turu"] = "Bütün Vücut Titreşimi"
                break
        
        # A(8) DEĞERİ
        a8_patterns = [
            r"(?i)A\(8\)\s*[:=]?\s*([0-9]+[.,][0-9]+)\s*m/s2",
            r"(?i)GÜNLÜK\s*MARUZİYET\s*DEĞERİ\s*[:=]?\s*([0-9]+[.,][0-9]+)\s*m/s2",
            r"(?i)DAILY\s*EXPOSURE\s*VALUE\s*[:=]?\s*([0-9]+[.,][0-9]+)\s*m/s2"
        ]
        
        for pattern in a8_patterns:
            match = re.search(pattern, text)
            if match:
                values["a8_degeri"] = match.group(1).replace(',', '.')
                break
        
        # YASAL UYGUNLUK
        compliance_patterns = [
            r"(?i)\b(UYGUN|SUITABLE|CONFORM|GEÇERLİ|VALID|PASS)\b",
            r"(?i)\b(UYGUNSUZ|NOT\s*SUITABLE|NON[\\-\\s]*CONFORM|GEÇERSİZ|INVALID|FAIL)\b",
            r"(?i)(?:EYLEM\s*DEĞERİ\s*AŞIL|ACTION\s*VALUE\s*EXCEED|SINIR\s*DEĞERİ\s*AŞIL|LIMIT\s*VALUE\s*EXCEED)"
        ]
        
        for pattern in compliance_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(0).strip().upper()
                if any(word in result for word in ["UYGUN", "SUITABLE", "CONFORM", "GEÇERLİ", "VALID", "PASS"]):
                    values["yasal_uygunluk"] = "UYGUN"
                elif any(word in result for word in ["UYGUNSUZ", "NOT SUITABLE", "NON-CONFORM", "GEÇERSİZ", "INVALID", "FAIL", "AŞIL", "EXCEED"]):
                    values["yasal_uygunluk"] = "UYGUNSUZ"
                break
        
        return values

    def validate_vibration_document(self, text: str) -> bool:
        """Dokümanın mekanik titreşim ölçüm raporu olup olmadığını kontrol et"""
        
        vibration_keywords = [
            # Temel titreşim terimleri
            "titreşim", "vibration", "ivme", "acceleration", "mekanik", "mechanical",
            
            # Titreşim türleri
            "el-kol", "hand-arm", "hav", "bütün vücut", "whole body", "wbv",
            
            # Ölçüm terimleri
            "ölçüm", "measurement", "test", "analiz", "analysis", "değerlendirme", "assessment",
            
            # Cihaz terimleri
            "ivmeölçer", "accelerometer", "titreşim ölçer", "vibration meter", "sensör", "sensor",
            
            # Yasal terimler
            "maruziyet", "exposure", "eylem değeri", "action value", "sınır değeri", "limit value",
            
            # Standart terimleri
            "iso 5349", "iso 2631", "ts en", "standart", "standard",
            
            # Birimler
            "m/s2", "a(8)", "günlük maruziyet", "daily exposure",
            
            # Rapor terimleri
            "rapor", "report", "değerlendirme", "evaluation", "sonuç", "result"
        ]
        
        # En az 3 anahtar kelime bulunmalı
        found_keywords = 0
        found_words = []
        
        for keyword in vibration_keywords:
            if re.search(rf"\b{keyword}\b", text, re.IGNORECASE):
                found_keywords += 1
                found_words.append(keyword)
                
        logger.info(f"Doküman validasyonu: {found_keywords} anahtar kelime bulundu: {found_words[:10]}")
        
        return found_keywords >= 3

    def check_date_validity(self, measurement_date: str, report_date: str) -> Tuple[bool, str]:
        """Ölçüm ve rapor tarihlerini bugünkü tarih ile kontrol et (1 yıl kuralı)"""
        
        # Hiçbir tarih bulunamadıysa
        if measurement_date == "Bulunamadı" and report_date == "Bulunamadı":
            return False, "RAPOR GEÇERSİZ: Ne ölçüm ne de rapor tarihi bulunamadı"
        
        # Kontrol edilecek tarihleri topla
        dates_to_check = []
        
        if measurement_date != "Bulunamadı":
            dates_to_check.append(("Ölçüm", measurement_date))
        
        if report_date != "Bulunamadı" and report_date != "Rapor tarihi ayrı belirtilmemiş":
            dates_to_check.append(("Rapor", report_date))
        
        # Hiç kontrol edilecek tarih yoksa
        if not dates_to_check:
            return False, "RAPOR GEÇERSİZ: Geçerli tarih bulunamadı"
        
        try:
            date_formats = ['%d/%m/%Y', '%d.%m.%Y', '%d-%m-%Y', '%d/%m/%y', '%d.%m.%y', '%d-%m-%y']
            today = datetime.now()
            
            for date_type, date_str in dates_to_check:
                # Tarihi parse et
                parsed_date = None
                for fmt in date_formats:
                    try:
                        parsed_date = datetime.strptime(date_str, fmt)
                        break
                    except:
                        continue
                
                if not parsed_date:
                    return False, f"RAPOR GEÇERSİZ: {date_type} tarihi formatı tanınmadı ({date_str})"
                
                # 2 haneli yıl düzeltmesi
                if parsed_date.year < 1950:
                    parsed_date = parsed_date.replace(year=parsed_date.year + 2000)
                
                # Bugünden farkı hesapla
                diff = abs((today - parsed_date).days)
                
                # 1 yıldan (365 gün) eskiyse geçersiz
                if diff > 365:
                    return False, f"RAPOR GEÇERSİZ: {date_type} tarihi 1 yıldan eski ({diff} gün önce - {date_str})"
            
            # Tüm tarihler geçerliyse
            checked_dates = [f"{dt[0]}: {dt[1]}" for dt in dates_to_check]
            return True, f"Tüm tarihler geçerli - {', '.join(checked_dates)}"
            
        except Exception as e:
            return False, f"RAPOR GEÇERSİZ: Tarih kontrolü yapılamadı - {e}"
            

    def generate_improvement_actions(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Dinamik iyileştirme önerileri oluştur"""
        actions = []
        
        # Kategorileri puana göre sırala (düşükten yükseğe)
        sorted_categories = sorted(
            scores["category_scores"].items(), 
            key=lambda x: x[1]["percentage"]
        )
        
        category_actions = {
            "Rapor Kimlik Bilgileri": [
                "Rapor numarasını benzersiz ve takip edilebilir şekilde belirtiniz",
                "Rapor düzenleme tarihini açıkça yazınız",
                "Ölçüm tarih ve saat aralığını detaylandırınız",
                "Ölçümü yapan akredite kuruluş bilgilerini eksiksiz yazınız",
                "Ölçüm uzmanının adı, unvanı ve imzasını alınız"
            ],
            "Ölçüm Ortam, Makine ve Çalışan Bilgileri": [
                "Firma adı ve tam adresini belirtiniz",
                "Çalışma ortamını detaylı tanımlayınız (fabrika, atölye, şantiye vb.)",
                "Titreşim kaynağı makine/ekipman bilgilerini eksiksiz yazınız",
                "Maruz kalan çalışan bilgilerini (ad, sicil no, görev) ekleyiniz",
                "Günlük maruziyet süresini saat/dakika olarak belirtiniz"
            ],
            "Ölçüm Cihazı ve Kalibrasyon Bilgileri": [
                "Titreşim ölçer/ivmeölçer marka ve modelini yazınız",
                "Tüm cihazların seri numaralarını belirtiniz",
                "Güncel kalibrasyon sertifikalarını rapora ekleyiniz",
                "Cihazların ulusal standartlara göre izlenebilirliğini teyit ediniz"
            ],
            "Titreşim Türleri ve Yasal Değerler": [
                "6331 sayılı İSG Kanunu ve ilgili yönetmeliğe atıf yapınız",
                "El-Kol ve Bütün Vücut titreşim türlerini açıklayınız",
                "Maruziyet eylem ve sınır değerlerini tablo halinde sunınuz",
                "Yasal değerlerin anlamlarını (eylem/sınır değeri) açıklayınız"
            ],
            "Ölçüm Metodolojisi ve Uygulanan Standartlar": [
                "TS EN ISO 5349 (El-Kol) ve TS ISO 2631 (Bütün Vücut) standartlarını belirtiniz",
                "Tri-aksiyel ölçüm şeklini ve sensör yerleşimini açıklayınız",
                "A(8) günlük maruziyet değeri hesaplama yöntemini detaylandırınız",
                "8 saatlik referans çalışma süresi normalizasyonunu açıklayınız"
            ],
            "Ölçüm Sonuçları ve Analizler": [
                "İvme, hız ve yer değiştirme değerlerini tablo halinde sunınız",
                "Her çalışan/görev için A(8) değerlerini hesaplayınız",
                "Gerekiyorsa FFT spektral analizleri ve dalga formlarını ekleyiniz",
                "Sonuç tablosunda yasal değerlendirme sütunu bulundurunuz"
            ],
            "Değerlendirme, Yorum ve Önlemler": [
                "Yasal uygunluk durumunu açıkça belirtiniz",
                "Uzman yorumu ve çevresel faktör değerlendirmesi yapınız",
                "Mühendislik, idari ve KKD önlemlerini öncelik sırasıyla listeyiniz",
                "Titreşim sönümleyici, anti-titreşim eldiven gibi spesifik önerilerde bulununuz"
            ],
            "Ekler ve Görseller": [
                "Ölçüm noktaları kroki/yerleşim planını hazırlayınız",
                "Ölçüm sırasında çekilmiş fotoğrafları ekleyiniz",
                "Kalibrasyon sertifikalarının kopyalarını sunınız",
                "İlgili mevzuat ve standart listesini hazırlayınız",
                "A(8) hesaplama detaylarını gerektiğinde ekleyiniz"
            ]
        }
        
        # En düşük puanlı 5 kategori için öneriler ekle
        for category, score_data in sorted_categories[:5]:
            if score_data["percentage"] < 70:
                actions.extend(category_actions.get(category, []))
        
        # Genel öneriler
        if scores["percentage"] < 50:
            actions.insert(0, "ÖNCELİK: Rapor yapısını ve içeriğini kapsamlı olarak yeniden düzenleyiniz")
        
        return actions
    
    def generate_recommendations(self, analysis_results: Dict, scores: Dict, date_valid: bool = True, date_message: str = "") -> List[str]:
        """Öneriler oluştur - Güncellenmiş versiyon"""
        recommendations = []
        
        total_percentage = scores["percentage"]
        
        # Tarih geçerlilik kontrolü
        if not date_valid:
            # date_message'dan hangi tarih sorunu olduğunu belirle
            if "Ne ölçüm ne de rapor tarihi bulunamadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ne ölçüm ne de rapor tarihi bulunamadı")
            elif "Ölçüm tarihi formatı tanınmadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ölçüm tarihi formatı tanınmadı")
            elif "Rapor tarihi formatı tanınmadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Rapor tarihi formatı tanınmadı")
            elif "Ölçüm tarihi 1 yıldan eski" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ölçüm tarihi 1 yıldan eski")
            elif "Rapor tarihi 1 yıldan eski" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Rapor tarihi 1 yıldan eski")
            elif "Geçerli tarih bulunamadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Geçerli tarih bulunamadı")
            else:
                recommendations.append(f"❌ RAPOR GEÇERSİZ: {date_message}")
            return recommendations
        
        # Tarih geçerliyse normal değerlendirme
        if total_percentage >= 70:
            recommendations.append(f"✅ Mekanik Titreşim Ölçüm Raporu GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ Mekanik Titreşim Ölçüm Raporu GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 40:
                recommendations.append(f"🔴 {category} bölümü yetersiz (%{category_score:.1f})")
                missing_items = [name for name, result in results.items() if not result.found]
                if missing_items:
                    recommendations.append(f"   Eksik: {', '.join(missing_items[:3])}")
            elif category_score < 70:
                recommendations.append(f"🟡 {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"🟢 {category} bölümü yeterli (%{category_score:.1f})")
        
        return recommendations


# Ana analiz fonksiyonunda da güncelleme yapılmalı
    def analyze_vibration_report(self, file_path: str) -> Dict[str, Any]:
        """Ana Mekanik Titreşim Ölçüm Raporu analiz fonksiyonu - Güncellenmiş"""
        logger.info("Mekanik Titreşim Ölçüm Raporu analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        if detected_lang != 'tr':
            logger.info(f"{detected_lang.upper()} dilinden Türkçe'ye çeviriliyor...")
            text = self.translate_to_turkish(text, detected_lang)
        
        # Doküman türü kontrolü
        if not self.validate_vibration_document(text):
            return {
                "error": "YANLIŞ DOKÜMAN: Bu dosya mekanik titreşim ölçüm raporu değil!",
                "document_type": "UNKNOWN",
                "suggestion": "Lütfen geçerli bir mekanik titreşim ölçüm raporu yükleyiniz."
            }
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        
        # Tarih geçerlilik kontrolü
        date_valid, date_message = self.check_date_validity(
            extracted_values.get("olcum_tarihi", "Bulunamadı"),
            extracted_values.get("rapor_tarihi", "Bulunamadı")
        )
        
        # Önerileri date_message ile birlikte oluştur
        recommendations = self.generate_recommendations(analysis_results, scores, date_valid, date_message)
        improvement_actions = self.generate_improvement_actions(analysis_results, scores)
        
        # Tarih geçersizse rapor otomatik başarısız
        if not date_valid:
            final_status = "FAIL"
            final_percentage = 0
        else:
            final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
            final_percentage = scores["percentage"]
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "tarih_gecerlilik": {
                "valid": date_valid,
                "message": date_message
            },
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "iyilestirme_eylemleri": improvement_actions,
            "ozet": {
                "toplam_puan": scores["total_score"] if date_valid else 0,
                "yuzde": final_percentage,
                "durum": final_status,
                "rapor_tipi": "MEKANİK_TİTREŞİM_OLCUM_RAPORU",
                "tarih_gecersiz": not date_valid
            }
        }
        
        return report

def main():
    """Ana fonksiyon"""
    analyzer = VibrationReportAnalyzer()

    # Test için örnek dosya yolları
    file_path = r"C:\Users\nuvo_teknik_2\Desktop\PILZ DOCUMENTS\2.5 Titreşim\Coca Cola Titresim Olcum Raporu.pdf"
    # file_path = r"C:\Users\example\Desktop\VIBRATION_REPORTS\vibration_report.docx"

    if not os.path.exists(file_path):
        print(f"❌ Dosya bulunamadı: {file_path}")
        print("Lütfen geçerli bir dosya yolu belirtiniz.")
        return
    
    print("🔧 Mekanik Titreşim Ölçüm Raporu Analizi Başlatılıyor...")
    print("=" * 60)
    
    report = analyzer.analyze_vibration_report(file_path)
    
    if "error" in report:
        print(f"❌ Hata: {report['error']}")
        return
    
    print("\n📊 ANALİZ SONUÇLARI")
    print("=" * 60)
    
    print(f"📅 Analiz Tarihi: {report['analiz_tarihi']}")
    print(f"📄 Dosya Tipi: {report['dosya_bilgisi']['file_type'].upper()}")
    print(f"🔍 Tespit Edilen Dil: {report['dosya_bilgisi']['detected_language'].upper()}")
    
    # Tarih geçerlilik kontrolü (güncellenmiş anahtar)
    if report['ozet']['tarih_gecersiz']:
        print(f"⚠️ TARİH GEÇERSİZLİĞİ: {report['tarih_gecerlilik']['message']}")
        print(f"📋 Toplam Puan: 0/100 (Otomatik başarısız)")
        print(f"📈 Yüzde: %0")
        print(f"🎯 Durum: FAIL")
    else:
        print(f"📋 Toplam Puan: {report['ozet']['toplam_puan']}/100")
        print(f"📈 Yüzde: %{report['ozet']['yuzde']}")
        print(f"🎯 Durum: {report['ozet']['durum']}")
        print(f"📅 Tarih Geçerliliği: {report['tarih_gecerlilik']['message']}")
    
    print(f"📄 Rapor Tipi: {report['ozet']['rapor_tipi']}")
    
    print("\n📋 ÖNEMLİ ÇIKARILAN DEĞERLER")
    print("-" * 40)
    for key, value in report['cikarilan_degerler'].items():
        display_name = {
            "rapor_numarasi": "Rapor Numarası",
            "olcum_tarihi": "Ölçüm Tarihi",
            "rapor_tarihi": "Rapor Tarihi",
            "olcum_cihazi": "Ölçüm Cihazı",
            "firma_adi": "Firma Adı",
            "titresim_turu": "Titreşim Türü",
            "a8_degeri": "A(8) Değeri",
            "yasal_uygunluk": "Yasal Uygunluk"
        }.get(key, key.replace('_', ' ').title())
        print(f"{display_name}: {value}")
    
    if not report['ozet']['tarih_gecersiz']:
        print("\n📊 KATEGORİ PUANLARI")
        print("-" * 40)
        for category, score_data in report['puanlama']['category_scores'].items():
            print(f"{category}: {score_data['normalized']}/{score_data['max_weight']} (%{score_data['percentage']:.1f})")
    
    print("\n💡 ÖNERİLER VE DEĞERLENDİRME")
    print("-" * 40)
    for recommendation in report['oneriler']:
        print(recommendation)
    
    print("\n📋 GENEL DEĞERLENDİRME")
    print("=" * 60)
    
    if report['ozet']['tarih_gecersiz']:
        print("❌ SONUÇ: GEÇERSİZ (TARİH SORUNU)")
        print(f"📝 Değerlendirme: {report['tarih_gecerlilik']['message']}")
        print("\n📌 YAPILMASI GEREKENLER:")
        print("1. Geçerli tarihli ölçümleri yapınız")
        print("2. Rapor tarihini güncelleyiniz")  
        print("3. Tarihlerin 1 yıldan eski olmadığından emin olunuz")
    elif report['ozet']['yuzde'] >= 70:
        print("✅ SONUÇ: GEÇERLİ")
        print(f"🌟 Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Mekanik titreşim ölçüm raporu genel olarak yeterli kriterleri sağlamaktadır.")
    else:
        print("❌ SONUÇ: GEÇERSİZ")
        print(f"⚠️ Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Mekanik titreşim ölçüm raporu minimum gereklilikleri sağlamamaktadır.")
        
        print("\n⚠️ EKSİK GEREKLİLİKLER:")
        for category, results in report['kategori_analizleri'].items():
            missing_items = []
            for criterion, result in results.items():
                if not result.found:
                    missing_items.append(criterion)
            
            if missing_items:
                print(f"\n🔍 {category}:")
                for item in missing_items:
                    readable_name = {
                        "rapor_numarasi": "Rapor Numarası",
                        "rapor_tarihi": "Rapor Tarihi",
                        "olcum_tarihi_saati": "Ölçüm Tarih/Saati",
                        "hazırlayan_kurulus": "Hazırlayan Kuruluş",
                        "olcum_uzman": "Ölçüm Uzmanı",
                        "firma_adi_adres": "Firma Adı/Adres",
                        "ortam_tanimi": "Ortam Tanımı",
                        "makine_ekipman": "Makine/Ekipman",
                        "calisan_bilgileri": "Çalışan Bilgileri",
                        "maruziyet_suresi": "Maruziyet Süresi",
                        "cihaz_marka_model": "Cihaz Marka/Model",
                        "seri_numarasi": "Seri Numarası",
                        "kalibrasyon_durumu": "Kalibrasyon Durumu",
                        "yasal_dayanak": "Yasal Dayanak",
                        "titresim_turleri": "Titreşim Türleri",
                        "maruziyet_degerler": "Maruziyet Değerler",
                        "deger_tablolari": "Değer Tabloları",
                        "uygulanan_standartlar": "Uygulanan Standartlar",
                        "olcum_sekli": "Ölçüm Şekli",
                        "a8_hesaplaması": "A(8) Hesaplaması",
                        "olcum_verileri": "Ölçüm Verileri",
                        "a8_degerleri": "A(8) Değerleri",
                        "makine_sagligi": "Makine Sağlığı",
                        "sonuc_tablolari": "Sonuç Tabloları",
                        "yasal_uygunluk": "Yasal Uygunluk",
                        "uzman_yorumu": "Uzman Yorumu",
                        "iyilestirici_onlemler": "İyileştirici Önlemler",
                        "onlem_detaylari": "Önlem Detayları",
                        "kroki_plan": "Kroki/Plan",
                        "fotograflar": "Fotoğraflar",
                        "kalibrasyon_sertifikalari": "Kalibrasyon Sertifikaları",
                        "mevzuat_standart_listesi": "Mevzuat/Standart Listesi",
                        "hesaplama_detaylari": "Hesaplama Detayları"
                    }.get(item, item.replace('_', ' ').title())
                    print(f"   ❌ {readable_name}")
        
        print("\n📌 YAPILMASI GEREKENLER:")
        if "iyilestirme_eylemleri" in report:
            for i, action in enumerate(report['iyilestirme_eylemleri'], 1):
                print(f"{i}. {action}")

if __name__ == "__main__":
    main()
