# ============================================
# PNÖMATİK DEVRE ŞEMASI ANALİZ SERVİSİ
# Standalone Service - Azure App Service Ready
# Port: 8010
# ============================================

# ============================================
# IMPORTS
# ============================================
import os
import json
import io
from datetime import datetime
import re
from typing import Dict, List, Tuple, Any, Optional
import PyPDF2
import pytesseract
from PIL import Image
from dataclasses import dataclass, asdict
import logging
import math
import cv2
import numpy as np
import fitz  # PyMuPDF
from docx import Document

from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import pdf2image

# ============================================
# LOGGING CONFIGURATION
# ============================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================
# TESSERACT CHECK
# ============================================
def check_tesseract_installation():
    """Check if Tesseract is properly installed"""
    try:
        version = pytesseract.get_tesseract_version()
        logger.info(f"Tesseract OCR kurulu - Sürüm: {version}")
        return True, f"Tesseract v{version}"
    except Exception as e:
        logger.error(f"Tesseract OCR kurulu değil: {e}")
        return False, str(e)

tesseract_available, tesseract_info = check_tesseract_installation()

# ============================================
# ANALİZ SINIFI - DATA CLASSES
# ============================================
@dataclass
class ComponentDetection:
    """Detected component information"""
    component_type: str
    label: str
    position: Tuple[int, int]
    confidence: float
    bounding_box: Tuple[int, int, int, int]

@dataclass
class CircuitAnalysisResult:
    """Analysis result for each criterion"""
    criteria_name: str
    found: bool
    content: str
    score: float
    max_score: float
    details: Dict[str, Any]
    visual_evidence: List[ComponentDetection]

# ============================================
# ANALİZ SINIFI - MAIN ANALYZER
# ============================================
class PneumaticCircuitAnalyzer:
    """Gelişmiş pnömatik devre şeması analiz sistemi"""
    
    def __init__(self):
        self.pneumatic_criteria_weights = {
            "Temel Sistem Bileşenleri": 25,
            "Pnömatik Semboller ve Vana Sistemleri": 30,
            "Akış Yönü ve Bağlantı Hatları": 20,
            "Sistem Bilgileri ve Teknik Parametreler": 15,
            "Dokümantasyon ve Standart Uygunluk": 10
        }
        
        self.pneumatic_criteria_details = {
            "Temel Sistem Bileşenleri": {
                "hava_kaynagi_ve_hazirlama": {
                    "pattern": r"(?i)(?:pressure\s*source|hava\s*kaynağı|basınçlı\s*hava|air\s*supply|P\s*=\s*\d+.*?bar|kompresör|pneumatic|pnömatik|bar)",
                    "weight": 5,
                    "description": "Hava kaynağı ve hazırlama ünitesi"
                },
                "filtre_regulator_lubrikator": {
                    "pattern": r"(?i)(?:FRL|filtre|filter|regulator|regülatör|lubricator|yağlayıcı|kondisyoner|hava\s*hazırlama)",
                    "weight": 5,
                    "description": "Hava hazırlama grubu (FRL)"
                },
                "basinc_gosterge_sensoru": {
                    "pattern": r"(?i)(?:manometre|pressure.*?gauge|gösterge|indicator|PI|PT|PS|basınç.*?sensör|ölçüm|pressure|basınç)",
                    "weight": 4,
                    "description": "Basınç gösterge ve sensörleri"
                },
                "susturucu_egzoz": {
                    "pattern": r"(?i)(?:susturucu|muffler|exhaust|egzoz|silencer|tahliye|vent|boşaltım)",
                    "weight": 4,
                    "description": "Susturucu ve egzoz elemanları"
                },
                "genel_pnomatik_sistem": {
                    "pattern": r"(?i)(?:pneumatic|pnömatik|circuit|devre|diagram|diyagram|şema|sistem|air|hava)",
                    "weight": 7,
                    "description": "Genel pnömatik sistem varlığı"
                }
            },
            "Pnömatik Semboller ve Vana Sistemleri": {
                "silindir_actuator": {
                    "pattern": r"(?i)(?:silindir|cylinder|piston|actuator|çift.*?etkili|tek.*?etkili|double.*?acting|single.*?acting|C\d+|MGF|CYL|SIL)",
                    "weight": 8,
                    "description": "Silindir ve aktüatör sembolleri"
                },
                "yon_kontrol_vanalar": {
                    "pattern": r"(?i)(?:VUVG|Y\d+|V\d+|valf|valve|5/2|4/2|3/2|2/2|yön.*?kontrol|directional.*?control|solenoid|vana|kontrol)",
                    "weight": 10,
                    "description": "Yön kontrol vanaları"
                },
                "hiz_kontrol_vanalar": {
                    "pattern": r"(?i)(?:hız.*?kontrol|speed.*?control|flow.*?control|akış.*?kontrol|throttle|kısıcı|flow|akış)",
                    "weight": 6,
                    "description": "Hız kontrol vanaları"
                },
                "basinc_kontrol_vanalar": {
                    "pattern": r"(?i)(?:basınç.*?kontrol|pressure.*?control|relief|emniyet|PRV|basınç.*?azaltıcı|pressure)",
                    "weight": 6,
                    "description": "Basınç kontrol vanaları"
                }
            },
            "Akış Yönü ve Bağlantı Hatları": {
                "hava_besleme_hatlari": {
                    "pattern": r"(?i)(?:besleme|supply.*?line|hava.*?hattı|ana.*?hat|pressure.*?line|P|feed|input|giriş)",
                    "weight": 5,
                    "description": "Hava besleme hatları"
                },
                "calisma_hatlari": {
                    "pattern": r"(?i)(?:A|B|çalışma.*?hattı|working.*?line|port|output|çıkış)",
                    "weight": 5,
                    "description": "Çalışma hatları (A, B portları)"
                },
                "egzoz_tahliye_hatlari": {
                    "pattern": r"(?i)(?:R|S|EA|EB|egzoz|exhaust|tahliye|drain|vent|return|boşaltım)",
                    "weight": 3,
                    "description": "Egzoz ve tahliye hatları"
                },
                "yon_oklari_akim_gosterimi": {
                    "pattern": r"(?i)(?:→|←|↑|↓|⇒|⇐|⇑|⇓|yön|direction|ok|arrow|akış|flow|hat|line)",
                    "weight": 4,
                    "description": "Yön okları ve akış gösterimi"
                },
                "baglanti_hatlari": {
                    "pattern": r"(?i)(?:bağlantı|connection|hat|line|pipe|boru|tube|hose)",
                    "weight": 3,
                    "description": "Genel bağlantı hatları"
                }
            },
            "Sistem Bilgileri ve Teknik Parametreler": {
                "calisma_basinci": {
                    "pattern": r"(?i)(?:P\s*=\s*\d+(?:\.\d+)?.*?bar|\d+(?:\.\d+)?.*?bar|çalışma.*?basınç|working.*?pressure|4-6.*?bar|basınç|pressure|\d+\s*bar)",
                    "weight": 4,
                    "description": "Çalışma basıncı değerleri"
                },
                "hava_tuketimi": {
                    "pattern": r"(?i)(?:Q\s*=\s*\d+.*?l/min|\d+.*?l/min|hava.*?tüketim|air.*?consumption|flow.*?rate|tüketim|consumption|l/min|flow)",
                    "weight": 3,
                    "description": "Hava tüketimi değerleri"
                },
                "strok_boyutlari": {
                    "pattern": r"(?i)(?:strok|stroke|s\s*=\s*\d+.*?mm|\d+.*?mm|boyut|dimension|mesafe|size|mm|cm)",
                    "weight": 3,
                    "description": "Strok ve boyut bilgileri"
                },
                "vana_tipleri_ozellikler": {
                    "pattern": r"(?i)(?:normalde.*?kapalı|normalde.*?açık|NC|NO|spring.*?return|yay.*?geri|5/2|4/2|3/2|2/2)",
                    "weight": 3,
                    "description": "Vana tipleri ve özellikleri"
                },
                "teknik_parametreler": {
                    "pattern": r"(?i)(?:teknik|technical|parametre|parameter|özellik|specification|spec|sistem|system)",
                    "weight": 2,
                    "description": "Genel teknik parametre varlığı"
                }
            },
            "Dokümantasyon ve Standart Uygunluk": {
                "sembol_standartlari": {
                    "pattern": r"(?i)(?:ISO.*?1219|DIN.*?ISO|pnömatik.*?sembol|pneumatic.*?symbol|standart|standard|ISO|DIN)",
                    "weight": 2,
                    "description": "Sembol standartları"
                },
                "cizim_bilgileri": {
                    "pattern": r"(?i)(?:çizim.*?tarih|drawing.*?date|tasarım.*?tarih|design.*?date|\d{2}.\d{2}.\d{4}|\d{2}/\d{2}/\d{4}|created|designed|tarih|date)",
                    "weight": 2,
                    "description": "Çizim bilgileri ve tarihler"
                },
                "proje_bilgileri": {
                    "pattern": r"(?i)(?:proje.*?adı|project.*?name|sistem.*?adı|description|açıklama)",
                    "weight": 2,
                    "description": "Proje bilgileri"
                },
                "firma_logo_imza": {
                    "pattern": r"(?i)(?:ACT|festo|FESTO|onay.*?tarih|approval|kontrol.*?tarih|check|created.*?by|checked.*?by|approved|firma|company)",
                    "weight": 2,
                    "description": "Firma logosu ve imza bilgileri"
                },
                "dokumantasyon_genel": {
                    "pattern": r"(?i)(?:revision|rev|circuit.*?number|customer|müşteri|sheet|size|boyut|diagram|diyagram|şema|schema)",
                    "weight": 2,
                    "description": "Genel dokümantasyon varlığı"
                }
            }
        }
        
        # Görsel tanıma için şablon desenler
        self.visual_templates = {
            "pneumatic_cylinders": ["cylinder", "piston", "MGF", "silindir"],
            "pneumatic_valves": ["VUVG", "valve", "valf", "Y01", "Y02"],
            "pneumatic_frl": ["FRL", "filtre", "regulator", "MS6"],
            "pneumatic_connections": ["connection", "T", "junction", "bağlantı"],
            "flow_arrows": ["arrow", "ok", "→", "←"],
            "pressure_gauges": ["gauge", "manometer", "PI", "pressure"]
        }

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den metin çıkarma"""
        try:
            text = ""
            doc = fitz.open(pdf_path)
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            
            if not text.strip():
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = text.replace('|', ' ')
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""

    def extract_text_with_ocr(self, pdf_path: str) -> str:
        """Extract text from PDF using OCR with enhanced settings"""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            
            for page_num in range(min(5, len(doc))):
                page = doc[page_num]
                mat = fitz.Matrix(3.0, 3.0)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.pil_tobytes(format="PNG")
                img = Image.open(io.BytesIO(img_data))
                img_array = np.array(img)
                
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                
                temp_path = f"temp_page_{page_num}.png"
                cv2.imwrite(temp_path, thresh)
                
                # Multiple OCR passes with different settings
                ocr_configs = [
                    '--psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,:-/+()[]{}ÇĞİÖŞÜçğıöşü ',
                    '--psm 4',
                    '--psm 11',
                    '--psm 8',
                    '--psm 3'
                ]
                
                page_text = ""
                for config in ocr_configs:
                    try:
                        ocr_result = pytesseract.image_to_string(temp_path, lang='tur+eng+deu', config=config)
                        if len(ocr_result.strip()) > len(page_text.strip()):
                            page_text = ocr_result
                    except:
                        continue
                
                text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"
                
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"OCR extraction error: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""

    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1254', 'iso-8859-9']
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""

    def extract_text_from_image(self, image_path: str) -> str:
        """OCR ile metin çıkarma"""
        try:
            img = cv2.imread(image_path)
            if img is None:
                return ""
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            denoised = cv2.medianBlur(enhanced, 3)
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyzÇçĞğİıÖöŞşÜü. '
            text = pytesseract.image_to_string(thresh, lang='tur+eng', config=custom_config)
            
            text = re.sub(r'[^\w\s\.\-\(\)\[\]\{\}\+\=\*\/]', '', text)
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text
            
        except Exception as e:
            logger.error(f"OCR hatası: {e}")
            return ""
        
    def detect_visual_components(self, image_path: str) -> List[ComponentDetection]:
        """Görsel bileşen tespiti"""
        detections = []
        try:
            img = cv2.imread(image_path)
            if img is None:
                return detections
            
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            denoised = cv2.medianBlur(enhanced, 3)
            edges = cv2.Canny(denoised, 30, 100)
            
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,3))
            dilated = cv2.dilate(edges, kernel, iterations=1)
            contours, _ = cv2.findContours(dilated, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
            
            for i, contour in enumerate(contours):
                area = cv2.contourArea(contour)
                if area < 20:
                    continue
                    
                perimeter = cv2.arcLength(contour, True)
                if perimeter == 0:
                    continue
                    
                approx = cv2.approxPolyDP(contour, 0.04 * perimeter, True)
                x, y, w, h = cv2.boundingRect(contour)
                aspect_ratio = w / h if h > 0 else 0
                
                component_type = "unknown"
                confidence = 0.2
                
                if len(approx) == 4:
                    if 0.7 < aspect_ratio < 1.4:
                        component_type = "valve"
                        confidence = 0.7
                    elif aspect_ratio > 1.5:
                        component_type = "cylinder"
                        confidence = 0.8
                    else:
                        component_type = "component"
                        confidence = 0.5
                elif len(approx) == 3:
                    component_type = "flow_control"
                    confidence = 0.6
                elif len(approx) > 10:
                    component_type = "connection"
                    confidence = 0.6
                
                if w > 10 and h > 10:
                    detection = ComponentDetection(
                        component_type=component_type,
                        label=f"{component_type}_{len(detections)+1}",
                        position=(x + w//2, y + h//2),
                        confidence=confidence,
                        bounding_box=(x, y, w, h)
                    )
                    detections.append(detection)
            
            unique_detections = []
            seen_positions = set()
            for detection in detections:
                pos_key = (detection.position[0] // 30, detection.position[1] // 30)
                if pos_key not in seen_positions and detection.confidence > 0.2:
                    unique_detections.append(detection)
                    seen_positions.add(pos_key)
            
            return unique_detections[:30]
            
        except Exception as e:
            logger.error(f"Görsel bileşen tespiti hatası: {e}")
            return detections

    def analyze_criteria(self, text: str, ocr_text: str, visual_detections: List[ComponentDetection], category: str) -> Dict[str, CircuitAnalysisResult]:
        """Kriter analizi"""
        results = {}
        criteria = self.pneumatic_criteria_details.get(category, {})
        
        combined_text = f"{text} {ocr_text}".lower()
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            description = criterion_data["description"]
            
            text_matches = re.findall(pattern, combined_text, re.IGNORECASE | re.MULTILINE)
            
            visual_matches = []
            if criterion_name in ["silindir_actuator", "yon_kontrol_vanalar"]:
                relevant_detections = [d for d in visual_detections 
                                     if d.component_type in ["cylinder", "valve", "filter", "connection"] and d.confidence > 0.3]
                visual_matches = relevant_detections[:5]
            
            base_score = 0
            
            if text_matches:
                text_score = min(weight * 0.9, len(text_matches) * (weight * 0.4))
                base_score += text_score
            
            if visual_matches:
                visual_score = min(weight * 0.6, len(visual_matches) * (weight * 0.25))
                base_score += visual_score
            
            if visual_detections and len(visual_detections) > 2:
                base_score = max(base_score, weight * 0.4)
            
            if visual_detections:
                base_score += weight * 0.1
            
            if len(text_matches) >= 2:
                base_score += weight * 0.15
            
            total_score = min(weight, base_score)
            found = len(text_matches) > 0 or len(visual_matches) > 0 or len(visual_detections) > 0
            
            content_parts = []
            if text_matches:
                content_parts.append(f"Metin: {str(text_matches[:3])}")
            if visual_matches:
                content_parts.append(f"Görsel: {len(visual_matches)} tespit")
            if visual_detections and not visual_matches:
                content_parts.append(f"Görsel Bileşenler: {len(visual_detections)} adet")
            
            content = " | ".join(content_parts) if content_parts else "Görsel bileşenler algılandı"
            
            results[criterion_name] = CircuitAnalysisResult(
                criteria_name=description,
                found=found,
                content=content,
                score=total_score,
                max_score=weight,
                details={
                    "pattern_used": pattern,
                    "text_matches": len(text_matches),
                    "visual_matches": len(visual_matches),
                    "total_visual_components": len(visual_detections)
                },
                visual_evidence=visual_matches
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, CircuitAnalysisResult]]) -> Dict[str, Any]:
        """Puan hesaplama"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.pneumatic_criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                raw_percentage = category_earned / category_possible
                adjusted_percentage = math.pow(raw_percentage, 0.8)
                normalized_score = adjusted_percentage * category_max
            else:
                normalized_score = 0
            
            category_scores[category] = {
                "earned": round(category_earned, 2),
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round((category_earned / category_possible * 100), 2) if category_possible > 0 else 0
            }
            
            total_score += normalized_score
        
        final_score = min(100, total_score * 1.05)
        
        return {
            "category_scores": category_scores,
            "total_score": round(final_score, 2),
            "total_max_score": 100,
            "overall_percentage": round(final_score, 2)
        }

    def extract_specific_values(self, text: str, ocr_text: str) -> Dict[str, Any]:
        """Özel değer çıkarma"""
        combined_text = f"{text} {ocr_text}"
        values = {
            "proje_adi": "Belirtilmemiş",
            "cizim_tarihi": "Belirtilmemiş",
            "calisma_basinci": "Belirtilmemiş",
            "sistem_tipi": "Belirtilmemiş",
            "vana_sayisi": 0,
            "silindir_sayisi": 0,
            "frl_mevcut": False,
            "firma": "Belirtilmemiş"
        }
        
        # Basınç bilgisi
        pressure_match = re.search(r"(\d+(?:\.\d+)?[-\s]*\d*)\s*bar", combined_text, re.IGNORECASE)
        if pressure_match:
            values["calisma_basinci"] = f"{pressure_match.group(1)} bar"
        
        # Vana sayısı
        vana_patterns = [r"(?:VUVG|Y\d+|V\d+|SV\d+)", r"(?:5/2|4/2|3/2|2/2)", r"(?:valf|valve|vana)"]
        vana_matches = []
        for pattern in vana_patterns:
            matches = re.findall(pattern, combined_text, re.IGNORECASE)
            vana_matches.extend(matches)
        
        values["vana_sayisi"] = len(set(vana_matches)) if vana_matches else 0
        
        # Silindir sayısı
        silindir_patterns = [r"(?:C\d+|CYL\d*|SIL\d*)", r"(?:silindir|cylinder|piston)", r"(?:MGF|actuator)"]
        silindir_matches = []
        for pattern in silindir_patterns:
            matches = re.findall(pattern, combined_text, re.IGNORECASE)
            silindir_matches.extend(matches)
        
        values["silindir_sayisi"] = len(set(silindir_matches)) if silindir_matches else 0
        
        # FRL kontrolü
        values["frl_mevcut"] = bool(re.search(r"(?:FRL|MS6-LFR|FILTRE|REGULATOR)", combined_text, re.IGNORECASE))
        
        return values

    def extract_project_information(self, text: str, ocr_text: str) -> Dict[str, Any]:
        """Proje bilgilerini çıkarır (Türkçe ve İngilizce destekli)"""
        combined_text = f"{text} {ocr_text}"
        project_info = {}
        
        # Created by / Oluşturan
        created_patterns = [
            r"Created\s+by[:\s]*([A-ZÇĞÜŞİÖ\s]+(?:[A-ZÇĞÜŞİÖ\s]*[A-ZÇĞÜŞİÖ]))",
            r"Oluşturan[:\s]*([A-ZÇĞÜŞİÖ\s]+(?:[A-ZÇĞÜŞİÖ\s]*[A-ZÇĞÜŞİÖ]))",
            r"Tasarlayan[:\s]*([A-ZÇĞÜŞİÖ\s]+(?:[A-ZÇĞÜŞİÖ\s]*[A-ZÇĞÜŞİÖ]))",
            r"Designer[:\s]*([A-ZÇĞÜŞİÖ\s]+(?:[A-ZÇĞÜŞİÖ\s]*[A-ZÇĞÜŞİÖ]))"
        ]
        for pattern in created_patterns:
            match = re.search(pattern, combined_text, re.IGNORECASE)
            if match:
                project_info["created_by"] = match.group(1).strip()
                break
        
        # Checked by / Kontrol eden
        checked_patterns = [
            r"Checked\s+by[:\s]*([A-ZÇĞÜŞİÖ\s]+(?:[A-ZÇĞÜŞİÖ\s]*[A-ZÇĞÜŞİÖ]))",
            r"Kontrol\s+eden[:\s]*([A-ZÇĞÜŞİÖ\s]+(?:[A-ZÇĞÜŞİÖ\s]*[A-ZÇĞÜŞİÖ]))"
        ]
        for pattern in checked_patterns:
            match = re.search(pattern, combined_text, re.IGNORECASE)
            if match:
                project_info["checked_by"] = match.group(1).strip()
                break
        
        # Tarihler
        date_patterns = [r"(\d{1,2}[./]\d{1,2}[./]\d{4})", r"(\d{4}-\d{1,2}-\d{1,2})"]
        dates_found = []
        for pattern in date_patterns:
            dates = re.findall(pattern, combined_text)
            dates_found.extend(dates)
        
        if dates_found:
            project_info["creation_date"] = dates_found[-1] if dates_found else None
        
        # Circuit Number
        circuit_patterns = [
            r"Circuit\s+Number[:\s]*([A-Z0-9\-\.]+)",
            r"Devre\s+Numarası[:\s]*([A-Z0-9\-\.]+)"
        ]
        for pattern in circuit_patterns:
            match = re.search(pattern, combined_text, re.IGNORECASE)
            if match:
                project_info["circuit_number"] = match.group(1).strip()
                break
        
        # Description
        description_patterns = [
            r"Description[:\s]*([A-ZÇĞÜŞİÖ0-9\s]+(?:PNOMATİK|PNEUMATIC|DİYAGRAM|DIAGRAM)[A-ZÇĞÜŞİÖ0-9\s]*)",
            r"(PRES\s+PNOMATİK\s+DİYAGRAM\s*\d*)"
        ]
        for pattern in description_patterns:
            match = re.search(pattern, combined_text, re.IGNORECASE)
            if match:
                project_info["description"] = match.group(1).strip()
                break
        
        return project_info

    def generate_recommendations(self, analysis_results: Dict, scores: Dict, extracted_values: Dict) -> List[str]:
        """Öneri üretimi"""
        recommendations = []
        
        overall_score = scores["overall_percentage"]
        
        if overall_score >= 80:
            recommendations.append("✅ ÇİZİM KALİTESİ: Mükemmel")
        elif overall_score >= 70:
            recommendations.append("✅ ÇİZİM KALİTESİ: İyi")
        elif overall_score >= 60:
            recommendations.append("⚠️ ÇİZİM KALİTESİ: Orta")
        else:
            recommendations.append("❌ ÇİZİM KALİTESİ: Yetersiz")
        
        recommendations.append(f"📊 Toplam Puan: {overall_score:.1f}/100")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 50:
                recommendations.append(f"❌ {category}: Ciddi eksiklikler var (%{category_score:.1f})")
            elif category_score < 70:
                recommendations.append(f"⚠️ {category}: Geliştirme gerekli (%{category_score:.1f})")
            else:
                recommendations.append(f"✅ {category}: Yeterli (%{category_score:.1f})")
        
        return recommendations

    def get_quality_level(self, percentage: float) -> str:
        """Kalite seviyesi belirleme"""
        if percentage >= 90:
            return "MÜKEMMEL"
        elif percentage >= 80:
            return "ÇOK İYİ"
        elif percentage >= 70:
            return "İYİ"
        elif percentage >= 60:
            return "ORTA"
        elif percentage >= 40:
            return "YETERSIZ"
        else:
            return "KÖTÜ"

    def analyze_file(self, file_path: str) -> Dict[str, Any]:
        """Ana dosya analiz fonksiyonu"""
        logger.info(f"Dosya analiz ediliyor: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        text = ""
        ocr_text = ""
        visual_detections = []
        
        try:
            if file_ext == '.pdf':
                text = self.extract_text_from_pdf(file_path)
                try:
                    doc = fitz.open(file_path)
                    for page_num in range(min(3, len(doc))):
                        page = doc[page_num]
                        pix = page.get_pixmap()
                        img_data = pix.pil_tobytes(format="PNG")
                        img = Image.open(io.BytesIO(img_data))
                        img_array = np.array(img)
                        temp_path = f"temp_page_{page_num}.png"
                        cv2.imwrite(temp_path, cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR))
                        ocr_text += self.extract_text_from_image(temp_path) + " "
                        visual_detections.extend(self.detect_visual_components(temp_path))
                        os.remove(temp_path)
                    doc.close()
                except:
                    pass
            
            elif file_ext in ['.docx', '.doc']:
                text = self.extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                text = self.extract_text_from_txt(file_path)
                    
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                ocr_text = self.extract_text_from_image(file_path)
                visual_detections = self.detect_visual_components(file_path)
            else:
                return {"error": f"Desteklenmeyen dosya formatı: {file_ext}"}
            
            if not text and not ocr_text:
                return {"error": "Dosyadan metin çıkarılamadı"}
            
            analysis_results = {}
            for category in self.pneumatic_criteria_weights.keys():
                analysis_results[category] = self.analyze_criteria(text, ocr_text, visual_detections, category)
            
            scores = self.calculate_scores(analysis_results)
            extracted_values = self.extract_specific_values(text, ocr_text)
            project_info = self.extract_project_information(text, ocr_text)
            recommendations = self.generate_recommendations(analysis_results, scores, extracted_values)
            
            report = {
                "analysis_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "file_info": {
                    "file_path": file_path,
                    "file_type": file_ext,
                    "text_length": len(text),
                    "ocr_text_length": len(ocr_text),
                    "visual_components": len(visual_detections)
                },
                "project_information": project_info,
                "extracted_values": extracted_values,
                "category_analyses": analysis_results,
                "scoring": scores,
                "recommendations": recommendations,
                "summary": {
                    "total_score": scores["total_score"],
                    "percentage": scores["overall_percentage"],
                    "status": "GEÇERLİ" if scores["overall_percentage"] >= 70 else "GEÇERSİZ",
                    "quality_level": self.get_quality_level(scores["overall_percentage"])
                }
            }
            
            return report
            
        except Exception as e:
            logger.error(f"Analiz hatası: {e}")
            return {"error": str(e)}
        
# ============================================
# HELPER FUNCTIONS (Server Validasyon)
# ============================================
def validate_document_server(text):
    """Server kodunda doküman validasyonu - Pnömatik Devre Şeması için"""
    critical_terms = [
        ["pnömatik", "pnomatik", "pneumatic", "hava", "air", "basınçlı hava", "compressed air"],
        ["silindir", "cylinder", "valf", "valve", "vana", "frl", "lubricator", "regulator", "filter"],
        ["basınç", "pressure", "psi", "bar", "debi", "flow", "cfm", "l/min"],
        ["kontrol", "control", "yön kontrol", "directional control", "hız kontrol", "speed control"],
        ["iso 5599", "5599", "iso 1219", "sembol", "symbol", "bağlantı", "connection", "port"]
    ]
    
    category_found = [
        any(re.search(rf"\b{term}\b", text, re.IGNORECASE) for term in category)
        for category in critical_terms
    ]
    
    valid_categories = sum(category_found)
    logger.info(f"Doküman validasyonu: {valid_categories}/5 kritik kategori bulundu")
    
    return valid_categories >= 4


def check_strong_keywords_first_pages(filepath):
    """İlk 1-2 sayfada özgü kelimeleri OCR ile ara - Pnömatik Devre Şeması için"""
    strong_keywords = [
        "pnömatik", "pnomatik", "pneumatic", "lubricator", "inflate", "psi", "bar", "regis", "r102", "regulator", "dump valve"
    ]
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=400, first_page=1, last_page=1)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 3 -l tur+eng')
            all_text += text.lower() + " "
        
        found_keywords = [kw for kw in strong_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa kontrol: {len(found_keywords)} özgü kelime")
        return len(found_keywords) >= 1
        
    except Exception as e:
        logger.warning(f"İlk sayfa kontrol hatası: {e}")
        return False


def check_excluded_keywords_first_pages(filepath):
    """İlk 1-2 sayfada istenmeyen rapor türlerinin kelimelerini ara"""
    excluded_keywords = [
        "aydınlatma", "lighting", "illumination", "lux", "lümen", "lumen", "ts en 12464", "en 12464", "ışık", "ışık şiddeti",
        "hidrolik", "HİDROLİK", "hydraulic", "hidrolik yağ", "hydraulic oil", "iso 1219", "1219",
        "hrc", "cobot", "robot", "çarpışma", "collaborative", "kolaboratif", "sd conta",
        "elektrik", "devre", "şema", "circuit", "electrical", "voltage", "amper", "ohm","enclosure","wrp-","light curtain","contactors","controller",
        "espe",
        "gürültü", "noise", "ses", "sound", "decibel", "db", "akustik", "acoustic",
        "kullanma", "kılavuz", "manual", "instruction", "talimat", "guide", "kılavuzu",
        "loto",
        "lvd", "TOPRAKLAMA SÜREKLİLİK",  "topraklama süreklilik", "TOPRAKLAMA İLETKENLERİ", "topraklama iletkenleri",
        "uygunluk", "beyan", "muayene", "conformity", "declaration", "declare",
        "isg", "periyodik", "kontrol", "periodic", "inspection", "denetim",
        "montaj", "assembly",
        "topraklama direnci", "grounding", "earthing", "60204", "topraklama","TOPRAKLAMA DİRENCİ",
        "bakım", "maintenance", "servis", "service","bakim","MAINTENCE",
        "titreşim", "vibration","TİTREŞİM",
        "AT TİP", "at tip", "ec type", "SERTİFİKA", "sertifika", "certificate"
    ]
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=1)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 3 -l tur+eng')
            all_text += text.lower() + " "
        
        found_excluded = [kw for kw in excluded_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa excluded kontrol: {len(found_excluded)} istenmeyen kelime")
        return len(found_excluded) >= 1
        
    except Exception as e:
        logger.warning(f"İlk sayfa excluded kontrol hatası: {e}")
        return False


def get_conclusion_message_pneumatic(status, percentage):
    """Sonuç mesajını döndür - Pnömatik için"""
    if status == "PASS":
        return f"Pnömatik devre şeması ISO 5599 ve ilgili standartlara uygun ve teknik açıdan yeterlidir (%{percentage:.0f})"
    elif status == "CONDITIONAL":
        return f"Pnömatik devre şeması kabul edilebilir ancak bazı eksiklikler var (%{percentage:.0f})"
    else:
        return f"Pnömatik devre şeması standartlara uygun değil, kapsamlı revizyon gerekli (%{percentage:.0f})"


def get_main_issues_pneumatic(analysis_result):
    """Ana sorunları listele - Pnömatik için"""
    issues = []
    
    if 'scoring' in analysis_result and 'category_scores' in analysis_result['scoring']:
        for category, score_data in analysis_result['scoring']['category_scores'].items():
            if isinstance(score_data, dict) and score_data.get('percentage', 0) < 50:
                issues.append(f"{category} kategorisinde ciddi eksiklikler")
    
    if not issues and analysis_result['summary']['total_score'] < 50:
        issues = [
            "Pnömatik semboller ISO 5599 standardına uygun değil",
            "FRL ünitesi ve hava hazırlama eksik",
            "Basınç ve debi değerleri eksik veya hatalı",
            "Vana tipleri ve kontrol sistemleri yetersiz",
            "Güvenlik elemanları ve acil durdurma eksik"
        ]
    
    return issues[:4]


# ============================================
# FLASK SERVİS KATMANI - CONFIGURATION
# ============================================
app = Flask(__name__)

UPLOAD_FOLDER = 'temp_uploads_pnomatic'
ALLOWED_EXTENSIONS = {'pdf', 'jpg', 'jpeg', 'png', 'docx', 'doc', 'txt'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ============================================
# FLASK SERVİS KATMANI - API ENDPOINTS
# ============================================
@app.route('/api/pnomatic-control', methods=['POST'])
def analyze_pnomatic_control():
    """Pnömatik Devre Şeması analiz API endpoint'i - 3 Aşamalı Validasyon"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided', 'message': 'Lütfen bir pnömatik devre şeması sağlayın'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type', 'message': 'Sadece PDF, JPG, JPEG, PNG, DOCX, TXT dosyaları kabul edilir'}), 400

        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            logger.info(f"Pnömatik Devre Şeması kontrol ediliyor: {filename}")

            analyzer = PneumaticCircuitAnalyzer()
            file_ext = os.path.splitext(filepath)[1].lower()

            # 3 AŞAMALI KONTROL (sadece PDF için)
            if file_ext == '.pdf':
                logger.info("Aşama 1: İlk sayfa pnömatik özgü kelime kontrolü...")
                if check_strong_keywords_first_pages(filepath):
                    logger.info("✅ Aşama 1 geçti")
                else:
                    logger.info("Aşama 2: İlk sayfa excluded kelime kontrolü...")
                    if check_excluded_keywords_first_pages(filepath):
                        logger.info("❌ Aşama 2'de excluded kelimeler bulundu")
                        try:
                            os.remove(filepath)
                        except:
                            pass
                        return jsonify({
                            'error': 'Invalid document type',
                            'message': 'Bu dosya pnömatik devre şeması değil'
                        }), 400
                    else:
                        logger.info("Aşama 3: Tam doküman critical terms kontrolü...")
                        try:
                            with open(filepath, 'rb') as pdf_file:
                                pdf_reader = PyPDF2.PdfReader(pdf_file)
                                text = "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
                            
                            if not text or len(text.strip()) < 50:
                                try:
                                    os.remove(filepath)
                                except:
                                    pass
                                return jsonify({
                                    'error': 'Text extraction failed',
                                    'message': 'Dosyadan yeterli metin çıkarılamadı'
                                }), 400
                            
                            if not validate_document_server(text):
                                try:
                                    os.remove(filepath)
                                except:
                                    pass
                                return jsonify({
                                    'error': 'Invalid document type',
                                    'message': 'Yüklediğiniz dosya pnömatik devre şeması değil!',
                                    'details': {'filename': filename, 'document_type': 'NOT_PNEUMATIC_CIRCUIT'}
                                }), 400
                                
                        except Exception as e:
                            logger.error(f"Aşama 3 hatası: {e}")
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({'error': 'Analysis failed'}), 500

            elif file_ext in ['.docx', '.doc', '.txt']:
                # DOCX/TXT için sadece tam doküman kontrolü
                logger.info(f"DOCX/TXT dosyası için tam doküman kontrolü: {file_ext}")
                text = ""
                if file_ext in ['.docx', '.doc']:
                    text = analyzer.extract_text_from_docx(filepath)
                elif file_ext == '.txt':
                    text = analyzer.extract_text_from_txt(filepath)
                
                if not text or len(text.strip()) == 0:
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Text extraction failed',
                        'message': 'Dosyadan yeterli metin çıkarılamadı'
                    }), 400
                
                if not validate_document_server(text):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Invalid document type',
                        'message': 'Yüklediğiniz dosya pnömatik devre şeması değil!',
                        'details': {'filename': filename, 'document_type': 'NOT_PNEUMATIC_CIRCUIT'}
                    }), 400

            elif file_ext in ['.jpg', '.jpeg', '.png']:
                if not tesseract_available:
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'OCR not available',
                        'message': 'Tesseract OCR kurulu değil',
                        'details': {'tesseract_error': tesseract_info}
                    }), 500

            logger.info(f"Pnömatik devre şeması doğrulandı, analiz başlatılıyor: {filename}")
            analysis_result = analyzer.analyze_file(filepath)
            
            try:
                os.remove(filepath)
            except:
                pass

            if 'error' in analysis_result:
                return jsonify({'error': 'Analysis failed', 'message': analysis_result['error']}), 400

            overall_percentage = analysis_result['summary']['percentage']
            status = "PASS" if overall_percentage >= 70 else "FAIL"
            
            response_data = {
                'analysis_date': analysis_result.get('analysis_date'),
                'analysis_id': f"pneumatic_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                'category_scores': {},
                'extracted_values': analysis_result.get('extracted_values', {}),
                'file_type': 'PNOMATIK_DEVRE_SEMASI',
                'filename': filename,
                'overall_score': {
                    'percentage': round(overall_percentage, 2),
                    'total_points': analysis_result['summary']['total_score'],
                    'max_points': 100,
                    'status': status,
                    'status_tr': 'GEÇERLİ' if status == "PASS" else 'GEÇERSİZ',
                    'quality_level': analysis_result['summary']['quality_level']
                },
                'project_information': analysis_result.get('project_information', {}),
                'recommendations': analysis_result.get('recommendations', []),
                'summary': {
                    'is_valid': status == "PASS",
                    'conclusion': get_conclusion_message_pneumatic(status, overall_percentage),
                    'main_issues': [] if status == "PASS" else get_main_issues_pneumatic(analysis_result)
                }
            }
            
            if 'scoring' in analysis_result and 'category_scores' in analysis_result['scoring']:
                for category, score_data in analysis_result['scoring']['category_scores'].items():
                    if isinstance(score_data, dict):
                        response_data['category_scores'][category] = {
                            'score': score_data.get('normalized', 0),
                            'max_score': score_data.get('max_weight', 0),
                            'percentage': score_data.get('percentage', 0),
                            'status': 'PASS' if score_data.get('percentage', 0) >= 70 else 'FAIL'
                        }

            return jsonify({
                'success': True,
                'message': 'Pnömatik Devre Şeması başarıyla analiz edildi',
                'analysis_service': 'pneumatic_circuit',
                'data': response_data
            })

        except Exception as analysis_error:
            try:
                if 'filepath' in locals():
                    os.remove(filepath)
            except:
                pass
            logger.error(f"Analiz hatası: {str(analysis_error)}")
            return jsonify({'error': 'Analysis failed', 'message': str(analysis_error)}), 500

    except Exception as e:
        logger.error(f"API endpoint hatası: {str(e)}")
        return jsonify({'error': 'Server error', 'message': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'Pneumatic Circuit Diagram Analyzer API',
        'version': '1.0.0',
        'tesseract_available': tesseract_available,
        'report_type': 'PNOMATIK_DEVRE_SEMASI'
    })


@app.route('/', methods=['GET'])
def index():
    """Ana sayfa - API bilgileri"""
    return jsonify({
        'service': 'Pneumatic Circuit Diagram Analyzer API',
        'version': '1.0.0',
        'description': 'Pnömatik Devre Şemalarını analiz eden REST API servisi',
        'endpoints': {
            'POST /api/pnomatic-control': 'Pnömatik devre şeması analizi',
            'GET /api/health': 'Servis sağlık kontrolü',
            'GET /': 'Bu bilgi sayfası'
        }
    })


# ============================================
# APPLICATION ENTRY POINT (Azure-Friendly)
# ============================================
if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("Pnömatik Devre Şeması Analiz Servisi")
    logger.info("=" * 60)
    
    port = int(os.environ.get('PORT', 8010))
    
    logger.info(f"🚀 Servis başlatılıyor - Port: {port}")
    logger.info("=" * 60)
    
    app.run(host='0.0.0.0', port=port, debug=False)