# ============================================
# IMPORTS
# ============================================
import os
import json
from datetime import datetime, timedelta
import re
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
import pandas as pd
from dataclasses import dataclass, asdict
import logging
import cv2
import numpy as np
from PIL import Image
import pdf2image
import pytesseract

from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename

# ============================================
# DATABASE IMPORTS (YENİ)
# ============================================
from flask import current_app
from database import db, init_db
from db_loader import load_service_config

# ============================================
# CONFIG IMPORT (YENİ)
# ============================================
from config import Config

# ============================================
# LOGGING CONFIGURATION
# ============================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================
# LANGUAGE DETECTION (Optional)
# ============================================
try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False
    logger.warning("langdetect modülü bulunamadı - dil tespiti devre dışı")


# ============================================
# ANALİZ SINIFI - DATA CLASSES
# ============================================
@dataclass
class VibrationCriteria:
    """Mekanik Titreşim Ölçüm Raporu kriterleri veri sınıfı"""
    rapor_kimlik_bilgileri: Dict[str, Any]
    olcum_ortam_makine_calisan: Dict[str, Any]
    olcum_cihazi_kalibrasyon: Dict[str, Any]
    titresim_turleri_yasal: Dict[str, Any]
    olcum_metodolojisi_standartlar: Dict[str, Any]
    olcum_sonuclari_analizler: Dict[str, Any]
    degerlendirme_yorum_onlemler: Dict[str, Any]
    ekler_gorseller: Dict[str, Any]


@dataclass
class VibrationAnalysisResult:
    """Mekanik Titreşim Ölçüm Raporu analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]


# ============================================
# ANALİZ SINIFI - MAIN ANALYZER
# ============================================
class VibrationReportAnalyzer:
    """Mekanik Titreşim Ölçüm Raporu analiz sınıfı"""
    
    def __init__(self, app=None):
        logger.info("Mekanik Titreşim Ölçüm Raporu analiz sistemi başlatılıyor...")
        
        # Flask app context varsa DB'den yükle, yoksa boş başlat
        if app:
            with app.app_context():
                try:
                    config = load_service_config('vibration_report')
                    
                    # DB'den yüklenen veriler
                    self.criteria_weights = config.get('criteria_weights', {})
                    self.criteria_details = config.get('criteria_details', {})
                    self.pattern_definitions = config.get('pattern_definitions', {})
                    self.validation_keywords = config.get('validation_keywords', {})
                    self.category_actions = config.get('category_actions', {})
                    
                    logger.info(f"✅ Veritabanından yüklendi: {len(self.criteria_weights)} kategori")
                    
                except Exception as e:
                    logger.error(f"⚠️ Veritabanından yükleme başarısız: {e}")
                    logger.warning("⚠️ Fallback: Boş config kullanılıyor")
                    self.criteria_weights = {}
                    self.criteria_details = {}
                    self.pattern_definitions = {}
                    self.validation_keywords = {}
                    self.category_actions = {}
        else:
            # Flask app yoksa boş başlat (eski davranış)
            logger.warning("⚠️ Flask app context yok, boş config kullanılıyor")
            self.criteria_weights = {}
            self.criteria_details = {}
            self.pattern_definitions = {}
            self.validation_keywords = {}
            self.category_actions = {}
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni Türkçe'ye çevir - şimdilik devre dışı"""
        if source_lang != 'tr':
            logger.info(f"Tespit edilen dil: {source_lang.upper()} - Çeviri yapılmıyor, orijinal metin kullanılıyor")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosya formatına göre metin çıkarma"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.doc':
            logger.warning("DOC formatı için DOCX'e dönüştürme gerekiyor veya OCR kullanılacak")
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatı: {file_ext}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den sadece PyPDF2 ile metin çıkarma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF analizi başlatılıyor - Toplam sayfa: {total_pages}")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = page_text.replace('|', ' ')
                    page_text = page_text.strip()
                    
                    all_text += page_text + "\n"
                
                # Metni temizle
                all_text = all_text.replace('—', '-')
                all_text = all_text.replace('"', '"').replace('"', '"')
                all_text = all_text.replace('´', "'")
                all_text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]+', ' ', all_text)
                all_text = all_text.strip()
                
                logger.info(f"✅ PDF analizi tamamlandı:")
                logger.info(f"   📊 Toplam metin uzunluğu: {len(all_text):,} karakter")
                logger.info(f"   📄 Toplam sayfa: {total_pages}")
                
                return all_text
                
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Paragrafları oku
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Tablolardan metin çıkar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            logger.info(f"DOCX'den {len(text)} karakter metin çıkarıldı")
            return text
            
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasından metin çıkarma"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text:
                # UTF-8 başarısız olursa diğer encoding'leri dene
                encodings = ['cp1254', 'iso-8859-9', 'latin1']
                for encoding in encodings:
                    try:
                        with open(txt_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        if text:
                            break
                    except:
                        continue
            
            logger.info(f"TXT'den {len(text)} karakter metin çıkarıldı")
            return text.strip()
            
        except Exception as e:
            logger.error(f"TXT metin çıkarma hatası: {e}")
            return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, VibrationAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                # Scoring algoritması
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
                    
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = VibrationAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={
                    "pattern_used": pattern,
                    "matches_count": len(matches) if matches else 0
                }
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, VibrationAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """Titreşim raporuna özgü değerleri çıkar"""
        values = {
            "rapor_numarasi": "Bulunamadı",
            "olcum_tarihi": "Bulunamadı",
            "rapor_tarihi": "Bulunamadı",
            "olcum_cihazi": "Bulunamadı",
            "firma_adi": "Bulunamadı",
            "titresim_turu": "Bulunamadı",
            "a8_degeri": "Bulunamadı",
            "yasal_uygunluk": "Bulunamadı"
        }

        # DB'den pattern'leri al
        report_no_patterns = self.pattern_definitions.get('report_no_patterns', {}).get('rapor_numarasi', [])
        measurement_date_patterns = self.pattern_definitions.get('measurement_date_patterns', {}).get('olcum_tarihi', [])
        report_date_patterns = self.pattern_definitions.get('report_date_patterns', {}).get('rapor_tarihi', [])
        device_patterns = self.pattern_definitions.get('device_patterns', {}).get('olcum_cihazi', [])
        company_patterns = self.pattern_definitions.get('company_patterns', {}).get('firma_adi', [])
        vibration_type_patterns = self.pattern_definitions.get('vibration_type_patterns', {}).get('titresim_turu', [])
        a8_patterns = self.pattern_definitions.get('a8_patterns', {}).get('a8_degeri', [])
        compliance_patterns = self.pattern_definitions.get('compliance_patterns', {}).get('yasal_uygunluk', [])
            
        # RAPOR NUMARASI
        for pattern in report_no_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 3 <= len(result) <= 20:
                    values["rapor_numarasi"] = result
                    break
    
        # ÖLÇÜM TARİHİ
        for pattern in measurement_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["olcum_tarihi"] = match.group(1)
                break
        
        # RAPOR TARİH
        for pattern in report_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["rapor_tarihi"] = match.group(1)
                break
        
        if values["rapor_tarihi"] == "Bulunamadı" and values["olcum_tarihi"] != "Bulunamadı":
            values["rapor_tarihi"] = "Rapor tarihi ayrı belirtilmemiş"
        
        # ÖLÇÜM CİHAZI
        for pattern in device_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 2 <= len(result) <= 35:
                    values["olcum_cihazi"] = result
                    break
        
        # FİRMA ADI
        for pattern in company_patterns:
            matches = re.findall(pattern, text)
            for result in matches:
                result = result.strip()
                if 3 <= len(result) <= 60:
                    values["firma_adi"] = result
                    break
            if values["firma_adi"] != "Bulunamadı":
                break
        
        # TİTREŞİM TÜRÜ    
        for pattern in vibration_type_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip().upper()
                if "EL" in result or "HAND" in result or "HAV" in result:
                    values["titresim_turu"] = "El-Kol Titreşimi"
                elif "VÜCUT" in result or "BODY" in result or "WBV" in result:
                    values["titresim_turu"] = "Bütün Vücut Titreşimi"
                break
        
        # A(8) DEĞERİ
        for pattern in a8_patterns:
            match = re.search(pattern, text)
            if match:
                values["a8_degeri"] = match.group(1).replace(',', '.')
                break
        
        # YASAL UYGUNLUK
        for pattern in compliance_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(0).strip().upper()
                if any(word in result for word in ["UYGUN", "SUITABLE", "CONFORM", "GEÇERLİ", "VALID", "PASS"]):
                    values["yasal_uygunluk"] = "UYGUN"
                elif any(word in result for word in ["UYGUNSUZ", "NOT SUITABLE", "NON-CONFORM", "GEÇERSİZ", "INVALID", "FAIL", "AŞIL", "EXCEED"]):
                    values["yasal_uygunluk"] = "UYGUNSUZ"
                break
        
        return values

    def check_date_validity(self, measurement_date: str, report_date: str) -> Tuple[bool, str]:
        """Ölçüm ve rapor tarihlerini bugünkü tarih ile kontrol et (1 yıl kuralı)"""
        
        if measurement_date == "Bulunamadı" and report_date == "Bulunamadı":
            return False, "RAPOR GEÇERSİZ: Ne ölçüm ne de rapor tarihi bulunamadı"
        
        dates_to_check = []
        
        if measurement_date != "Bulunamadı":
            dates_to_check.append(("Ölçüm", measurement_date))
        
        if report_date != "Bulunamadı" and report_date != "Rapor tarihi ayrı belirtilmemiş":
            dates_to_check.append(("Rapor", report_date))
        
        if not dates_to_check:
            return False, "RAPOR GEÇERSİZ: Geçerli tarih bulunamadı"
        
        try:
            date_formats = ['%d/%m/%Y', '%d.%m.%Y', '%d-%m-%Y', '%d/%m/%y', '%d.%m.%y', '%d-%m-%y']
            today = datetime.now()
            
            for date_type, date_str in dates_to_check:
                parsed_date = None
                for fmt in date_formats:
                    try:
                        parsed_date = datetime.strptime(date_str, fmt)
                        break
                    except:
                        continue
                
                if not parsed_date:
                    return False, f"RAPOR GEÇERSİZ: {date_type} tarihi formatı tanınmadı ({date_str})"
                
                if parsed_date.year < 1950:
                    parsed_date = parsed_date.replace(year=parsed_date.year + 2000)
                
                diff = abs((today - parsed_date).days)
                
                if diff > 365:
                    return False, f"RAPOR GEÇERSİZ: {date_type} tarihi 1 yıldan eski ({diff} gün önce - {date_str})"
            
            checked_dates = [f"{dt[0]}: {dt[1]}" for dt in dates_to_check]
            return True, f"Tüm tarihler geçerli - {', '.join(checked_dates)}"
            
        except Exception as e:
            return False, f"RAPOR GEÇERSİZ: Tarih kontrolü yapılamadı - {e}"

    def generate_improvement_actions(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Dinamik iyileştirme önerileri oluştur - DB'den gelen actions ile"""
        actions = []
        
        sorted_categories = sorted(
            scores["category_scores"].items(), 
            key=lambda x: x[1]["percentage"]
        )
        
        # DB'den category_actions al (self.category_actions zaten __init__'te yüklendi)
        category_actions = self.category_actions
        
        # Eğer DB'den veri gelmediyse, boş liste döndür
        if not category_actions:
            logger.warning("⚠️ Category actions bulunamadı, boş öneri listesi döndürülüyor")
            if scores["percentage"] < 50:
                return ["ÖNCELİK: Rapor yapısını ve içeriğini kapsamlı olarak yeniden düzenleyiniz"]
            return []
        
        for category, score_data in sorted_categories[:5]:
            if score_data["percentage"] < 70:
                actions.extend(category_actions.get(category, []))
        
        if scores["percentage"] < 50:
            actions.insert(0, "ÖNCELİK: Rapor yapısını ve içeriğini kapsamlı olarak yeniden düzenleyiniz")
        
        return actions
        
    def generate_recommendations(self, analysis_results: Dict, scores: Dict, date_valid: bool = True, date_message: str = "") -> List[str]:
        """Öneriler oluştur"""
        recommendations = []
        
        total_percentage = scores["percentage"]
        
        if not date_valid:
            if "Ne ölçüm ne de rapor tarihi bulunamadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ne ölçüm ne de rapor tarihi bulunamadı")
            elif "Ölçüm tarihi formatı tanınmadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ölçüm tarihi formatı tanınmadı")
            elif "Rapor tarihi formatı tanınmadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Rapor tarihi formatı tanınmadı")
            elif "Ölçüm tarihi 1 yıldan eski" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Ölçüm tarihi 1 yıldan eski")
            elif "Rapor tarihi 1 yıldan eski" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Rapor tarihi 1 yıldan eski")
            elif "Geçerli tarih bulunamadı" in date_message:
                recommendations.append("❌ RAPOR GEÇERSİZ: Geçerli tarih bulunamadı")
            else:
                recommendations.append(f"❌ RAPOR GEÇERSİZ: {date_message}")
            return recommendations
        
        if total_percentage >= 70:
            recommendations.append(f"✅ Mekanik Titreşim Ölçüm Raporu GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ Mekanik Titreşim Ölçüm Raporu GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 40:
                recommendations.append(f"🔴 {category} bölümü yetersiz (%{category_score:.1f})")
                missing_items = [name for name, result in results.items() if not result.found]
                if missing_items:
                    recommendations.append(f"   Eksik: {', '.join(missing_items[:3])}")
            elif category_score < 70:
                recommendations.append(f"🟡 {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"🟢 {category} bölümü yeterli (%{category_score:.1f})")
        
        return recommendations

    def analyze_vibration_report(self, file_path: str) -> Dict[str, Any]:
        """Ana Mekanik Titreşim Ölçüm Raporu analiz fonksiyonu"""
        logger.info("Mekanik Titreşim Ölçüm Raporu analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        if detected_lang != 'tr':
            logger.info(f"{detected_lang.upper()} dilinden Türkçe'ye çeviriliyor...")
            text = self.translate_to_turkish(text, detected_lang)
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        
        date_valid, date_message = self.check_date_validity(
            extracted_values.get("olcum_tarihi", "Bulunamadı"),
            extracted_values.get("rapor_tarihi", "Bulunamadı")
        )
        
        recommendations = self.generate_recommendations(analysis_results, scores, date_valid, date_message)
        improvement_actions = self.generate_improvement_actions(analysis_results, scores)
        
        if not date_valid:
            final_status = "FAIL"
            final_percentage = 0
        else:
            final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
            final_percentage = scores["percentage"]
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "tarih_gecerlilik": {
                "valid": date_valid,
                "message": date_message
            },
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "iyilestirme_eylemleri": improvement_actions,
            "ozet": {
                "toplam_puan": scores["total_score"] if date_valid else 0,
                "yuzde": final_percentage,
                "durum": final_status,
                "rapor_tipi": "MEKANİK_TİTREŞİM_OLCUM_RAPORU",
                "tarih_gecersiz": not date_valid
            }
        }
        
        return report


# ============================================
# HELPER FUNCTIONS (Server Validasyon)
# ============================================
def map_language_code(lang_code):
    """Dil kodunu tam isme çevir"""
    lang_mapping = {
        'tr': 'turkish',
        'en': 'english', 
        'de': 'german',
        'fr': 'french',
        'es': 'spanish',
        'it': 'italian'
    }
    return lang_mapping.get(lang_code, 'turkish')


def validate_document_server(text, validation_keywords):
    """Server kodunda doküman validasyonu - DB'den gelen keywords ile"""
    
    # DB'den gelen critical_terms
    critical_terms_data = validation_keywords.get('critical_terms', [])
    
    # Liste formatına dönüştür (eski format uyumluluğu)
    critical_terms = []
    for item in critical_terms_data:
        if isinstance(item, dict) and 'keywords' in item:
            critical_terms.append(item['keywords'])
        elif isinstance(item, list):
            critical_terms.append(item)
    
    # Eğer DB'den veri gelmediyse, boş validasyon (hata verme)
    if not critical_terms:
        logger.warning("⚠️ Critical terms bulunamadı, validasyon atlanıyor")
        return True  # Varsayılan: geçerli kabul et
    
    category_found = []
    
    for i, category in enumerate(critical_terms):
        found_in_category = False
        for term in category:
            if re.search(rf"\b{term}\b", text, re.IGNORECASE):
                found_in_category = True
                logger.info(f"Kategori {i+1} bulundu: '{term}'")
                break
        category_found.append(found_in_category)
    
    valid_categories = sum(category_found)
    
    logger.info(f"Doküman validasyonu: {valid_categories}/{len(critical_terms)} kritik kategori bulundu")
    
    return valid_categories >= len(critical_terms) - 1  # En az (n-1) kategori bulunmalı


def check_strong_keywords_first_pages(filepath, validation_keywords):
    """İlk 1-2 sayfada özgü kelimeleri OCR ile ara - DB'den gelen keywords ile"""
    
    # DB'den strong keywords al
    strong_keywords = validation_keywords.get('strong_keywords', [])
    
    # Eğer DB'den veri gelmediyse, validasyon atla
    if not strong_keywords:
        logger.warning("⚠️ Strong keywords bulunamadı, validasyon atlanıyor")
        return True  # Varsayılan: geçerli kabul et
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=2)
        
        all_text = ""
        for i, page in enumerate(pages):
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        
        found_keywords = []
        for keyword in strong_keywords:
            if re.search(rf"\b{keyword.lower()}\b", all_text):
                found_keywords.append(keyword)
        
        logger.info(f"İlk sayfa kontrol: {len(found_keywords)} özgü kelime: {found_keywords}")
        return len(found_keywords) >= 2
        
    except Exception as e:
        logger.warning(f"İlk sayfa kontrol hatası: {e}")
        return False


def check_excluded_keywords_first_pages(filepath, validation_keywords):
    """İlk 1-2 sayfada istenmeyen rapor türlerinin kelimelerini ara - DB'den"""
    
    # DB'den excluded keywords al
    excluded_keywords = validation_keywords.get('excluded_keywords', [])
    
    # Eğer DB'den veri gelmediyse, validasyon atla
    if not excluded_keywords:
        logger.warning("⚠️ Excluded keywords bulunamadı, validasyon atlanıyor")
        return False  # Varsayılan: excluded yok kabul et
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=2)
        
        all_text = ""
        for i, page in enumerate(pages):
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        
        logger.info(f"OCR okunan text: {all_text[:500]}...")
        
        found_excluded = []
        for keyword in excluded_keywords:
            if re.search(rf"\b{keyword.lower()}\b", all_text):
                found_excluded.append(keyword)
        
        logger.info(f"İlk sayfa excluded kontrol: {len(found_excluded)} istenmeyen kelime: {found_excluded}")
        return len(found_excluded) >= 2
        
    except Exception as e:
        logger.warning(f"İlk sayfa excluded kontrol hatası: {e}")
        return False


def get_conclusion_message_titresim(status, percentage):
    """Sonuç mesajını döndür - Titreşim için"""
    if status == "PASS":
        return f"Mekanik titreşim ölçüm raporu TS EN ISO 5349 ve TS ISO 2631 standartlarına uygundur (%{percentage:.0f})"
    elif status == "CONDITIONAL":
        return f"Titreşim raporu kabul edilebilir ancak bazı eksiklikler var (%{percentage:.0f})"
    else:
        return f"Titreşim raporu standartlara uygun değil, kapsamlı revizyon gerekli (%{percentage:.0f})"


def get_main_issues_titresim(analysis_result):
    """Ana sorunları listele - Titreşim için"""
    issues = []
    
    for category, score_data in analysis_result['puanlama']['category_scores'].items():
        if score_data['percentage'] < 50:
            issues.append(f"{category} kategorisinde ciddi eksiklikler")
    
    if not issues:
        if analysis_result['puanlama']['total_score'] < 50:
            issues = [
                "A(8) günlük maruziyet değeri eksik",
                "Titreşim ölçüm cihazı kalibrasyon bilgileri eksik",
                "El-kol veya bütün vücut titreşim türü tanımı eksik",
                "Yasal uygunluk değerlendirmesi yapılmamış",
                "TS EN ISO 5349 ve TS ISO 2631 standart kontrolü eksik"
            ]
    
    return issues[:4]


# ============================================
# FLASK SERVİS KATMANI - CONFIGURATION
# ============================================
app = Flask(__name__)

# Database configuration (YENİ)
app.config['SQLALCHEMY_DATABASE_URI'] = Config.SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = Config.SQLALCHEMY_TRACK_MODIFICATIONS
app.config['SQLALCHEMY_ECHO'] = Config.SQLALCHEMY_ECHO

# Upload configuration
UPLOAD_FOLDER = 'temp_uploads_titresim'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

# Create upload folder if it doesn't exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ============================================
# FLASK SERVİS KATMANI - API ENDPOINTS
# ============================================
@app.route('/api/titresim-report', methods=['POST'])
def analyze_titresim_report():
    """Mekanik Titreşim Ölçüm Raporu analiz API endpoint'i - 3 Aşamalı Validasyon"""
    try:
        if 'file' not in request.files:
            return jsonify({
                'error': 'No file provided',
                'message': 'Lütfen analiz edilmek üzere bir mekanik titreşim raporu sağlayın'
            }), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({
                'error': 'No file selected',
                'message': 'Lütfen bir dosya seçin'
            }), 400

        if not allowed_file(file.filename):
            return jsonify({
                'error': 'Invalid file type',
                'message': 'Sadece PDF, DOCX, DOC ve TXT dosyaları kabul edilir'
            }), 400

        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            logger.info(f"Mekanik Titreşim Raporu kontrol ediliyor: {filename}")

            # Create analyzer instance
            analyzer = VibrationReportAnalyzer(app=app) 
            
            # ÜÇ AŞAMALI TİTREŞİM KONTROLÜ
            logger.info(f"Üç aşamalı titreşim kontrolü başlatılıyor: {filename}")
            file_ext = os.path.splitext(filepath)[1].lower()

            if file_ext == '.pdf':
                # PDF için üç aşamalı kontrol (OCR dahil)
                logger.info("Aşama 1: İlk sayfa titreşim özgü kelime kontrolü...")
                if check_strong_keywords_first_pages(filepath, analyzer.validation_keywords):
                    logger.info("✅ Aşama 1 geçti - Titreşim özgü kelimeler bulundu")
                else:
                    logger.info("Aşama 2: İlk sayfa excluded kelime kontrolü...")
                    if check_excluded_keywords_first_pages(filepath, analyzer.validation_keywords):
                        logger.info("❌ Aşama 2'de excluded kelimeler bulundu - Titreşim değil")
                        try:
                            os.remove(filepath)
                        except:
                            pass
                        return jsonify({
                            'error': 'Invalid document type',
                            'message': 'Bu dosya titreşim raporu değil (farklı rapor türü tespit edildi). Lütfen mekanik titreşim ölçüm raporu yükleyiniz.',
                            'details': {
                                'filename': filename,
                                'document_type': 'OTHER_REPORT_TYPE',
                                'required_type': 'MEKANIK_TITRESIM_OLCUM_RAPORU'
                            }
                        }), 400
                    else:
                        # AŞAMA 3: PyPDF2 ile tam doküman kontrolü
                        logger.info("Aşama 3: Tam doküman critical terms kontrolü...")
                        try:
                            import PyPDF2
                            with open(filepath, 'rb') as file:
                                pdf_reader = PyPDF2.PdfReader(file)
                                text = ""
                                for page in pdf_reader.pages:
                                    text += page.extract_text() + "\n"
                            
                            if not text or len(text.strip()) < 50:
                                try:
                                    os.remove(filepath)
                                except:
                                    pass
                                return jsonify({
                                    'error': 'Text extraction failed',
                                    'message': 'Dosyadan yeterli metin çıkarılamadı'
                                }), 400
                            
                            if not validate_document_server(text, analyzer.validation_keywords):
                                try:
                                    os.remove(filepath)
                                except:
                                    pass
                                return jsonify({
                                    'error': 'Invalid document type',
                                    'message': 'Yüklediğiniz dosya mekanik titreşim ölçüm raporu değil! Lütfen geçerli bir titreşim raporu yükleyiniz.',
                                    'details': {
                                        'filename': filename,
                                        'document_type': 'NOT_TITRESIM_REPORT',
                                        'required_type': 'MEKANIK_TITRESIM_OLCUM_RAPORU'
                                    }
                                }), 400
                                
                        except Exception as e:
                            logger.error(f"Aşama 3 hatası: {e}")
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({
                                'error': 'Analysis failed',
                                'message': 'Dosya analizi sırasında hata oluştu'
                            }), 500

            elif file_ext in ['.docx', '.doc', '.txt']:
                # DOCX/TXT için sadece tam doküman kontrolü
                logger.info(f"DOCX/TXT dosyası için tam doküman kontrolü: {file_ext}")
                text = ""
                if file_ext in ['.docx', '.doc']:
                    text = analyzer.extract_text_from_docx(filepath)
                elif file_ext == '.txt':
                    text = analyzer.extract_text_from_txt(filepath)
                
                if not text or len(text.strip()) == 0:
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Text extraction failed',
                        'message': 'Dosyadan yeterli metin çıkarılamadı'
                    }), 400
                
                if not validate_document_server(text):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Invalid document type',
                        'message': 'Yüklediğiniz dosya mekanik titreşim ölçüm raporu değil! Lütfen geçerli bir titreşim raporu yükleyiniz.',
                        'details': {
                            'filename': filename,
                            'document_type': 'NOT_TITRESIM_REPORT',
                            'required_type': 'MEKANIK_TITRESIM_OLCUM_RAPORU'
                        }
                    }), 400

            # Buraya kadar geldiyse titreşim raporu, şimdi analizi yap
            logger.info(f"Titreşim raporu doğrulandı, analiz başlatılıyor: {filename}")
            analysis_result = analyzer.analyze_vibration_report(filepath)
            
            # Clean up the uploaded file
            try:
                os.remove(filepath)
                logger.info(f"Uploaded file {filename} cleaned up")
            except Exception as e:
                logger.warning(f"Could not remove uploaded file {filename}: {e}")

            # Check if analysis was successful
            if 'error' in analysis_result:
                return jsonify({
                    'error': 'Analysis failed',
                    'message': analysis_result['error'],
                    'details': {
                        'filename': filename,
                        'text_length': analysis_result.get('text_length', 0)
                    }
                }), 400

            # Extract key results for API response
            overall_percentage = analysis_result['puanlama']['percentage']
            status = "PASS" if overall_percentage >= 70 else "FAIL"
            
            response_data = {
                'analysis_date': analysis_result.get('analiz_tarihi', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                'analysis_details': {
                    'found_criteria': len([r for results in analysis_result.get('kategori_analizleri', {}).values() 
                                         for r in results.values() if r.found]),
                    'total_criteria': len([r for results in analysis_result.get('kategori_analizleri', {}).values() 
                                         for r in results.values()]),
                    'percentage': round((len([r for results in analysis_result.get('kategori_analizleri', {}).values() 
                                            for r in results.values() if r.found]) / 
                                       max(1, len([r for results in analysis_result.get('kategori_analizleri', {}).values() 
                                                 for r in results.values()])) * 100), 1)
                },
                'analysis_id': f"titresim_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                'category_scores': {},
                'date_validity': 'N/A',
                'extracted_values': analysis_result.get('cikarilan_degerler', {}),
                'file_type': 'MEKANIK_TITRESIM_OLCUM_RAPORU',
                'filename': filename,
                'language_info': {
                    'detected_language': map_language_code(analysis_result['dosya_bilgisi']['detected_language']),
                    'text_length': analysis_result.get('text_length', 0)
                },
                'overall_score': {
                    'percentage': round(overall_percentage, 2),
                    'total_points': analysis_result['puanlama']['total_score'],
                    'max_points': 100,
                    'status': status,
                    'status_tr': 'GEÇERLİ' if status == "PASS" else 'GEÇERSİZ',
                    'text_quality': 'good' if len(analysis_result.get('cikarilan_degerler', {})) > 3 else 'fair'
                },
                'recommendations': analysis_result.get('oneriler', []),
                'summary': {
                    'is_valid': status == "PASS",
                    'conclusion': get_conclusion_message_titresim(status, overall_percentage),
                    'main_issues': [] if status == "PASS" else get_main_issues_titresim(analysis_result)
                }
            }
            
            # Add category scores
            for category, score_data in analysis_result['puanlama']['category_scores'].items():
                response_data['category_scores'][category] = {
                    'score': score_data['normalized'],
                    'max_score': score_data['max_weight'],
                    'percentage': score_data['percentage'],
                    'status': 'PASS' if score_data['percentage'] >= 70 else 'CONDITIONAL' if score_data['percentage'] >= 50 else 'FAIL'
                }

            return jsonify({
                'success': True,
                'message': 'Mekanik Titreşim Ölçüm Raporu başarıyla analiz edildi',
                'analysis_service': 'titresim',
                'service_description': 'Mekanik Titreşim Ölçüm Raporu Analizi',
                'data': response_data
            })

        except Exception as analysis_error:
            # Clean up file on error
            try:
                if 'filepath' in locals():
                    os.remove(filepath)
            except:
                pass
            
            logger.error(f"Analiz hatası: {str(analysis_error)}")
            return jsonify({
                'error': 'Analysis failed',
                'message': f'Titreşim raporu analizi sırasında hata oluştu: {str(analysis_error)}',
                'details': {
                    'error_type': type(analysis_error).__name__,
                    'file_processed': filename if 'filename' in locals() else 'unknown'
                }
            }), 500

    except Exception as e:
        logger.error(f"API endpoint hatası: {str(e)}")
        return jsonify({
            'error': 'Server error',
            'message': f'Sunucu hatası: {str(e)}'
        }), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint - Titreşim için"""
    return jsonify({
        'status': 'healthy',
        'service': 'Mechanical Vibration Report Analyzer API',
        'version': '1.0.0',
        'upload_folder': UPLOAD_FOLDER,
        'max_file_size_mb': app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024),
        'supported_formats': list(ALLOWED_EXTENSIONS),
        'ocr_support': True,
        'report_type': 'MEKANIK_TITRESIM_OLCUM_RAPORU'
    })


@app.route('/api/test-titresim', methods=['GET'])
def test_titresim_analysis():
    """Test endpoint for debugging - Titreşim için"""
    try:
        analyzer = VibrationReportAnalyzer()
        test_info = {
            'analyzer_initialized': True,
            'available_methods': [method for method in dir(analyzer) if not method.startswith('_')],
            'criteria_categories': list(analyzer.criteria_weights.keys()),
            'criteria_weights': analyzer.criteria_weights,
            'total_possible_score': sum(analyzer.criteria_weights.values()),
            'ocr_support': True
        }
        
        return jsonify({
            'success': True,
            'message': 'Test başarılı',
            'data': test_info
        })
    except Exception as e:
        return jsonify({
            'error': 'Test failed',
            'message': str(e)
        }), 500


@app.route('/api/titresim-categories', methods=['GET'])
def get_titresim_categories():
    """Titreşim analiz kategorilerini döndür"""
    try:
        analyzer = VibrationReportAnalyzer()
        return jsonify({
            'success': True,
            'data': {
                'categories': analyzer.criteria_weights,
                'total_weight': sum(analyzer.criteria_weights.values()),
                'criteria_details': {
                    category: list(criteria.keys()) 
                    for category, criteria in analyzer.criteria_details.items()
                },
                'standard_reference': 'TS EN ISO 5349 ve TS ISO 2631'
            }
        })
    except Exception as e:
        return jsonify({
            'error': 'Failed to get categories',
            'message': str(e)
        }), 500


@app.route('/', methods=['GET'])
def index():
    """Ana sayfa - API bilgileri - Titreşim için"""
    return jsonify({
        'service': 'Mechanical Vibration Report Analyzer API',
        'version': '1.0.0',
        'description': 'Mekanik Titreşim Ölçüm Raporlarını analiz eden REST API servisi',
        'endpoints': {
            'POST /api/titresim-report': 'Titreşim raporu analizi',
            'GET /api/health': 'Servis sağlık kontrolü',
            'GET /api/test-titresim': 'Test analizi',
            'GET /api/titresim-categories': 'Analiz kategorileri',
            'GET /': 'Bu bilgi sayfası'
        },
        'usage': {
            'upload_format': 'multipart/form-data',
            'file_field': 'file',
            'supported_types': list(ALLOWED_EXTENSIONS),
            'max_size_mb': app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024)
        },
        'analysis_categories': [
            'Genel Bilgiler',
            'Test Koşulları ve Senaryo Tanımı',
            'Ölçüm Noktaları ve Metodoloji',
            'Titreşim Ölçüm Sonuçları',
            'Sınır Değerlerle Karşılaştırma',
            'Risk Değerlendirmesi ve Sonuç',
            'Öneriler ve Önlemler',
            'Ekler ve Kalibrasyon Belgeleri'
        ],
        'scoring': {
            'PASS': '≥70% - Standarta uygun',
            'CONDITIONAL': '50-69% - Kabul edilebilir',
            'FAIL': '<50% - Standarta uygun değil'
        },
        'example_curl': 'curl -X POST -F "file=@titresim_raporu.pdf" http://localhost:8000/api/titresim-report'
    })

# ============================================
# DATABASE INITIALIZATION
# ============================================
with app.app_context():
    db.init_app(app)


# ============================================
# APPLICATION ENTRY POINT (Azure-Friendly)
# ============================================
if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("Mekanik Titreşim Ölçüm Raporu Analiz Servisi")
    logger.info("=" * 60)
    logger.info(f"📁 Upload klasörü: {UPLOAD_FOLDER}")
    logger.info(f"📊 Desteklenen formatlar: {', '.join(ALLOWED_EXTENSIONS)}")
    logger.info(f"📏 Maksimum dosya boyutu: {app.config['MAX_CONTENT_LENGTH'] // (1024 * 1024)} MB")
    logger.info("")
    logger.info("🔗 API Endpoints:")
    logger.info("  POST /api/titresim-report - Titreşim raporu analizi")
    logger.info("  GET /api/health - Servis sağlık kontrolü")
    logger.info("  GET /api/test-titresim - Test analizi")
    logger.info("  GET /api/titresim-categories - Analiz kategorileri")
    logger.info("  GET / - API bilgileri")
    logger.info("=" * 60)
    
    # Azure App Service PORT environment variable (default: 8000)
    port = int(os.environ.get('PORT', 8000))
    
    logger.info(f"🚀 Servis başlatılıyor - Port: {port}")
    logger.info("=" * 60)
    
    app.run(
        host='0.0.0.0',
        port=port,
        debug=False  # Production için debug=False
    )