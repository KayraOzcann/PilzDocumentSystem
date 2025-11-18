import re
import os
import json
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
import pandas as pd
from dataclasses import dataclass, asdict
import logging

try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class MaintenanceCriteria:
    """Bakım Talimatları kriterleri veri sınıfı"""
    genel_bilgiler: Dict[str, Any]
    guvenlik_onlemleri: Dict[str, Any]
    bakim_turleri_operasyon: Dict[str, Any]
    adim_adim_bakim_islemleri: Dict[str, Any]
    teknik_veriler: Dict[str, Any]
    yedek_parca_sarf_malzemeler: Dict[str, Any]
    kayit_izlenebilirlik: Dict[str, Any]
    cevre_atik_yonetimi: Dict[str, Any]
    ekler: Dict[str, Any]

@dataclass
class MaintenanceAnalysisResult:
    """Bakım Talimatları analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

class MaintenanceReportAnalyzer:
    """Bakım Talimatları rapor analiz sınıfı"""
    
    def __init__(self):
        logger.info("Bakım Talimatları analiz sistemi başlatılıyor...")
        
        self.criteria_weights = {
            "Genel Bilgiler": 10,
            "Güvenlik Önlemleri": 15,
            "Bakım Türleri ve Operasyon Çeşitleri": 20,
            "Adım Adım Bakım İşlemleri": 10,
            "Teknik Veriler": 10,
            "Yedek Parça ve Sarf Malzemeleri": 15,
            "Kayıt ve İzlenebilirlik": 10,
            "Çevre ve Atık Yönetimi": 5,
            "Ekler": 5
        }
        
        self.criteria_details = {
            "Genel Bilgiler": {
                "makine_tanimi": {"pattern": r"(?:Makine|Machine|Model|Seri\s*No|Serial\s*Number|Üretici|Manufacturer|Marka|Brand|Ekipman|Equipment|Cihaz|Device|Sistem|System|Ürün|Product|Tip|Type)", "weight": 4},
                "belge_kapsami": {"pattern": r"(?:Kapsam|Scope|Geçerli|Valid|Bu\s*talimat|This\s*instruction|Coverage|Amaç|Purpose|Hedef|Target|Kullan|Use|Uygula|Apply|Alan|Field|Bölüm|Section)", "weight": 3},
                "yetkili_personel": {"pattern": r"(?:Yetkili|Authorized|Teknisyen|Technician|Eğitim|Training|Yetkin|Qualified|Personel|Personnel|Sorumlu|Responsible|Operatör|Operator|Uzman|Expert|Mühendis|Engineer)", "weight": 3}
            },
            "Güvenlik Önlemleri": {
                "genel_guvenlik_uyari": {"pattern": r"(?:Güvenlik\s*Uyar[ıi]|Safety\s*Warning|Elektrik\s*Çarp|Electric\s*Shock|Hareketli\s*Parça|Moving\s*Parts|S[ıi]cak\s*Yüzey|Hot\s*Surface|Tehlike|Danger|Hazard|Risk|Dikkat|Caution|Warning|Alert)", "weight": 4},
                "kkd_zorunluluk": {"pattern": r"(?:KKD|PPE|Gözlük|Goggle|Eldiven|Glove|Kulaklık|Earplug|Baret|Helmet|Koruyucu\s*Donan[ıi]m|Protective\s*Equipment|Maske|Mask|İş\s*Güvenlik|Work\s*Safety|Koruyucu|Protective)", "weight": 4},
                "loto_kilitleme": {"pattern": r"(?:LOTO|Lock\s*Out|Tag\s*Out|Kilitle|Lock|Kapama|Shutdown|Enerji\s*Kes|Energy\s*Cut|İzolasyon|Isolation|Güvenli\s*Durdur|Safe\s*Stop|Devreden\s*Çıkar|Disconnect|Anahtar|Switch)", "weight": 4},
                "artik_riskler": {"pattern": r"(?:Art[ıi]k\s*Risk|Residual\s*Risk|Kalan\s*Tehlike|Remaining\s*Hazard|Tehlike|Hazard|Risk\s*Değerlendirme|Risk\s*Assessment|Güvenlik\s*Riski|Safety\s*Risk)", "weight": 3}
            },
            "Bakım Türleri ve Operasyon Çeşitleri": {
                "periyodik_bakim": {"pattern": r"(?:Periyodik|Periodic|Günlük|Daily|Haftal[ıi]k|Weekly|Ayl[ıi]k|Monthly|Y[ıi]ll[ıi]k|Yearly|Yağlama|Lubrication|Filtre|Filter|Rutin|Routine|Düzenli|Regular|Planlı|Planned|Programlı|Scheduled)", "weight": 6},
                "duzeltici_bakim": {"pattern": r"(?:Düzeltici|Corrective|Ar[ıi]za|Failure|Müdahale|Intervention|Onar[ıi]m|Repair|Tamir|Fix|Çözüm|Solution|Giderme|Troubleshoot|Acil|Emergency|Reaktif|Reactive)", "weight": 5},
                "ongorul_bakim": {"pattern": r"(?:Öngörülü|Predictive|Sensör|Sensor|Titreşim|Vibration|S[ıi]cakl[ıi]k|Temperature|Analiz|Analysis|İzleme|Monitoring|Tahmin|Prediction|Kondisyon|Condition|Durum\s*İzleme|Condition\s*Monitoring)", "weight": 4},
                "operasyon_cesitleri": {"pattern": r"(?:Ayarlama|Adjustment|Temizlik|Cleaning|Dezenfeksiyon|Disinfection|Sökme|Disassembly|İzolasyon|Isolation|Reset|Kalibrasyon|Calibration|Test|Kontrol|Check|Muayene|Inspection|Değişim|Replacement)", "weight": 5}
            },
            "Adım Adım Bakım İşlemleri": {
                "islem_sirasi": {"pattern": r"(?:Ad[ıi]m|Step|Sıra|Order|İşlem|Process|Prosedür|Procedure|Resim|Picture|Görsel|Image|Sıral[ıi]|Sequential|Aşama|Phase|Basamak|Stage|Numara|Number)", "weight": 3},
                "gerekli_aletler": {"pattern": r"(?:Alet|Tool|Cihaz|Device|Yedek\s*Parça|Spare\s*Part|Malzeme|Material|Equipment|Teçhizat|Ekipman|Gereç|Apparatus|Takım|Kit|Set|Donan[ıi]m|Hardware)", "weight": 3},
                "kontrol_listesi": {"pattern": r"(?:Kontrol\s*Liste|Check\s*List|Liste|List|Form|Checklist|Doğrulama|Verification|Onay|Approval|İmza|Signature|Teyit|Confirm|Sınar|Validate)", "weight": 2},
                "fonksiyon_testleri": {"pattern": r"(?:Test|Fonksiyon|Function|Doğrulama|Verification|Kontrol|Check|Muayene|Inspection|Deneme|Trial|Çal[ıi]şma|Operation|Performans|Performance|Kalite|Quality)", "weight": 2}
            },
            "Teknik Veriler": {
                "yaglama_bilgileri": {"pattern": r"(?:Yağlama|Lubrication|Yağ\s*Tür|Oil\s*Type|Yağ\s*Miktar|Oil\s*Amount|Gres|Grease|Lubricant|Yağl[ıi]|Oiled|Viscosity|Viskozite|Motor\s*Yağ|Engine\s*Oil)", "weight": 3},
                "tork_ayar_degerleri": {"pattern": r"(?:Tork|Torque|Ayar|Setting|Ölçü|Measure|Tolerans|Tolerance|Boşluk|Clearance|Değer|Value|Parametre|Parameter|Spec|Specification|Limit|S[ıi]n[ıi]r)", "weight": 3},
                "basinc_degerleri": {"pattern": r"(?:Bas[ıi]nç|Pressure|Elektrik|Electric|Pnömatik|Pneumatic|Hidrolik|Hydraulic|Bar|PSI|Volt|Ampere|Watt|kPa|MPa|V|A|W|Power|Güç)", "weight": 2},
                "sarf_malzeme_omru": {"pattern": r"(?:Ömür|Life|Sarf|Consumable|Filtre|Filter|Kay[ıi]ş|Belt|Conta|Gasket|Değişim|Change|Replace|Renewal|Yenileme|Service\s*Life|Working\s*Life|Kullan[ıi]m\s*Ömrü)", "weight": 2}
            },
            "Yedek Parça ve Sarf Malzemeleri": {
                "orijinal_parca_listesi": {"pattern": r"(?:Orijinal|Original|Yedek\s*Parça|Spare\s*Part|Parça\s*No|Part\s*Number|Liste|List|Katalog|Catalog|Code|Kod|OEM|Genuine|Gerçek|Authentic)", "weight": 6},
                "kritik_stok": {"pattern": r"(?:Kritik|Critical|Stok|Stock|Bulundur|Keep|Tavsiye|Recommend|Rezerv|Reserve|Minimum|Emergency|Acil|Zorunlu|Must|Required|Gerekli)", "weight": 5},
                "yanlis_parca_riski": {"pattern": r"(?:Yanl[ıi]ş|Wrong|Risk|Tehlike|Hazard|Uyar[ıi]|Warning|Dikkat|Attention|Caution|Uyumlu|Compatible|Uyumsuz|Incompatible|Doğru|Correct)", "weight": 4}
            },
            "Kayıt ve İzlenebilirlik": {
                "bakim_formu": {"pattern": r"(?:Bak[ıi]m\s*Form|Maintenance\s*Form|Kay[ıi]t|Record|Tarih|Date|Sorumlu|Responsible|Log|Rapor|Report|Dosya|File|Belge|Document|Archive|Arşiv)", "weight": 4},
                "ariza_kayitlari": {"pattern": r"(?:Ar[ıi]za\s*Kay[ıi]t|Failure\s*Record|Takip|Track|Değiştir|Replace|Geçmiş|History|Trend|İstatistik|Statistics|Analiz|Analysis|Log)", "weight": 3},
                "yasal_izlenebilirlik": {"pattern": r"(?:Yasal|Legal|İzlenebilir|Traceable|Gereklilik|Requirement|Uygunluk|Compliance|Standart|Standard|Sertifika|Certificate|Audit|Denetim|Regulation|Regulasyon)", "weight": 3}
            },
            "Çevre ve Atık Yönetimi": {
                "atik_bertaraf": {"pattern": r"(?:At[ıi]k|Waste|Bertaraf|Disposal|Yağ|Oil|Filtre|Filter|Akü|Battery|İmha|Destruction|Geri\s*Dönüşüm|Recycling|Çevre|Environment)", "weight": 2},
                "cevre_koruma": {"pattern": r"(?:Çevre|Environment|Zarar|Damage|İmha|Destruction|Geri\s*Dönüşüm|Recycling|Kirletici|Pollutant|Temiz|Clean|Yeşil|Green|Sürdürülebilir|Sustainable)", "weight": 2},
                "talimat_yontem": {"pattern": r"(?:Talimat|Instruction|Yöntem|Method|Prosedür|Procedure|Guide|Kılavuz|Guideline|Protocol|Protokol|Manual|El\s*Kitab[ıi])", "weight": 1}
            },
            "Ekler": {
                "resimli_sema": {"pattern": r"(?:Resim|Picture|Şema|Scheme|Diyagram|Diagram|Çizim|Drawing|Plan|Grafik|Graphic|Chart|Tablo|Table|Figure|Şekil|Image|Görsel)", "weight": 1},
                "ariza_teshis": {"pattern": r"(?:Ar[ıi]za\s*Teşhis|Fault\s*Diagnosis|Sorun|Problem|Neden|Cause|Çözüm|Solution|Troubleshoot|Debug|Hata|Error|Issue|Mesele)", "weight": 1},
                "iletisim_bilgi": {"pattern": r"(?:İletişim|Contact|Üretici|Manufacturer|Servis|Service|Telefon|Phone|E-mail|Mail|Address|Adres|Support|Destek|Help|Yard[ıi]m|Hotline)", "weight": 3}
            }
        }
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni Türkçe'ye çevir - şimdilik devre dışı"""
        if source_lang != 'tr':
            logger.info(f"Tespit edilen dil: {source_lang.upper()} - Çeviri yapılmıyor, orijinal metin kullanılıyor")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosya formatına göre metin çıkarma"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.doc':
            logger.warning("DOC formatı için DOCX'e dönüştürme gerekiyor veya OCR kullanılacak")
            return self.extract_text_from_docx(file_path)  # Deneme, başarısız olursa OCR
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatı: {file_ext}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den sadece PyPDF2 ile metin çıkarma - OCR Kaldırıldı"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                total_pages = len(pdf_reader.pages)
                pages_with_images = 0
                
                logger.info(f"PDF analizi başlatılıyor - Toplam sayfa: {total_pages}")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    # Sadece PyPDF2 ile metin çıkar
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = page_text.replace('|', ' ')
                    page_text = page_text.strip()
                    
                    # Görsel tespit (sadece sayma için)
                    has_images = len(page.images) > 0 if hasattr(page, 'images') else False
                    if has_images:
                        pages_with_images += 1
                    
                    all_text += page_text + "\n"
                
                # Metni temizle
                all_text = all_text.replace('—', '-')
                all_text = all_text.replace('"', '"').replace('"', '"')
                all_text = all_text.replace('´', "'")
                all_text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]+', ' ', all_text)
                all_text = all_text.strip()
                
                logger.info(f"✅ PDF analizi tamamlandı:")
                logger.info(f"   📊 Toplam metin uzunluğu: {len(all_text):,} karakter")
                logger.info(f"   📄 Toplam sayfa: {total_pages}")
                
                return all_text
                
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            
            # Paragrafları oku
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Tablolardan metin çıkar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            logger.info(f"DOCX'den {len(text)} karakter metin çıkarıldı")
            return text
            
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasından metin çıkarma"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text:
                # UTF-8 başarısız olursa diğer encoding'leri dene
                encodings = ['cp1254', 'iso-8859-9', 'latin1']
                for encoding in encodings:
                    try:
                        with open(txt_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        if text:
                            break
                    except:
                        continue
            
            logger.info(f"TXT'den {len(text)} karakter metin çıkarıldı")
            return text.strip()
            
        except Exception as e:
            logger.error(f"TXT metin çıkarma hatası: {e}")
            return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, MaintenanceAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                # DÜZELTME: Weight'e göre scoring algoritması
                if weight <= 2:
                    # Düşük weight'ler için basit scoring (minimum 1 puan garantisi)
                    score = min(weight, len(matches))
                else:
                    # Yüksek weight'ler için gelişmiş scoring
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
                    
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = MaintenanceAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={
                    "pattern_used": pattern,
                    "matches_count": len(matches) if matches else 0
                }
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, MaintenanceAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """Spesifik değerleri çıkar - Genel doküman uyumlu pattern'ler"""
        values = {
            "makine_adi": "Bulunamadı",
            "makine_modeli": "Bulunamadı",
            "seri_numarasi": "Bulunamadı",
            "bakim_turu": "Bulunamadı",
            "yetkili_personel": "Bulunamadı"
        }
        
        # ========================= MAKİNE ADI =========================
        
        # 1. Başlık formatlarından makine adı çıkarma
        machine_title_patterns = [
            r"([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,40})\s+(?:MAKİNESİ|MACHINE)\s+(?:BAKIM|KULLANIM|SERVİS|MAINTENANCE|SERVICE|OPERATION)",
            r"([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,40})\s+(?:CİHAZI|DEVICE)\s+(?:BAKIM|KULLANIM|SERVİS|MAINTENANCE|SERVICE)",
            r"([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,40})\s+(?:SİSTEMİ|SYSTEM)\s+(?:BAKIM|KULLANIM|SERVİS|MAINTENANCE|SERVICE)",
            r"([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,40})\s+(?:EKİPMANI|EQUIPMENT)\s+(?:BAKIM|KULLANIM|SERVİS|MAINTENANCE|SERVICE)",
            r"([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,40})\s+(?:BAKIM|MAINTENANCE)\s+(?:TALİMAT|KILAVUZ|PROSEDÜR|INSTRUCTION|MANUAL|PROCEDURE)",
            r"([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,40})\s+(?:SERVİS|SERVICE)\s+(?:TALİMAT|KILAVUZ|MANUAL|GUIDE)",
            r"([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,40})\s+(?:KULLANIM|OPERATION|USER)\s+(?:KILAVUZ|MANUAL|GUIDE)"
        ]
        
        # 2. Standart field formatlarından makine adı
        machine_field_patterns = [
            r"(?i)(?:ÜRÜN\s*ADI|ÜRÜN\s*TANIM|PRODUCT\s*NAME|PRODUCT\s*DESCRIPTION)\s*[:=]\s*([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,50})(?=\s*$|\s+(?:PROJE|MODEL|TİP|PROJECT|TYPE|\n))",
            r"(?i)(?:MAKİNE\s*ADI|MACHINE\s*NAME|MACHINE\s*TITLE)\s*[:=]\s*([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,50})(?=\s*$|\s+(?:PROJE|MODEL|PROJECT|\n))",
            r"(?i)(?:EKİPMAN\s*ADI|EQUIPMENT\s*NAME|EQUIPMENT\s*TITLE)\s*[:=]\s*([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,50})(?=\s*$|\s+(?:PROJE|MODEL|PROJECT|\n))",
            r"(?i)(?:CİHAZ\s*ADI|DEVICE\s*NAME|DEVICE\s*TITLE)\s*[:=]\s*([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,50})(?=\s*$|\s+(?:PROJE|MODEL|PROJECT|\n))",
            r"(?i)(?:SİSTEM\s*ADI|SYSTEM\s*NAME|SYSTEM\s*TITLE)\s*[:=]\s*([A-ZÇĞİÖŞÜ][A-Z0-9ÇĞİÖŞÜ\s-]{3,50})(?=\s*$|\s+(?:PROJE|MODEL|PROJECT|\n))"
        ]
        
        # Önce başlık pattern'lerini dene
        for pattern in machine_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 3 <= len(result) <= 50 and not result.lower().startswith(("bu", "şu", "o ")):
                    values["makine_adi"] = result
                    break
        
        # Başlık pattern'lerinde bulunamadıysa field pattern'leri dene
        if values["makine_adi"] == "Bulunamadı":
            for pattern in machine_field_patterns:
                match = re.search(pattern, text)
                if match:
                    result = match.group(1).strip()
                    # Post-processing temizlik
                    cleanup_words = ["PROJESİ", "PROJECT", "SİSTEMİ", "SYSTEM", "EKİPMANI", "EQUIPMENT", "TALİMATI", "INSTRUCTION"]
                    for cleanup_word in cleanup_words:
                        result = re.sub(rf"\b{cleanup_word}\b", "", result, flags=re.IGNORECASE).strip()
                    result = re.sub(r'\s+', ' ', result).strip()
                    
                    if 3 <= len(result) <= 50 and not result.lower().startswith(("bu", "şu", "o ")):
                        values["makine_adi"] = result
                        break
        
        # ========================= MODEL NUMARASI =========================
        
        # 1. Başlık formatlarından model çıkarma
        model_title_patterns = [
            r"(?:MODEL|TİP|TYPE)\s+([A-Z0-9-]{2,20})\s+(?:MAKİNESİ|CİHAZI|SİSTEMİ|BAKIM|SERVİS)",
            r"([A-Z0-9-]{2,20})\s+(?:MODEL|TİP|TYPE)\s+(?:BAKIM|SERVİS|KULLANIM|MAINTENANCE)",
            r"([A-Z0-9-]{2,20})\s+(?:SERİSİ|SERIES)\s+(?:BAKIM|SERVİS|KULLANIM)"
        ]
        
        # 2. Standart field formatlarından model
        model_field_patterns = [
            r"(?i)(?:MODEL\s*NO|MODEL\s*NUMARASI|MODEL\s*NUMBER)\s*[:=]\s*([A-Z0-9-]{2,20})(?=\s|$|\n)",
            r"(?i)(?:MODEL\s*KODU|MODEL\s*CODE)\s*[:=]\s*([A-Z0-9-]{2,20})(?=\s|$|\n)",
            r"(?i)(?:MODEL)\s*[:=]\s*([A-Z0-9-]{2,20})(?=\s|$|\n)",
            r"(?i)(?:TİP|TYPE)\s*[:=]\s*([A-Z0-9-]{2,20})(?=\s|$|\n)",
            r"(?i)(?:TİP\s*NO|TYPE\s*NO|TİP\s*NUMARASI|TYPE\s*NUMBER)\s*[:=]\s*([A-Z0-9-]{2,20})(?=\s|$|\n)"
        ]
        
        # Önce başlık pattern'lerini dene
        for pattern in model_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 2 <= len(result) <= 25:
                    values["makine_modeli"] = result
                    break
        
        # Başlık pattern'lerinde bulunamadıysa field pattern'leri dene
        if values["makine_modeli"] == "Bulunamadı":
            for pattern in model_field_patterns:
                match = re.search(pattern, text)
                if match:
                    result = match.group(1).strip()
                    filter_words = ["konveyör", "tipi", "çeşit", "türü", "xxx", "tbd", "n/a"]
                    if not any(word in result.lower() for word in filter_words):
                        if 2 <= len(result) <= 25:
                            values["makine_modeli"] = result
                            break
        
        # ========================= SERİ NUMARASI =========================
        
        # 1. Başlık formatlarından seri numarası
        serial_title_patterns = [
            r"(?:SERİ|SERIAL)\s+(?:NO|NUMBER|NUMARASI)[\s:=]*([A-Z0-9-]{3,25})",
            r"(?:S/N|SN)[\s:=]*([A-Z0-9-]{3,25})"
        ]
        
        # 2. Standart field formatlarından seri numarası
        serial_field_patterns = [
            r"(?i)(?:SERİ\s*NO|SERİ\s*NUMARASI|SERIAL\s*NUMBER|S/N)\s*[:=]\s*([A-Z0-9-]{3,25})(?=\s+(?:Üretim|Tarih|Rev|Revizyon|\n)|$)",
            r"(?i)(?:SERİ\s*KODU|SERIAL\s*CODE)\s*[:=]\s*([A-Z0-9-]{3,25})(?=\s+(?:Üretim|Tarih|Rev|\n)|$)",
            r"(?i)(?:SERIAL)\s*[:=]\s*([A-Z0-9-]{3,25})(?=\s|$)"
        ]
        
        # Placeholder pattern'leri
        placeholder_patterns = [
            r"^[X]{2,}$", r"^[X-]{2,}$", r".*[X]{3,}.*", r"^[-]{2,}$",
            r"(?i)^(tbd|n/a|na|null|none|boş|yok)$", r"(?i).*rev[x0-9].*",
            r"^[0]{3,}$", r"(?i)^(sample|örnek|example).*"
        ]
        
        # Başlık pattern'lerini dene
        for pattern in serial_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                is_placeholder = any(re.match(p, result) for p in placeholder_patterns)
                if not is_placeholder and 3 <= len(result) <= 30:
                    values["seri_numarasi"] = result
                    break
        
        # Field pattern'leri dene
        if values["seri_numarasi"] == "Bulunamadı":
            for pattern in serial_field_patterns:
                match = re.search(pattern, text)
                if match:
                    result = match.group(1).strip()
                    is_placeholder = any(re.match(p, result) for p in placeholder_patterns)
                    if not is_placeholder and 3 <= len(result) <= 30:
                        values["seri_numarasi"] = result
                        break
        
        # ========================= BAKIM TÜRÜ =========================
        
        # 1. Başlık formatlarından bakım türü
        maintenance_title_patterns = [
            r"(PERİYODİK\s*BAKIM|PERIODIC\s*MAINTENANCE)\s+(?:TALİMAT|PROSEDÜR|MANUAL)",
            r"(DÜZELTİCİ\s*BAKIM|CORRECTIVE\s*MAINTENANCE)\s+(?:TALİMAT|PROSEDÜR|MANUAL)",
            r"(ÖNGÖRÜLÜ\s*BAKIM|PREDICTIVE\s*MAINTENANCE)\s+(?:TALİMAT|PROSEDÜR|MANUAL)",
            r"(GÜNLÜK\s*BAKIM|DAILY\s*MAINTENANCE)\s+(?:TALİMAT|PROSEDÜR|MANUAL)",
            r"(HAFTALIK\s*BAKIM|WEEKLY\s*MAINTENANCE)\s+(?:TALİMAT|PROSEDÜR|MANUAL)",
            r"(AYLIK\s*BAKIM|MONTHLY\s*MAINTENANCE)\s+(?:TALİMAT|PROSEDÜR|MANUAL)",
            r"(YILLIK\s*BAKIM|YEARLY\s*MAINTENANCE|ANNUAL\s*MAINTENANCE)\s+(?:TALİMAT|PROSEDÜR|MANUAL)"
        ]
        
        # 2. Standart word boundary pattern'leri
        maintenance_word_patterns = [
            r"\b(Periyodik\s*Bakım|Periodic\s*Maintenance)\b",
            r"\b(Düzeltici\s*Bakım|Corrective\s*Maintenance)\b", 
            r"\b(Öngörülü\s*Bakım|Predictive\s*Maintenance)\b",
            r"\b(Önleyici\s*Bakım|Preventive\s*Maintenance)\b",
            r"\b(Günlük\s*Bakım|Daily\s*Maintenance)\b",
            r"\b(Haftalık\s*Bakım|Weekly\s*Maintenance)\b",
            r"\b(Aylık\s*Bakım|Monthly\s*Maintenance)\b",
            r"\b(Yıllık\s*Bakım|Yearly\s*Maintenance|Annual\s*Maintenance)\b",
            r"\b(Acil\s*Bakım|Emergency\s*Maintenance)\b",
            r"\b(Planlı\s*Bakım|Planned\s*Maintenance)\b"
        ]
        
        # Önce başlık pattern'lerini dene
        for pattern in maintenance_title_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result = match.group(1).strip()
                if len(result) <= 35:
                    values["bakim_turu"] = result
                    break
        
        # Başlık pattern'lerinde bulunamadıysa word pattern'leri dene
        if values["bakim_turu"] == "Bulunamadı":
            for pattern in maintenance_word_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    result = match.group().strip()
                    if len(result) <= 35:
                        values["bakim_turu"] = result
                        break
        
        # ========================= YETKİLİ PERSONEL =========================
        
        # 1. Başlık formatlarından personel
        personnel_title_patterns = [
            r"(?:SORUMLU|RESPONSIBLE)[\s:=]*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$|\n)",
            r"(?:TEKNİSYEN|TECHNICIAN)[\s:=]*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$|\n)",
            r"(?:OPERATÖR|OPERATOR)[\s:=]*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$|\n)"
        ]
        
        # 2. Standart field formatlarından personel
        personnel_field_patterns = [
            r"(?i)(?:YETKİLİ\s*PERSONEL|AUTHORIZED\s*PERSONNEL|AUTHORIZED\s*STAFF)\s*[:=]\s*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$)",
            r"(?i)(?:TEKNİSYEN|TECHNICIAN|TECHNICAL\s*STAFF)\s*[:=]\s*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$)",
            r"(?i)(?:SORUMLU\s*PERSONEL|RESPONSIBLE\s*PERSONNEL|RESPONSIBLE\s*STAFF|RESPONSIBLE\s*PERSON)\s*[:=]\s*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$)",
            r"(?i)(?:BAKIM\s*SORUMLUSU|MAINTENANCE\s*RESPONSIBLE|MAINTENANCE\s*SUPERVISOR|MAINTENANCE\s*MANAGER)\s*[:=]\s*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$)",
            r"(?i)(?:OPERATÖR|OPERATOR|MACHINE\s*OPERATOR)\s*[:=]\s*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$)",
            r"(?i)(?:UZMAN|EXPERT|SPESİYALİST|SPECIALIST|TECHNICAL\s*EXPERT)\s*[:=]\s*([A-ZÇĞİÖŞÜ][a-züçğıöş\s]{2,30})(?=\s|$)"
        ]
        
        # Önce başlık pattern'lerini dene
        for pattern in personnel_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                filter_words = ["saklanması", "gereken", "yapılması", "bulunması", "olması", "edilmesi", "sağlanması"]
                if not any(word in result.lower() for word in filter_words):
                    if len(result.split()) >= 2 or (len(result.split()) == 1 and len(result) >= 4):
                        if 3 <= len(result) <= 35:
                            values["yetkili_personel"] = result
                            break
        
        # Field pattern'leri dene
        if values["yetkili_personel"] == "Bulunamadı":
            for pattern in personnel_field_patterns:
                match = re.search(pattern, text)
                if match:
                    result = match.group(1).strip()
                    filter_words = ["saklanması", "gereken", "yapılması", "bulunması", "olması", "edilmesi", "sağlanması"]
                    if not any(word in result.lower() for word in filter_words):
                        if len(result.split()) >= 2 or (len(result.split()) == 1 and len(result) >= 4):
                            if 3 <= len(result) <= 35:
                                values["yetkili_personel"] = result
                                break
        
        return values

    def validate_maintenance_document(self, text: str) -> bool:
        """Dokümanın bakım talimatları olup olmadığını kontrol et"""
        
        # Bakım talimatlarında olması gereken genişletilmiş kelimeler
        maintenance_keywords = [
            # Temel bakım terimleri
            "bakım", "maintenance", "talimat", "instruction", "kılavuz", "manual", "guide",
            
            # Ekipman terimleri  
            "makine", "machine", "ekipman", "equipment", "cihaz", "device", "sistem", "system",
            
            # Güvenlik terimleri
            "güvenlik", "safety", "tehlike", "hazard", "risk", "dikkat", "caution", "warning",
            
            # Operasyon terimleri
            "prosedür", "procedure", "işlem", "process", "operasyon", "operation", "çalışma", "work",
            
            # Servis terimleri
            "onarım", "repair", "servis", "service", "tamir", "fix", "düzeltici", "corrective",
            
            # Kontrol terimleri
            "kontrol", "check", "muayene", "inspection", "test", "doğrulama", "verification",
            
            # Periyodik terimler
            "periyodik", "periodic", "günlük", "daily", "haftalık", "weekly", "aylık", "monthly",
            
            # Parça terimleri
            "parça", "part", "yedek", "spare", "malzeme", "material", "sarf", "consumable",
            
            # Teknik terimler
            "yağlama", "lubrication", "filtre", "filter", "ayarlama", "adjustment", "kalibrasyon", "calibration",
            
            # Alet terimleri
            "alet", "tool", "teçhizat", "apparatus", "donanım", "hardware"
        ]
        
        # En az 2 anahtar kelime bulunmalı (eşik düşürüldü)
        found_keywords = 0
        found_words = []
        
        for keyword in maintenance_keywords:
            if re.search(rf"\b{keyword}\b", text, re.IGNORECASE):
                found_keywords += 1
                found_words.append(keyword)
                
        logger.info(f"Doküman validasyonu: {found_keywords} anahtar kelime bulundu: {found_words[:10]}")  # İlk 10 kelimeyi göster
        
        return found_keywords >= 2  # Minimum 2 kelime (3'ten düşürüldü)

    def generate_improvement_actions(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Dinamik iyileştirme önerileri oluştur"""
        actions = []
        
        # Kategorileri puana göre sırala (düşükten yükseğe)
        sorted_categories = sorted(
            scores["category_scores"].items(), 
            key=lambda x: x[1]["percentage"]
        )
        
        category_actions = {
            "Genel Bilgiler": [
                "Makine tanımını (ad, model, seri numarası, üretici) netleştiriniz",
                "Belge kapsamını ve hangi makine(ler) için geçerli olduğunu belirtiniz",
                "Yetkili personel niteliklerini (eğitilmiş teknisyen vb.) tanımlayınız"
            ],
            "Güvenlik Önlemleri": [
                "Genel güvenlik uyarılarını (elektrik çarpması, hareketli parçalar) ekleyiniz",
                "KKD gerekliliklerini (gözlük, eldiven, kulaklık) detaylandırınız",
                "LOTO prosedürlerini ve makineyi güvenli duruma getirme adımlarını belirtiniz",
                "Artık riskleri ve bakım sırasındaki tehlikeleri açıklayınız"
            ],
            "Bakım Türleri ve Operasyon Çeşitleri": [
                "Periyodik bakım türlerini (günlük, haftalık, aylık, yıllık) tanımlayınız",
                "Düzeltici bakım prosedürlerini ve arıza sonrası müdahaleleri açıklayınız",
                "Öngörülü bakım yöntemlerini (sensör verileri, titreşim analizi) ekleyiniz",
                "Operasyon çeşitlerini (ayarlama, temizlik, yağlama, parça değişimi) listeleyiniz"
            ],
            "Adım Adım Bakım İşlemleri": [
                "İşlem sırasını açık ve resimli olarak hazırlayınız",
                "Gerekli alet ve yedek parçaları belirtiniz",
                "Kontrol listelerini (checklist) oluşturunuz",
                "Fonksiyon testlerini ve işlem sonrası kontrolleri açıklayınız"
            ],
            "Teknik Veriler": [
                "Yağlama noktaları, yağ türleri ve miktarlarını belirtiniz",
                "Tork değerleri, ayar ölçüleri ve boşluk toleranslarını ekleyiniz",
                "Elektrik/Pnömatik/Hidrolik bağlantı basınçlarını tanımlayınız",
                "Sarf malzeme ömürlerini (filtre, kayış, conta) listeleyiniz"
            ],
            "Yedek Parça ve Sarf Malzemeleri": [
                "Orijinal yedek parça listesini parça numaralarıyla hazırlayınız",
                "Kritik yedeklerin stokta bulundurulma tavsiyelerini ekleyiniz",
                "Yanlış parça kullanımının risklerini açıklayınız"
            ],
            "Kayıt ve İzlenebilirlik": [
                "Bakım formu/tablosunu (tarih, sorumlu, yapılan işlemler) oluşturunuz",
                "Arıza kayıtları ve değiştirilen parça takip sistemini kurunuz",
                "Yasal gerekliliklere uygun izlenebilirlik sağlayınız"
            ],
            "Çevre ve Atık Yönetimi": [
                "Kullanılmış yağ, filtre, akü atıklarının bertaraf yöntemlerini belirtiniz",
                "Çevreye zarar vermeyecek imha ve geri dönüşüm talimatlarını ekleyiniz",
                "Atık yönetimi prosedürlerini ve çevresel uygunluk kriterlerini tanımlayınız"
            ],
            "Ekler": [
                "Resimli şema ve diyagramları ekleyiniz",
                "Arıza teşhis tablosunu (sorun-neden-çözüm) hazırlayınız",
                "İletişim bilgilerini (üretici, servis sağlayıcı) güncel tutunuz"
            ]
        }
        
        # En düşük puanlı 5 kategori için öneriler ekle
        for category, score_data in sorted_categories[:5]:
            if score_data["percentage"] < 70:  # %70'in altındaki kategoriler için
                actions.extend(category_actions.get(category, []))
        
        # Genel öneriler
        if scores["percentage"] < 50:
            actions.insert(0, "ÖNCE: Doküman yapısını ve formatını yeniden gözden geçiriniz")
        
        return actions
    
    def generate_recommendations(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Öneriler oluştur"""
        recommendations = []
        
        total_percentage = scores["percentage"]
        
        if total_percentage >= 70:
            recommendations.append(f"✅ Bakım Talimatları GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ Bakım Talimatları GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 40:
                recommendations.append(f"🔴 {category} bölümü yetersiz (%{category_score:.1f})")
                missing_items = [name for name, result in results.items() if not result.found]
                if missing_items:
                    recommendations.append(f"   Eksik: {', '.join(missing_items[:3])}")
            elif category_score < 70:
                recommendations.append(f"🟡 {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"🟢 {category} bölümü yeterli (%{category_score:.1f})")
        
        return recommendations

    def analyze_maintenance_report(self, file_path: str) -> Dict[str, Any]:
        """Ana Bakım Talimatları analiz fonksiyonu"""
        logger.info("Bakım Talimatları analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        if detected_lang != 'tr':
            logger.info(f"{detected_lang.upper()} dilinden Türkçe'ye çeviriliyor...")
            text = self.translate_to_turkish(text, detected_lang)
        
        # Doküman türü kontrolü - YANLIŞ DOKÜMAN TESPİTİ
        if not self.validate_maintenance_document(text):
            return {
                "error": "YANLIŞ DOKÜMAN: Bu dosya bakım talimatları dökümanı değil!",
                "document_type": "UNKNOWN",
                "suggestion": "Lütfen geçerli bir bakım talimatları dökümanı yükleyiniz."
            }
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        recommendations = self.generate_recommendations(analysis_results, scores)
        improvement_actions = self.generate_improvement_actions(analysis_results, scores)
        
        final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "iyilestirme_eylemleri": improvement_actions,
            "ozet": {
                "toplam_puan": scores["total_score"],
                "yuzde": scores["percentage"],
                "durum": final_status,
                "rapor_tipi": "BAKIM_TALIMATLARI"
            }
        }
        
        return report

def main():
    """Ana fonksiyon"""
    analyzer = MaintenanceReportAnalyzer()

    # Test için örnek dosya yolları
    file_path = r"C:\Users\nuvo_teknik_2\Desktop\PILZ DOCUMENTS\3.3 Bakım Talimatları\Bakım Talimatı.pdf"
    # file_path = r"C:\Users\nuvo_teknik_2\Desktop\MAINTENANCE_DOCS\bakim_talimati.docx"

    if not os.path.exists(file_path):
        print(f"❌ Dosya bulunamadı: {file_path}")
        return
    
    print("🔧 Bakım Talimatları Analizi Başlatılıyor...")
    print("=" * 60)
    
    report = analyzer.analyze_maintenance_report(file_path)
    
    if "error" in report:
        print(f"❌ Hata: {report['error']}")
        return
    
    print("\n📊 ANALİZ SONUÇLARI")
    print("=" * 60)
    
    print(f"📅 Analiz Tarihi: {report['analiz_tarihi']}")
    print(f"📄 Dosya Tipi: {report['dosya_bilgisi']['file_type'].upper()}")
    print(f"🔍 Tespit Edilen Dil: {report['dosya_bilgisi']['detected_language'].upper()}")
    print(f"📋 Toplam Puan: {report['ozet']['toplam_puan']}/100")
    print(f"📈 Yüzde: %{report['ozet']['yuzde']}")
    print(f"🎯 Durum: {report['ozet']['durum']}")
    print(f"📄 Rapor Tipi: {report['ozet']['rapor_tipi']}")
    
    print("\n📋 ÖNEMLİ ÇIKARILAN DEĞERLER")
    print("-" * 40)
    for key, value in report['cikarilan_degerler'].items():
        display_name = {
            "makine_adi": "Makine Adı",
            "makine_modeli": "Makine Modeli",
            "seri_numarasi": "Seri Numarası",
            "bakim_turu": "Bakım Türü",
            "yetkili_personel": "Yetkili Personel"
        }.get(key, key.replace('_', ' ').title())
        print(f"{display_name}: {value}")
    
    print("\n📊 KATEGORİ PUANLARI")
    print("-" * 40)
    for category, score_data in report['puanlama']['category_scores'].items():
        print(f"{category}: {score_data['normalized']}/{score_data['max_weight']} (%{score_data['percentage']:.1f})")
    
    print("\n💡 ÖNERİLER VE DEĞERLENDİRME")
    print("-" * 40)
    for recommendation in report['oneriler']:
        print(recommendation)
    
    print("\n📋 GENEL DEĞERLENDİRME")
    print("=" * 60)
    
    if report['ozet']['yuzde'] >= 70:
        print("✅ SONUÇ: GEÇERLİ")
        print(f"🌟 Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Bakım talimatları genel olarak yeterli kriterleri sağlamaktadır.")
    else:
        print("❌ SONUÇ: GEÇERSİZ")
        print(f"⚠️ Toplam Başarı: %{report['ozet']['yuzde']:.1f}")
        print("📝 Değerlendirme: Bakım talimatları minimum gereklilikleri sağlamamaktadır.")
        
        print("\n⚠️ EKSİK GEREKLİLİKLER:")
        for category, results in report['kategori_analizleri'].items():
            missing_items = []
            for criterion, result in results.items():
                if not result.found:
                    missing_items.append(criterion)
            
            if missing_items:
                print(f"\n🔍 {category}:")
                for item in missing_items:
                    readable_name = {
                        "makine_tanimi": "Makine Tanımı",
                        "belge_kapsami": "Belge Kapsamı", 
                        "yetkili_personel": "Yetkili Personel",
                        "genel_guvenlik_uyari": "Genel Güvenlik Uyarıları",
                        "kkd_zorunluluk": "KKD Zorunlulukları",
                        "loto_kilitleme": "LOTO/Kilitleme Prosedürleri",
                        "artik_riskler": "Artık Riskler",
                        "periyodik_bakim": "Periyodik Bakım Türleri",
                        "duzeltici_bakim": "Düzeltici Bakım",
                        "ongorul_bakim": "Öngörülü Bakım",
                        "operasyon_cesitleri": "Operasyon Çeşitleri",
                        "islem_sirasi": "İşlem Sırası",
                        "gerekli_aletler": "Gerekli Aletler",
                        "kontrol_listesi": "Kontrol Listesi",
                        "fonksiyon_testleri": "Fonksiyon Testleri",
                        "yaglama_bilgileri": "Yağlama Bilgileri",
                        "tork_ayar_degerleri": "Tork/Ayar Değerleri",
                        "basinc_degerleri": "Basınç Değerleri",
                        "sarf_malzeme_omru": "Sarf Malzeme Ömrü",
                        "orijinal_parca_listesi": "Orijinal Parça Listesi",
                        "kritik_stok": "Kritik Stok Tavsiyeleri",
                        "yanlis_parca_riski": "Yanlış Parça Kullanım Riski",
                        "bakim_formu": "Bakım Formu/Tablosu",
                        "ariza_kayitlari": "Arıza Kayıtları",
                        "yasal_izlenebilirlik": "Yasal İzlenebilirlik",
                        "atik_bertaraf": "Atık Bertaraf Yöntemleri",
                        "cevre_koruma": "Çevre Koruma",
                        "talimat_yontem": "Bertaraf Talimat/Yöntemleri",
                        "resimli_sema": "Resimli Şema/Diyagramlar",
                        "ariza_teshis": "Arıza Teşhis Tablosu",
                        "iletisim_bilgi": "İletişim Bilgileri"
                    }.get(item, item.replace('_', ' ').title())
                    print(f"   ❌ {readable_name}")
        
        print("\n📌 YAPILMASI GEREKENLER:")
        if "iyilestirme_eylemleri" in report:
            for i, action in enumerate(report['iyilestirme_eylemleri'], 1):
                print(f"{i}. {action}")
        else:
            # Fallback - eski statik liste
            print("1. Güvenlik önlemleri ve LOTO prosedürlerini detaylandırın")
            print("2. Adım adım bakım talimatlarını numaralandırarak açıklayın") 
            print("3. Bakım periyodu ve zamanlaması netleştirin")
            print("4. Gerekli ekipman ve malzeme listesini eksiksiz hazırlayın")
            print("5. Kontrol ve onay prosedürlerini güçlendirin")
            print("6. Teknik görseller, fotoğraflar ve şemalar ekleyin")
            print("7. Referans standartları ve belgelerini belirtin")
            print("8. Sonuç değerlendirmesi ve iyileştirme önerilerini ekleyin")

if __name__ == "__main__":
    main()