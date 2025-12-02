# ============================================
# AYDINLATMA Ã–LÃ‡ÃœM RAPORU ANALÄ°Z SERVÄ°SÄ°
# Standalone Service - Azure App Service Ready
# Port: 8008
# Database-Driven Configuration
# ============================================

# ============================================
# IMPORTS
# ============================================
import os
import json
from datetime import datetime, timedelta
import re
from typing import Dict, List, Tuple, Any
import PyPDF2
from docx import Document
from dataclasses import dataclass
import logging
import cv2
import numpy as np
from PIL import Image
import pdf2image
import pytesseract

from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename

# ============================================
# DATABASE IMPORTS (YENÄ°)
# ============================================
from flask import current_app
from database import db, init_db
from db_loader import load_service_config

# ============================================
# CONFIG IMPORT (YENÄ°)
# ============================================
from config import Config

# ============================================
# LOGGING CONFIGURATION
# ============================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================
# LANGUAGE DETECTION (Optional)
# ============================================
try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False
    logger.warning("langdetect modÃ¼lÃ¼ bulunamadÄ± - dil tespiti devre dÄ±ÅŸÄ±")

# ============================================
# TESSERACT CHECK
# ============================================
def check_tesseract_installation():
    """Check if Tesseract is properly installed"""
    try:
        version = pytesseract.get_tesseract_version()
        logger.info(f"Tesseract OCR kurulu - SÃ¼rÃ¼m: {version}")
        return True, f"Tesseract v{version}"
    except Exception as e:
        logger.error(f"Tesseract OCR kurulu deÄŸil: {e}")
        return False, str(e)

tesseract_available, tesseract_info = check_tesseract_installation()

# ============================================
# ANALÄ°Z SINIFI - DATA CLASSES
# ============================================
@dataclass
class LightingAnalysisResult:
    """AydÄ±nlatma Ã–lÃ§Ã¼m Raporu analiz sonucu veri sÄ±nÄ±fÄ±"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

# ============================================
# ANALÄ°Z SINIFI - MAIN ANALYZER
# ============================================
class LightingReportAnalyzer:
    """AydÄ±nlatma Ã–lÃ§Ã¼m Raporu analiz sÄ±nÄ±fÄ±"""
    
    def __init__(self, app=None):
        logger.info("AydÄ±nlatma Ã–lÃ§Ã¼m Raporu analiz sistemi baÅŸlatÄ±lÄ±yor...")
        
        # Flask app context varsa DB'den yÃ¼kle, yoksa boÅŸ baÅŸlat
        if app:
            with app.app_context():
                try:
                    config = load_service_config('lighting_report')
                    
                    # DB'den yÃ¼klenen veriler
                    self.criteria_weights = config.get('criteria_weights', {})
                    self.criteria_details = config.get('criteria_details', {})
                    self.pattern_definitions = config.get('pattern_definitions', {})
                    self.validation_keywords = config.get('validation_keywords', {})
                    self.category_actions = config.get('category_actions', {})
                    
                    logger.info(f"âœ… VeritabanÄ±ndan yÃ¼klendi: {len(self.criteria_weights)} kategori")
                    
                except Exception as e:
                    logger.error(f"âš ï¸ VeritabanÄ±ndan yÃ¼kleme baÅŸarÄ±sÄ±z: {e}")
                    logger.warning("âš ï¸ Fallback: BoÅŸ config kullanÄ±lÄ±yor")
                    self.criteria_weights = {}
                    self.criteria_details = {}
                    self.pattern_definitions = {}
                    self.validation_keywords = {}
                    self.category_actions = {}
        else:
            # Flask app yoksa boÅŸ baÅŸlat (eski davranÄ±ÅŸ)
            logger.warning("âš ï¸ Flask app context yok, boÅŸ config kullanÄ±lÄ±yor")
            self.criteria_weights = {}
            self.criteria_details = {}
            self.pattern_definitions = {}
            self.validation_keywords = {}
            self.category_actions = {}
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti baÅŸarÄ±sÄ±z: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        """Metni TÃ¼rkÃ§e'ye Ã§evir - ÅŸimdilik devre dÄ±ÅŸÄ±"""
        if source_lang != 'tr':
            logger.info(f"Tespit edilen dil: {source_lang.upper()} - Ã‡eviri yapÄ±lmÄ±yor, orijinal metin kullanÄ±lÄ±yor")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosya formatÄ±na gÃ¶re metin Ã§Ä±karma"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.doc':
            logger.warning("DOC formatÄ± iÃ§in DOCX'e dÃ¶nÃ¼ÅŸtÃ¼rme gerekiyor veya OCR kullanÄ±lacak")
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatÄ±: {file_ext}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den metin Ã§Ä±karma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF analizi baÅŸlatÄ±lÄ±yor - Toplam sayfa: {total_pages}")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    page_text = page_text.replace('|', ' ')
                    page_text = page_text.strip()
                    all_text += page_text + "\n"
                
                all_text = all_text.replace('â€”', '-')
                all_text = all_text.replace('"', '"').replace('"', '"')
                all_text = all_text.replace('Â´', "'")
                all_text = re.sub(r'[^\x00-\x7F\u00C0-\u00FF\u0100-\u017F\u0180-\u024F]+', ' ', all_text)
                all_text = all_text.strip()
                
                logger.info(f"âœ… PDF analizi tamamlandÄ±: {len(all_text):,} karakter")
                return all_text
                
        except Exception as e:
            logger.error(f"PDF metin Ã§Ä±karma hatasÄ±: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin Ã§Ä±karma"""
        try:
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            text = re.sub(r'\s+', ' ', text).strip()
            logger.info(f"DOCX'den {len(text)} karakter metin Ã§Ä±karÄ±ldÄ±")
            return text
            
        except Exception as e:
            logger.error(f"DOCX metin Ã§Ä±karma hatasÄ±: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasÄ±ndan metin Ã§Ä±karma"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text:
                encodings = ['cp1254', 'iso-8859-9', 'latin1']
                for encoding in encodings:
                    try:
                        with open(txt_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        if text:
                            break
                    except:
                        continue
            
            logger.info(f"TXT'den {len(text)} karakter metin Ã§Ä±karÄ±ldÄ±")
            return text.strip()
            
        except Exception as e:
            logger.error(f"TXT metin Ã§Ä±karma hatasÄ±: {e}")
            return ""

    def analyze_criteria(self, text: str, category: str) -> Dict[str, LightingAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
                    
            else:
                content = "BulunamadÄ±"
                found = False
                score = 0
            
            results[criterion_name] = LightingAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={"pattern_used": pattern, "matches_count": len(matches) if matches else 0}
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, LightingAnalysisResult]]) -> Dict[str, Any]:
        """PuanlarÄ± hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """AydÄ±nlatma raporuna Ã¶zgÃ¼ deÄŸerleri Ã§Ä±kar"""
        values = {
            "rapor_numarasi": "BulunamadÄ±",
            "proje_adi": "BulunamadÄ±",
            "olcum_tarihi": "BulunamadÄ±", 
            "rapor_tarihi": "BulunamadÄ±",
            "olcum_cihazi": "BulunamadÄ±",
            "tesis_adi": "BulunamadÄ±",
            "genel_uygunluk": "BulunamadÄ±"
        }
        
        # DB'den pattern'leri al
        extract_values = self.pattern_definitions.get('extract_values', {})
        
        # RAPOR NUMARASI
        report_no_patterns = extract_values.get('rapor_numarasi', [])
        for pattern in report_no_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 3 <= len(result) <= 20:
                    values["rapor_numarasi"] = result
                    break

        # PROJE ADI
        project_patterns = extract_values.get('proje_adi', [])
        for pattern in project_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                result = re.sub(r'\s+', ' ', result)
                if 3 <= len(result) <= 100:
                    values["proje_adi"] = result
                    break
        
        # Ã–LÃ‡ÃœM TARÄ°HÄ°
        measurement_date_patterns = extract_values.get('olcum_tarihi', [])
        for pattern in measurement_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["olcum_tarihi"] = match.group(1)
                break
        
        # RAPOR TARÄ°HÄ°
        report_date_patterns = extract_values.get('rapor_tarihi', [])
        for pattern in report_date_patterns:
            match = re.search(pattern, text)
            if match:
                values["rapor_tarihi"] = match.group(1)
                break
    
        if values["rapor_tarihi"] == "BulunamadÄ±" and values["olcum_tarihi"] != "BulunamadÄ±":
            values["rapor_tarihi"] = "Rapor tarihi ayrÄ± belirtilmemiÅŸ"

        # Ã–LÃ‡ÃœM CÄ°HAZI
        device_patterns = extract_values.get('olcum_cihazi', [])
        for pattern in device_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 2 <= len(result) <= 35:
                    values["olcum_cihazi"] = result
                    break
        
        # TESÄ°S ADI
        facility_patterns = extract_values.get('tesis_adi', [])
        for pattern in facility_patterns:
            matches = re.findall(pattern, text)
            for result in matches:
                result = result.strip()
                if 3 <= len(result) <= 60: 
                    values["tesis_adi"] = result
                    break
            if values["tesis_adi"] != "BulunamadÄ±":
                break
        
        # GENEL UYGUNLUK
        compliance_patterns = extract_values.get('genel_uygunluk', [])
        for pattern in compliance_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip().upper()
                if result in ["UYGUN", "SUITABLE", "CONFORM", "GEÃ‡ERLÄ°", "VALID", "PASS"]:
                    values["genel_uygunluk"] = "UYGUN"
                elif result in ["UYGUNSUZ", "NOT SUITABLE", "NON-CONFORM", "GEÃ‡ERSÄ°Z", "INVALID", "FAIL"]:
                    values["genel_uygunluk"] = "UYGUNSUZ"
                break
        
        return values

    def check_date_validity(self, measurement_date: str, report_date: str) -> Tuple[bool, str]:
        """Ã–lÃ§Ã¼m ve rapor tarihlerini bugÃ¼nkÃ¼ tarih ile kontrol et (1 yÄ±l kuralÄ±)"""
        if measurement_date == "BulunamadÄ±" and report_date == "BulunamadÄ±":
            return False, "Ne Ã¶lÃ§Ã¼m ne de rapor tarihi bulunamadÄ±"
        
        dates_to_check = []
        
        if measurement_date != "BulunamadÄ±":
            dates_to_check.append(("Ã–lÃ§Ã¼m", measurement_date))
        
        if report_date != "BulunamadÄ±" and report_date != "Rapor tarihi ayrÄ± belirtilmemiÅŸ":
            dates_to_check.append(("Rapor", report_date))
        
        if not dates_to_check:
            return False, "GeÃ§erli tarih bulunamadÄ±"
        
        try:
            date_formats = ['%d/%m/%Y', '%d.%m.%Y', '%d-%m-%Y', '%d/%m/%y', '%d.%m.%y', '%d-%m-%y']
            today = datetime.now()
            
            for date_type, date_str in dates_to_check:
                parsed_date = None
                for fmt in date_formats:
                    try:
                        parsed_date = datetime.strptime(date_str, fmt)
                        break
                    except:
                        continue
                
                if not parsed_date:
                    return False, f"{date_type} tarihi formatÄ± tanÄ±nmadÄ± ({date_str})"
                
                if parsed_date.year < 1950:
                    parsed_date = parsed_date.replace(year=parsed_date.year + 2000)
                
                diff = abs((today - parsed_date).days)
                
                if diff > 365:
                    return False, f"{date_type} tarihi 1 yÄ±ldan eski ({diff} gÃ¼n Ã¶nce - {date_str})"
            
            checked_dates = [f"{dt[0]}: {dt[1]}" for dt in dates_to_check]
            return True, f"TÃ¼m tarihler geÃ§erli - {', '.join(checked_dates)}"
            
        except Exception as e:
            return False, f"Tarih kontrolÃ¼ yapÄ±lamadÄ± - {e}"

    def generate_improvement_actions(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Dinamik iyileÅŸtirme Ã¶nerileri oluÅŸtur - DB'den gelen actions ile"""
        actions = []
        
        sorted_categories = sorted(scores["category_scores"].items(), key=lambda x: x[1]["percentage"])
        
        # DB'den category_actions al (self.category_actions zaten __init__'te yÃ¼klendi)
        category_actions = self.category_actions
        
        # EÄŸer DB'den veri gelmediyse, boÅŸ liste dÃ¶ndÃ¼r
        if not category_actions:
            logger.warning("âš ï¸ Category actions bulunamadÄ±, boÅŸ Ã¶neri listesi dÃ¶ndÃ¼rÃ¼lÃ¼yor")
            if scores["percentage"] < 50:
                return ["Ã–NCELÄ°K: Rapor yapÄ±sÄ±nÄ± ve iÃ§eriÄŸini kapsamlÄ± olarak yeniden dÃ¼zenleyiniz"]
            return []
        
        for category, score_data in sorted_categories[:5]:
            if score_data["percentage"] < 70:
                actions.extend(category_actions.get(category, []))
        
        if scores["percentage"] < 50:
            actions.insert(0, "Ã–NCELÄ°K: Rapor yapÄ±sÄ±nÄ± ve iÃ§eriÄŸini kapsamlÄ± olarak yeniden dÃ¼zenleyiniz")
        
        return actions

    def generate_recommendations(self, analysis_results: Dict, scores: Dict, date_valid: bool = True, date_message: str = "") -> List[str]:
        """Ã–neriler oluÅŸtur"""
        recommendations = []
        total_percentage = scores["percentage"]
        
        if not date_valid:
            if "Ne Ã¶lÃ§Ã¼m ne de rapor tarihi bulunamadÄ±" in date_message:
                recommendations.append("âŒ RAPOR GEÃ‡ERSÄ°Z: Ne Ã¶lÃ§Ã¼m ne de rapor tarihi bulunamadÄ±")
            elif "Ã–lÃ§Ã¼m tarihi formatÄ± tanÄ±nmadÄ±" in date_message:
                recommendations.append("âŒ RAPOR GEÃ‡ERSÄ°Z: Ã–lÃ§Ã¼m tarihi formatÄ± tanÄ±nmadÄ±")
            elif "Rapor tarihi formatÄ± tanÄ±nmadÄ±" in date_message:
                recommendations.append("âŒ RAPOR GEÃ‡ERSÄ°Z: Rapor tarihi formatÄ± tanÄ±nmadÄ±")
            elif "Ã–lÃ§Ã¼m tarihi 1 yÄ±ldan eski" in date_message:
                recommendations.append("âŒ RAPOR GEÃ‡ERSÄ°Z: Ã–lÃ§Ã¼m tarihi 1 yÄ±ldan eski")
            elif "Rapor tarihi 1 yÄ±ldan eski" in date_message:
                recommendations.append("âŒ RAPOR GEÃ‡ERSÄ°Z: Rapor tarihi 1 yÄ±ldan eski")
            elif "GeÃ§erli tarih bulunamadÄ±" in date_message:
                recommendations.append("âŒ RAPOR GEÃ‡ERSÄ°Z: GeÃ§erli tarih bulunamadÄ±")
            else:
                recommendations.append(f"âŒ RAPOR GEÃ‡ERSÄ°Z: {date_message}")
            return recommendations
        
        if total_percentage >= 70:
            recommendations.append(f"âœ… AydÄ±nlatma Ã–lÃ§Ã¼m Raporu GEÃ‡ERLÄ° (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"âŒ AydÄ±nlatma Ã–lÃ§Ã¼m Raporu GEÃ‡ERSÄ°Z (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 40:
                recommendations.append(f"ğŸ”´ {category} bÃ¶lÃ¼mÃ¼ yetersiz (%{category_score:.1f})")
                missing_items = [name for name, result in results.items() if not result.found]
                if missing_items:
                    recommendations.append(f"   Eksik: {', '.join(missing_items[:3])}")
            elif category_score < 70:
                recommendations.append(f"ğŸŸ¡ {category} bÃ¶lÃ¼mÃ¼ geliÅŸtirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"ğŸŸ¢ {category} bÃ¶lÃ¼mÃ¼ yeterli (%{category_score:.1f})")
        
        return recommendations

    def analyze_lighting_report(self, file_path: str) -> Dict[str, Any]:
        """Ana AydÄ±nlatma Ã–lÃ§Ã¼m Raporu analiz fonksiyonu"""
        logger.info("AydÄ±nlatma Ã–lÃ§Ã¼m Raporu analizi baÅŸlatÄ±lÄ±yor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadÄ±: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin Ã§Ä±karÄ±lamadÄ±"}
        
        detected_lang = self.detect_language(text)
        
        if detected_lang != 'tr':
            logger.info(f"{detected_lang.upper()} dilinden TÃ¼rkÃ§e'ye Ã§eviriliyor...")
            text = self.translate_to_turkish(text, detected_lang)
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        
        date_valid, date_message = self.check_date_validity(
            extracted_values.get("olcum_tarihi", "BulunamadÄ±"),
            extracted_values.get("rapor_tarihi", "BulunamadÄ±")
        )
        
        recommendations = self.generate_recommendations(analysis_results, scores, date_valid, date_message)
        improvement_actions = self.generate_improvement_actions(analysis_results, scores)
        
        if not date_valid:
            final_status = "FAIL"
            final_percentage = 0
        else:
            final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
            final_percentage = scores["percentage"]
        
        report = {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "tarih_gecerlilik": {
                "valid": date_valid,
                "message": date_message
            },
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "iyilestirme_eylemleri": improvement_actions,
            "ozet": {
                "toplam_puan": scores["total_score"] if date_valid else 0,
                "yuzde": final_percentage,
                "durum": final_status,
                "rapor_tipi": "AYDINLATMA_OLCUM_RAPORU",
                "tarih_gecersiz": not date_valid
            }
        }
        
        return report  
    
# ============================================
# HELPER FUNCTIONS (Server Validasyon)
# ============================================
def validate_document_server(text, validation_keywords):
    """Server kodunda dokÃ¼man validasyonu - DB'den gelen keywords ile"""
    
    # DB'den gelen critical_terms
    critical_terms_data = validation_keywords.get('critical_terms', [])
    
    # Liste formatÄ±na dÃ¶nÃ¼ÅŸtÃ¼r (eski format uyumluluÄŸu)
    critical_terms = []
    for item in critical_terms_data:
        if isinstance(item, dict) and 'keywords' in item:
            critical_terms.append(item['keywords'])
        elif isinstance(item, list):
            critical_terms.append(item)
    
    # EÄŸer DB'den veri gelmediyse, boÅŸ validasyon (hata verme)
    if not critical_terms:
        logger.warning("âš ï¸ Critical terms bulunamadÄ±, validasyon atlanÄ±yor")
        return True  # VarsayÄ±lan: geÃ§erli kabul et
    
    category_found = []
    
    for i, category in enumerate(critical_terms):
        found_in_category = False
        for term in category:
            if re.search(rf"\b{term}\b", text, re.IGNORECASE):
                found_in_category = True
                logger.info(f"AydÄ±nlatma Kategori {i+1} bulundu: '{term}'")
                break
        category_found.append(found_in_category)
    
    valid_categories = sum(category_found)
    logger.info(f"DokÃ¼man validasyonu: {valid_categories}/{len(critical_terms)} kritik kategori bulundu")
    
    return valid_categories >= len(critical_terms) - 1  # En az (n-1) kategori bulunmalÄ±


def check_strong_keywords_first_pages(filepath, validation_keywords):
    """Ä°lk 1-2 sayfada Ã¶zgÃ¼ kelimeleri OCR ile ara - DB'den gelen keywords ile"""
    
    # DB'den strong keywords al
    strong_keywords = validation_keywords.get('strong_keywords', [])
    
    # EÄŸer DB'den veri gelmediyse, validasyon atla
    if not strong_keywords:
        logger.warning("âš ï¸ Strong keywords bulunamadÄ±, validasyon atlanÄ±yor")
        return True  # VarsayÄ±lan: geÃ§erli kabul et
    
    try:
        Image.MAX_IMAGE_PIXELS = None
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=2)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        
        found_keywords = [kw for kw in strong_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"Ä°lk sayfa kontrol: {len(found_keywords)} Ã¶zgÃ¼ kelime: {found_keywords}")
        return len(found_keywords) >= 2
        
    except Exception as e:
        logger.warning(f"Ä°lk sayfa kontrol hatasÄ±: {e}")
        return False


def check_excluded_keywords_first_pages(filepath, validation_keywords):
    """Ä°lk 1-2 sayfada istenmeyen rapor tÃ¼rlerinin kelimelerini ara - DB'den"""
    
    # DB'den excluded keywords al
    excluded_keywords = validation_keywords.get('excluded_keywords', [])
    
    # EÄŸer DB'den veri gelmediyse, validasyon atla
    if not excluded_keywords:
        logger.warning("âš ï¸ Excluded keywords bulunamadÄ±, validasyon atlanÄ±yor")
        return False  # VarsayÄ±lan: excluded yok kabul et
    
    try:
        Image.MAX_IMAGE_PIXELS = None
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=2)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        
        found_excluded = [kw for kw in excluded_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"Ä°lk sayfa excluded kontrol: {len(found_excluded)} istenmeyen kelime: {found_excluded}")
        return len(found_excluded) >= 2
        
    except Exception as e:
        logger.warning(f"Ä°lk sayfa excluded kontrol hatasÄ±: {e}")
        return False


def get_conclusion_message_aydinlatma(status, percentage):
    """SonuÃ§ mesajÄ±nÄ± dÃ¶ndÃ¼r - AydÄ±nlatma iÃ§in"""
    if status == "PASS":
        return f"AydÄ±nlatma Ã¶lÃ§Ã¼m raporu TS EN 12464-1 ve Ä°SG mevzuatÄ±na uygundur (%{percentage:.0f})"
    else:
        return f"AydÄ±nlatma raporu standartlara uygun deÄŸil (%{percentage:.0f})"


def get_main_issues_aydinlatma(report):
    """Ana sorunlarÄ± listele - AydÄ±nlatma iÃ§in"""
    issues = []
    
    for category, score_data in report['puanlama']['category_scores'].items():
        if score_data['percentage'] < 50:
            issues.append(f"{category} kategorisinde ciddi eksiklikler")
    
    if not issues and report['puanlama']['total_score'] < 50:
        issues = [
            "AydÄ±nlatma seviyesi Ã¶lÃ§Ã¼m sonuÃ§larÄ± eksik",
            "Ã–lÃ§Ã¼m cihazÄ± kalibrasyon bilgileri eksik",
            "Ã‡alÄ±ÅŸma alanÄ± tanÄ±mÄ± ve sÄ±nÄ±flandÄ±rmasÄ± eksik",
            "Yasal uygunluk deÄŸerlendirmesi yapÄ±lmamÄ±ÅŸ"
        ]
    
    return issues[:4]


# ============================================
# FLASK SERVÄ°S KATMANI - CONFIGURATION
# ============================================
app = Flask(__name__)

# Database configuration (YENÄ°)
app.config['SQLALCHEMY_DATABASE_URI'] = Config.SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = Config.SQLALCHEMY_TRACK_MODIFICATIONS
app.config['SQLALCHEMY_ECHO'] = Config.SQLALCHEMY_ECHO

UPLOAD_FOLDER = 'temp_uploads_aydinlatma'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ============================================
# FLASK SERVÄ°S KATMANI - API ENDPOINTS
# ============================================
@app.route('/api/aydinlatma-report', methods=['POST'])
def analyze_aydinlatma_report():
    """AydÄ±nlatma Ã–lÃ§Ã¼m Raporu analiz API endpoint'i - 3 AÅŸamalÄ± Validasyon"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '' or not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file'}), 400

        try:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            logger.info(f"AydÄ±nlatma Ã–lÃ§Ã¼m Raporu kontrol ediliyor: {filename}")

            # Create analyzer instance
            analyzer = LightingReportAnalyzer(app=app)
            
            # ÃœÃ‡ AÅAMALI AYDINLATMA KONTROLÃœ
            logger.info(f"ÃœÃ§ aÅŸamalÄ± aydÄ±nlatma kontrolÃ¼ baÅŸlatÄ±lÄ±yor: {filename}")
            file_ext = os.path.splitext(filepath)[1].lower()

            if file_ext == '.pdf':
                # PDF iÃ§in Ã¼Ã§ aÅŸamalÄ± kontrol (OCR dahil)
                logger.info("AÅŸama 1: Ä°lk sayfa aydÄ±nlatma Ã¶zgÃ¼ kelime kontrolÃ¼...")
                if check_strong_keywords_first_pages(filepath, analyzer.validation_keywords):
                    logger.info("âœ… AÅŸama 1 geÃ§ti - AydÄ±nlatma Ã¶zgÃ¼ kelimeler bulundu")
                else:
                    logger.info("AÅŸama 2: Ä°lk sayfa excluded kelime kontrolÃ¼...")
                    if check_excluded_keywords_first_pages(filepath, analyzer.validation_keywords):
                        logger.info("âŒ AÅŸama 2'de excluded kelimeler bulundu - AydÄ±nlatma deÄŸil")
                        try:
                            os.remove(filepath)
                        except:
                            pass
                        return jsonify({
                            'error': 'Invalid document type',
                            'message': 'Bu dosya aydÄ±nlatma raporu deÄŸil (farklÄ± rapor tÃ¼rÃ¼ tespit edildi). LÃ¼tfen aydÄ±nlatma Ã¶lÃ§Ã¼m raporu yÃ¼kleyiniz.',
                            'details': {
                                'filename': filename,
                                'document_type': 'OTHER_REPORT_TYPE',
                                'required_type': 'AYDINLATMA_OLCUM_RAPORU'
                            }
                        }), 400
                    else:
                        # AÅAMA 3: PyPDF2 ile tam dokÃ¼man kontrolÃ¼
                        logger.info("AÅŸama 3: Tam dokÃ¼man critical terms kontrolÃ¼...")
                        try:
                            with open(filepath, 'rb') as pdf_file:
                                pdf_reader = PyPDF2.PdfReader(pdf_file)
                                text = ""
                                for page in pdf_reader.pages:
                                    text += page.extract_text() + "\n"
                            
                            if not text or len(text.strip()) < 50:
                                try:
                                    os.remove(filepath)
                                except:
                                    pass
                                return jsonify({
                                    'error': 'Text extraction failed',
                                    'message': 'Dosyadan yeterli metin Ã§Ä±karÄ±lamadÄ±'
                                }), 400
                            
                            if not validate_document_server(text, analyzer.validation_keywords):
                                try:
                                    os.remove(filepath)
                                except:
                                    pass
                                return jsonify({
                                    'error': 'Invalid document type',
                                    'message': 'YÃ¼klediÄŸiniz dosya aydÄ±nlatma Ã¶lÃ§Ã¼m raporu deÄŸil! LÃ¼tfen geÃ§erli bir aydÄ±nlatma raporu yÃ¼kleyiniz.',
                                    'details': {
                                        'filename': filename,
                                        'document_type': 'NOT_AYDINLATMA_REPORT',
                                        'required_type': 'AYDINLATMA_OLCUM_RAPORU'
                                    }
                                }), 400
                                
                        except Exception as e:
                            logger.error(f"AÅŸama 3 hatasÄ±: {e}")
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({
                                'error': 'Analysis failed',
                                'message': 'Dosya analizi sÄ±rasÄ±nda hata oluÅŸtu'
                            }), 500

            elif file_ext in ['.docx', '.doc', '.txt']:
                # DOCX/TXT iÃ§in sadece tam dokÃ¼man kontrolÃ¼
                logger.info(f"DOCX/TXT dosyasÄ± iÃ§in tam dokÃ¼man kontrolÃ¼: {file_ext}")
                text = ""
                if file_ext in ['.docx', '.doc']:
                    text = analyzer.extract_text_from_docx(filepath)
                elif file_ext == '.txt':
                    text = analyzer.extract_text_from_txt(filepath)
                
                if not text or len(text.strip()) == 0:
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Text extraction failed',
                        'message': 'Dosyadan yeterli metin Ã§Ä±karÄ±lamadÄ±'
                    }), 400
                
                if not validate_document_server(text, analyzer.validation_keywords):
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Invalid document type',
                        'message': 'YÃ¼klediÄŸiniz dosya aydÄ±nlatma Ã¶lÃ§Ã¼m raporu deÄŸil! LÃ¼tfen geÃ§erli bir aydÄ±nlatma raporu yÃ¼kleyiniz.',
                        'details': {
                            'filename': filename,
                            'document_type': 'NOT_AYDINLATMA_REPORT',
                            'required_type': 'AYDINLATMA_OLCUM_RAPORU'
                        }
                    }), 400

            # Buraya kadar geldiyse aydÄ±nlatma raporu, ÅŸimdi analizi yap
            logger.info(f"AydÄ±nlatma raporu doÄŸrulandÄ±, analiz baÅŸlatÄ±lÄ±yor: {filename}")
            report = analyzer.analyze_lighting_report(filepath)
            
            try:
                os.remove(filepath)
            except:
                pass

            if 'error' in report:
                return jsonify({'error': 'Analysis failed', 'message': report['error']}), 400

            if 'puanlama' in report and 'percentage' in report['puanlama']:
                overall_percentage = report['puanlama']['percentage']
            elif 'ozet' in report and 'yuzde' in report['ozet']:
                overall_percentage = report['ozet']['yuzde']
            else:
                overall_percentage = 0
            status = "PASS" if overall_percentage >= 70 else "FAIL"
            
            response_data = {
                'analysis_date': report.get('analiz_tarihi', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
                'analysis_details': {
                    'found_criteria': len([c for results in report['kategori_analizleri'].values() for c in results.values() if c.found]),
                    'total_criteria': len([c for results in report['kategori_analizleri'].values() for c in results.values()]),
                    'percentage': round(overall_percentage, 1)
                },
                'analysis_id': f"aydinlatma_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                'category_scores': {},
                'extracted_values': report['cikarilan_degerler'],
                'file_type': 'AYDINLATMA_OLCUM_RAPORU',
                'filename': filename,
                'language_info': {
                    'detected_language': report['dosya_bilgisi'].get('detected_language', 'turkish').upper(),
                    'file_type': file_ext.replace('.', '')
                },
                'overall_score': {
                    'percentage': round(overall_percentage, 2),
                    'total_points': report['ozet']['toplam_puan'],
                    'max_points': 100,
                    'status': status,
                    'status_tr': 'GEÃ‡ERLÄ°' if status == "PASS" else 'GEÃ‡ERSÄ°Z',
                    'text_quality': 'good'
                },
                'recommendations': report['oneriler'],
                'summary': {
                    'is_valid': status == "PASS",
                    'conclusion': get_conclusion_message_aydinlatma(status, overall_percentage),
                    'main_issues': [] if status == "PASS" else get_main_issues_aydinlatma(report)
                }
            }
            
            for category, score_data in report['puanlama']['category_scores'].items():
                response_data['category_scores'][category] = {
                    'score': score_data.get('normalized', 0),
                    'max_score': score_data.get('max_weight', 0),
                    'percentage': score_data.get('percentage', 0),
                    'status': 'PASS' if score_data.get('percentage', 0) >= 70 else 'CONDITIONAL' if score_data.get('percentage', 0) >= 50 else 'FAIL'
                }

            return jsonify({
                'success': True,
                'message': 'AydÄ±nlatma Ã–lÃ§Ã¼m Raporu baÅŸarÄ±yla analiz edildi',
                'analysis_service': 'aydinlatma',
                'service_description': 'AydÄ±nlatma Ã–lÃ§Ã¼m Raporu Analizi',
                'data': response_data
            })

        except Exception as analysis_error:
            try:
                if 'filepath' in locals():
                    os.remove(filepath)
            except:
                pass
            logger.error(f"Analiz hatasÄ±: {str(analysis_error)}")
            return jsonify({'error': 'Analysis failed', 'message': str(analysis_error)}), 500

    except Exception as e:
        logger.error(f"API endpoint hatasÄ±: {str(e)}")
        return jsonify({'error': 'Server error', 'message': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'Lighting Report Analyzer API',
        'version': '1.0.0',
        'tesseract_available': tesseract_available,
        'report_type': 'AYDINLATMA_OLCUM_RAPORU',
        'standard': 'TS EN 12464-1'
    })


@app.route('/', methods=['GET'])
def index():
    """API bilgileri"""
    return jsonify({
        'service': 'Lighting Report Analyzer API',
        'version': '1.0.0',
        'description': 'AydÄ±nlatma Ã–lÃ§Ã¼m RaporlarÄ±nÄ± analiz eden REST API servisi',
        'endpoints': {
            'POST /api/aydinlatma-report': 'AydÄ±nlatma raporu analizi',
            'GET /api/health': 'Servis saÄŸlÄ±k kontrolÃ¼',
            'GET /': 'Bu bilgi sayfasÄ±'
        }
    })


# ============================================
# DATABASE INITIALIZATION
# ============================================
with app.app_context():
    db.init_app(app)


# ============================================
# APPLICATION ENTRY POINT (Azure-Friendly)
# ============================================
if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("AydÄ±nlatma Ã–lÃ§Ã¼m Raporu Analiz Servisi")
    logger.info("=" * 60)
    
    port = int(os.environ.get('PORT', 8008))
    
    logger.info(f"ğŸš€ Servis baÅŸlatÄ±lÄ±yor - Port: {port}")
    logger.info(f"ğŸ“ Upload klasÃ¶rÃ¼: {UPLOAD_FOLDER}")
    logger.info(f"ğŸ”§ Tesseract: {tesseract_info}")
    logger.info("=" * 60)
    
    app.run(host='0.0.0.0', port=port, debug=False)