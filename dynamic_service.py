# ============================================
# DYNAMIC DOCUMENT TYPE ANALYZER
# Tüm yeni döküman türlerini yöneten tek servis
# Database-driven configuration
# Port: 8018
# ============================================

# ============================================
# IMPORTS
# ============================================
import os
import re
from datetime import datetime
from typing import Dict, List, Any
import PyPDF2
from docx import Document
from dataclasses import dataclass
import logging
import pytesseract
import cv2
import numpy as np
from PIL import Image
import pdf2image

from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename

# ============================================
# DATABASE IMPORTS
# ============================================
from database import db, init_db
from models import DocumentType, CriteriaWeight, CriteriaDetail, PatternDefinition, ValidationKeyword
from config import Config

# ============================================
# LOGGING CONFIGURATION
# ============================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================
# LANGUAGE DETECTION
# ============================================
try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False
    logger.warning("langdetect modülü bulunamadı - dil tespiti devre dışı")

# ============================================
# ANALİZ SINIFI - DATA CLASS
# ============================================
@dataclass
class DynamicAnalysisResult:
    """Dynamic analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

# ============================================
# ANALİZ SINIFI - MAIN ANALYZER
# ============================================
class DynamicReportAnalyzer:
    """Dynamic döküman türleri için analiz sınıfı"""
    
    def __init__(self, document_code: str, app=None):
        """
        Args:
            document_code: Döküman türü kodu (örn: 'yangin_sondurme')
            app: Flask app instance
        """
        self.document_code = document_code
        logger.info(f"Dynamic analyzer başlatılıyor: {document_code}")
        
        if app:
            with app.app_context():
                try:
                    # Database'den döküman türünü bul
                    self.doc_type = DocumentType.query.filter_by(code=document_code, is_active=True).first()
                    
                    if not self.doc_type:
                        raise ValueError(f"Döküman türü bulunamadı: {document_code}")
                    
                    # Config'i yükle
                    self.load_config()
                    
                    logger.info(f"✅ {self.doc_type.name} config yüklendi")
                    logger.info(f"   - OCR: {'Evet' if self.use_ocr else 'Hayır'}")
                    logger.info(f"   - Kategoriler: {len(self.criteria_weights)}")
                    
                except Exception as e:
                    logger.error(f"⚠️ Config yükleme hatası: {e}")
                    raise
        else:
            raise ValueError("Flask app context gerekli")
    
    def load_config(self):
        """Database'den config yükle"""
        # Criteria Weights
        self.criteria_weights = {}
        criteria_weights = CriteriaWeight.query.filter_by(
            document_type_id=self.doc_type.id
        ).order_by(CriteriaWeight.display_order).all()
        
        for cw in criteria_weights:
            self.criteria_weights[cw.category_name] = cw.weight
        
        # Criteria Details
        self.criteria_details = {}
        for cw in criteria_weights:
            details = CriteriaDetail.query.filter_by(
                criteria_weight_id=cw.id
            ).order_by(CriteriaDetail.display_order).all()
            
            self.criteria_details[cw.category_name] = {
                cd.criterion_name: {
                    "pattern": cd.pattern,
                    "weight": cd.weight
                }
                for cd in details
            }
        
        # Pattern Definitions (extract_specific_values)
        self.pattern_definitions = {}
        patterns = PatternDefinition.query.filter_by(
            document_type_id=self.doc_type.id
        ).all()
        
        for pattern in patterns:
            if pattern.pattern_group not in self.pattern_definitions:
                self.pattern_definitions[pattern.pattern_group] = {}
            self.pattern_definitions[pattern.pattern_group][pattern.field_name] = pattern.patterns
        
        # Validation Keywords
        self.validation_keywords = {}
        
        # Critical terms
        critical_terms = ValidationKeyword.query.filter_by(
            document_type_id=self.doc_type.id,
            keyword_type='critical_terms'
        ).all()
        self.validation_keywords['critical_terms'] = [vk.keywords for vk in critical_terms]
        
        # Strong keywords
        strong = ValidationKeyword.query.filter_by(
            document_type_id=self.doc_type.id,
            keyword_type='strong_keywords'
        ).first()
        self.validation_keywords['strong_keywords'] = strong.keywords if strong else []
        
        # Excluded keywords
        excluded = ValidationKeyword.query.filter_by(
            document_type_id=self.doc_type.id,
            keyword_type='excluded_keywords'
        ).first()
        self.validation_keywords['excluded_keywords'] = excluded.keywords if excluded else []
        
        # OCR kullanımı (service_file'a bakarak belirle - gelecekte config'e eklenebilir)
        # Şimdilik default: Varsa strong_keywords varsa OCR kullan
        self.use_ocr = self.doc_type.needs_ocr
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosyadan metin çıkar (OCR ayarına göre)"""
        try:
            text = ""
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                # OCR KONTROLÜ - ÖNCE KARAR VER! 👇
                if self.doc_type.needs_ocr:
                    # OCR İŞARETLİ → DİREKT OCR KULLAN!
                    logger.info("✅ OCR etkin - Direkt OCR kullanılıyor (PyPDF2 atlanıyor)")
                    return self.extract_text_with_ocr(file_path)
                else:
                    # OCR İŞARETSİZ → SADECE PyPDF2
                    logger.info("📄 OCR devre dışı - Sadece PyPDF2 kullanılıyor")
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    
                    text = text.strip()
                    
                    if len(text) > 100:
                        logger.info(f"✅ PyPDF2 başarılı - {len(text)} karakter")
                        return text
                    else:
                        logger.error("❌ PyPDF2 yetersiz metin çıkardı - OCR devre dışı!")
                        return ""
            
            elif file_ext in ['.jpg', '.jpeg', '.png']:
                # Görsel dosyalar için OCR ZORUNLU!
                if self.doc_type.needs_ocr:
                    logger.info("📷 Görsel dosya - OCR kullanılıyor...")
                    return self.extract_text_with_ocr(file_path)
                else:
                    logger.error("❌ OCR devre dışı - Görsel dosya işlenemez!")
                    return ""
            
            elif file_ext in ['.docx', '.doc']:
                return self.extract_text_from_docx(file_path)
            
            elif file_ext == '.txt':
                return self.extract_text_from_txt(file_path)
            
            else:
                logger.error(f"❌ Desteklenmeyen dosya formatı: {file_ext}")
                return ""
            
        except Exception as e:
            logger.error(f"❌ Metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den PyPDF2 ile metin çıkar"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    all_text += page_text + "\n"
                
                all_text = all_text.strip()
                logger.info(f"PDF'den {len(all_text)} karakter çıkarıldı")
                return all_text
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_with_ocr(self, file_path: str) -> str:
        """OCR ile metin çıkar"""
        try:
            pages = pdf2image.convert_from_path(file_path, dpi=200)
            all_text = ""
            
            for i, page in enumerate(pages):
                opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
                gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
                processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
                
                text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
                all_text += text + "\n"
                logger.info(f"OCR sayfa {i+1}/{len(pages)} tamamlandı")
            
            logger.info(f"OCR ile {len(all_text)} karakter çıkarıldı")
            return all_text.strip()
        except Exception as e:
            logger.error(f"OCR hatası: {e}")
            return ""
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkar"""
        try:
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasından metin çıkar"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except:
            try:
                with open(txt_path, 'r', encoding='cp1254') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"TXT metin çıkarma hatası: {e}")
                return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, DynamicAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            # 👇 YENİ - DEBUG LOG
            logger.info(f"   🔍 {criterion_name}:")
            logger.info(f"      Pattern: {pattern[:150]}...")
            logger.info(f"      Sonuç: {len(matches)} match bulundu")
            if matches:
                logger.info(f"      Örnek: {str(matches[:2])}")
                
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = DynamicAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={"pattern_used": pattern, "matches_count": len(matches) if matches else 0}
            )
        
        return results
    
    def calculate_scores(self, analysis_results: Dict[str, Dict[str, DynamicAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }
    
    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """Döküman'a özgü değerleri çıkar"""
        values = {}
        
        extract_values = self.pattern_definitions.get('extract_values', {})
        
        for field_name, patterns_list in extract_values.items():
            values[field_name] = "Bulunamadı"
            
            for pattern in patterns_list:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    try:
                        values[field_name] = match.group(1).strip()
                    except:
                        values[field_name] = match.group(0).strip()
                    break
        
        return values
    
    def generate_recommendations(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Detaylı öneriler oluştur - Eksik kriterleri göster"""
        recommendations = []
        
        total_percentage = scores["percentage"]
        
        # Genel durum
        if total_percentage >= 70:
            recommendations.append(f"✅ {self.doc_type.name} GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ {self.doc_type.name} GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        # Kategori bazlı detaylı analiz
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            category_earned = scores["category_scores"][category]["earned"]
            category_possible = scores["category_scores"][category]["possible"]
            
            # Kategori durumu
            if category_score < 40:
                status = "YETERSİZ"
            elif category_score < 70:
                status = "GELİŞTİRİLMELİ"
            else:
                status = "YETERLİ"
            
            # Eksik/Zayıf kriterleri topla
            criteria_details = []
            
            for criterion_name, result in results.items():
                if result.score == 0:
                    # Hiç puan alamamış
                    criteria_details.append(f"{criterion_name}: Bulunamadı (0/{result.max_score} puan)")
                elif result.score < result.max_score * 0.5:
                    # %50'den az puan almış
                    criteria_details.append(f"{criterion_name}: Yetersiz ({result.score}/{result.max_score} puan)")
                elif result.score < result.max_score:
                    # Tam puan almamış ama %50'den fazla
                    criteria_details.append(f"{criterion_name}: Geliştirilmeli ({result.score}/{result.max_score} puan)")
            
            # Kategori satırı - Emoji + Kategori + Puan + Kriterler
            if criteria_details:
                # Eksik kriterler varsa listele
                emoji = "🔴" if category_score < 40 else "🟡" if category_score < 70 else "🟢"
                criteria_str = ", ".join(criteria_details)
                recommendations.append(
                    f"{emoji} {category}: {criteria_str}"
                )
            else:
                # Tüm kriterler tam puan
                recommendations.append(
                    f"🟢 {category}: Tüm kriterler yeterli (%{category_score:.1f})"
                )
        
        return recommendations
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """Ana analiz fonksiyonu"""
        logger.info(f"{self.doc_type.name} analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        recommendations = self.generate_recommendations(analysis_results, scores)
        
        final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
        
        return {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "ozet": {
                "toplam_puan": scores["total_score"],
                "yuzde": scores["percentage"],
                "durum": final_status,
                "rapor_tipi": self.doc_type.code.upper()
            }
        }


# ============================================
# HELPER FUNCTIONS - VALIDATION
# ============================================
def validate_document_server(text, validation_keywords):
    """Döküman validasyonu"""
    critical_terms_data = validation_keywords.get('critical_terms', [])
    
    if not critical_terms_data:
        logger.warning("⚠️ Critical terms bulunamadı, validasyon atlanıyor")
        return True
    
    category_found = []
    
    for i, category in enumerate(critical_terms_data):
        found_in_category = False
        for term in category:
            if re.search(rf"\b{term}\b", text, re.IGNORECASE):
                found_in_category = True
                logger.info(f"Kategori {i+1} bulundu: '{term}'")
                break
        category_found.append(found_in_category)
    
    valid_categories = sum(category_found)
    logger.info(f"Döküman validasyonu: {valid_categories}/{len(critical_terms_data)} kritik kategori bulundu")
    
    return valid_categories >= len(critical_terms_data) - 1


def check_strong_keywords_first_pages(filepath, validation_keywords):
    """İlk 1-2 sayfada özgü kelimeleri OCR ile ara"""
    strong_keywords = validation_keywords.get('strong_keywords', [])
    
    if not strong_keywords:
        logger.warning("⚠️ Strong keywords bulunamadı, validasyon atlanıyor")
        return True
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=2)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        
        found_keywords = [kw for kw in strong_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa kontrol: {len(found_keywords)} özgü kelime")
        return len(found_keywords) >= 1
    except Exception as e:
        logger.warning(f"İlk sayfa kontrol hatası: {e}")
        return False


def check_excluded_keywords_first_pages(filepath, validation_keywords):
    """İlk 1-2 sayfada istenmeyen kelimeleri ara"""
    excluded_keywords = validation_keywords.get('excluded_keywords', [])
    
    if not excluded_keywords:
        logger.warning("⚠️ Excluded keywords bulunamadı, validasyon atlanıyor")
        return False
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=2)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        
        found_excluded = [kw for kw in excluded_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa excluded kontrol: {len(found_excluded)} istenmeyen kelime - Bulunanlar: {found_excluded}") 
        return len(found_excluded) >= 1
    except Exception as e:
        logger.warning(f"İlk sayfa excluded kontrol hatası: {e}")
        return False


# ============================================
# FLASK SERVICE
# ============================================
app = Flask(__name__)

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = Config.SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = Config.SQLALCHEMY_TRACK_MODIFICATIONS
app.config['SQLALCHEMY_ECHO'] = Config.SQLALCHEMY_ECHO

UPLOAD_FOLDER = 'temp_uploads_dynamic'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/api/dynamic-report', methods=['POST'])
def analyze_dynamic_report():
    """Dynamic döküman analiz endpoint - 3 Aşamalı Validasyon"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        # Document code al
        document_code = request.form.get('document_code')
        if not document_code:
            return jsonify({'error': 'document_code gerekli'}), 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        logger.info(f"Dynamic analiz başlatılıyor: {document_code} - {filename}")

        # Analyzer oluştur
        analyzer = DynamicReportAnalyzer(document_code, app=app)
        
        # ÜÇ AŞAMALI KONTROL (sadece PDF için)
        file_ext = os.path.splitext(filepath)[1].lower()
        
        if file_ext == '.pdf':
            logger.info("Aşama 1: İlk sayfa özgü kelime kontrolü...")
            if check_strong_keywords_first_pages(filepath, analyzer.validation_keywords):
                logger.info("✅ Aşama 1 geçti")
            else:
                logger.info("Aşama 2: İlk sayfa excluded kelime kontrolü...")
                if check_excluded_keywords_first_pages(filepath, analyzer.validation_keywords):
                    logger.info("❌ Aşama 2'de excluded kelimeler bulundu")
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Invalid document type',
                        'message': f'Bu dosya {analyzer.doc_type.name} değil (farklı rapor türü tespit edildi).'
                    }), 400
                else:
                    logger.info("Aşama 3: Tam doküman critical terms kontrolü...")
                    try:
                        with open(filepath, 'rb') as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            text = ""
                            for page in pdf_reader.pages:
                                text += page.extract_text() + "\n"
                        
                        if not text or len(text.strip()) == 0:
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({'error': 'Text extraction failed'}), 400
                        
                        if not validate_document_server(text, analyzer.validation_keywords):
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({
                                'error': 'Invalid document type',
                                'message': f'Yüklediğiniz dosya {analyzer.doc_type.name} değil!'
                            }), 400
                    except Exception as e:
                        logger.error(f"Aşama 3 hatası: {e}")
                        try:
                            os.remove(filepath)
                        except:
                            pass
                        return jsonify({'error': 'Analysis failed'}), 500

        # Analizi yap
        logger.info(f"Döküman doğrulandı, analiz başlatılıyor: {filename}")
        analysis_result = analyzer.analyze_document(filepath)
        
        try:
            os.remove(filepath)
        except:
            pass

        if 'error' in analysis_result:
            return jsonify({'error': 'Analysis failed', 'message': analysis_result['error']}), 400

        overall_percentage = analysis_result['puanlama']['percentage']
        status = "PASS" if overall_percentage >= 70 else "FAIL"
        
        response_data = {
            'analysis_date': analysis_result.get('analiz_tarihi'),
            'analysis_id': f"{document_code}_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            'category_scores': {},
            'extracted_values': analysis_result.get('cikarilan_degerler', {}),
            'file_type': document_code.upper(),
            'filename': filename,
            'overall_score': {
                'percentage': round(overall_percentage, 2),
                'total_points': analysis_result['puanlama']['total_score'],
                'max_points': 100,
                'status': status,
                'status_tr': 'GEÇERLİ' if status == "PASS" else 'GEÇERSİZ'
            },
            'recommendations': analysis_result.get('oneriler', []),
            'summary': {
                'is_valid': status == "PASS",
                'conclusion': f"{analyzer.doc_type.name} analiz sonucu: {status}",
                'main_issues': []
            }
        }
        
        for category, score_data in analysis_result['puanlama']['category_scores'].items():
            response_data['category_scores'][category] = {
                'score': score_data['normalized'],
                'max_score': score_data['max_weight'],
                'percentage': score_data['percentage'],
                'status': 'PASS' if score_data['percentage'] >= 70 else 'FAIL'
            }

        return jsonify({
            'success': True,
            'message': f'{analyzer.doc_type.name} başarıyla analiz edildi',
            'analysis_service': 'dynamic',
            'data': response_data
        })

    except Exception as e:
        logger.error(f"API error: {str(e)}")
        return jsonify({'error': 'Server error', 'message': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check"""
    return jsonify({
        'status': 'healthy',
        'service': 'Dynamic Document Analyzer API',
        'version': '1.0.0'
    })


@app.route('/', methods=['GET'])
def index():
    """Index page"""
    return jsonify({
        'service': 'Dynamic Document Analyzer API',
        'version': '1.0.0',
        'endpoints': {
            'POST /api/dynamic-report': 'Dynamic döküman analizi',
            'GET /api/health': 'Health check',
            'GET /': 'API info'
        }
    })


# ============================================
# DATABASE INITIALIZATION
# ============================================
with app.app_context():
    db.init_app(app)


# ============================================
# APPLICATION ENTRY POINT (Azure-Friendly)
# ============================================
if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("Dynamic Document Analyzer Service")
    logger.info("=" * 60)
    
    port = int(os.environ.get('PORT', 8018))
    logger.info(f"🚀 Servis başlatılıyor - Port: {port}")
    
    app.run(host='0.0.0.0', port=port, debug=False)