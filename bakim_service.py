"""
Bakım Talimatları Analiz Servisi
=================================
Azure App Service için optimize edilmiş standalone servis

Endpoint: POST /api/bakimtalimatlari-report
Health Check: GET /api/health
Port: 8014
"""

# ============================================
# IMPORTS
# ============================================
import os
import re
from datetime import datetime
from typing import Dict, List, Any
import PyPDF2
from docx import Document
from dataclasses import dataclass
import logging
import pytesseract
import cv2
import numpy as np
from PIL import Image
import pdf2image

from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename

# ============================================
# DATABASE IMPORTS
# ============================================
from flask import current_app
from database import db, init_db
from db_loader import load_service_config
from config import Config

# ============================================
# LOGGING CONFIGURATION
# ============================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================
# LANGUAGE DETECTION
# ============================================
try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False
    logger.warning("langdetect modülü bulunamadı - dil tespiti devre dışı")

# ============================================
# TESSERACT CHECK
# ============================================
def check_tesseract_installation():
    """Check if Tesseract is properly installed"""
    try:
        version = pytesseract.get_tesseract_version()
        logger.info(f"Tesseract OCR kurulu - Sürüm: {version}")
        return True, f"Tesseract v{version}"
    except Exception as e:
        logger.error(f"Tesseract OCR kurulu değil: {e}")
        return False, str(e)

tesseract_available, tesseract_info = check_tesseract_installation()

# ============================================
# DATA CLASSES
# ============================================
@dataclass
class MaintenanceAnalysisResult:
    """Bakım Talimatları analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

# ============================================
# ANALYZER CLASS
# ============================================
class MaintenanceReportAnalyzer:
    """Bakım Talimatları rapor analiz sınıfı"""
    
    def __init__(self, app=None):
        logger.info("Bakım Talimatları analiz sistemi başlatılıyor...")
        
        if app:
            with app.app_context():
                try:
                    config = load_service_config('maintenance_instructions')
                    
                    self.criteria_weights = config.get('criteria_weights', {})
                    self.criteria_details = config.get('criteria_details', {})
                    self.pattern_definitions = config.get('pattern_definitions', {})
                    self.validation_keywords = config.get('validation_keywords', {})
                    self.category_actions = config.get('category_actions', {})
                    
                    logger.info(f"✅ Veritabanından yüklendi: {len(self.criteria_weights)} kategori")
                    
                except Exception as e:
                    logger.error(f"⚠️ Veritabanından yükleme başarısız: {e}")
                    logger.warning("⚠️ Fallback: Boş config kullanılıyor")
                    self.criteria_weights = {}
                    self.criteria_details = {}
                    self.pattern_definitions = {}
                    self.validation_keywords = {}
                    self.category_actions = {}
        else:
            logger.warning("⚠️ Flask app context yok, boş config kullanılıyor")
            self.criteria_weights = {}
            self.criteria_details = {}
            self.pattern_definitions = {}
            self.validation_keywords = {}
            self.category_actions = {}
    
    def detect_language(self, text: str) -> str:
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def translate_to_turkish(self, text: str, source_lang: str) -> str:
        if source_lang != 'tr':
            logger.info(f"Tespit edilen dil: {source_lang.upper()} - Çeviri yapılmıyor, orijinal metin kullanılıyor")
        return text
    
    def extract_text_from_file(self, file_path: str) -> str:
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatı: {file_ext}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """PDF'den sadece PyPDF2 ile metin çıkarma"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                all_text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    page_text = re.sub(r'\s+', ' ', page_text)
                    all_text += page_text + "\n"
                return all_text.strip()
        except Exception as e:
            logger.error(f"PDF metin çıkarma hatası: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except:
            try:
                with open(txt_path, 'r', encoding='cp1254') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"TXT metin çıkarma hatası: {e}")
                return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, MaintenanceAnalysisResult]:
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data.get("pattern", "")
            weight = criterion_data.get("weight", 0)
            
            if not pattern:
                logger.warning(f"Pattern bulunamadı: {criterion_name}")
                continue
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = MaintenanceAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={"pattern_used": pattern, "matches_count": len(matches) if matches else 0}
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict) -> Dict[str, Any]:
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """Spesifik değerleri çıkar - DB'den pattern'lerle"""
        values = {
            "makine_adi": "Bulunamadı",
            "makine_modeli": "Bulunamadı",
            "seri_numarasi": "Bulunamadı",
            "bakim_turu": "Bulunamadı",
            "yetkili_personel": "Bulunamadı"
        }
        
        extract_patterns = self.pattern_definitions.get('extract_values', {})
        
        # Makine Adı
        machine_title_patterns = extract_patterns.get('makine_adi', [])
        for pattern in machine_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 3 <= len(result) <= 50 and not result.lower().startswith(("bu", "şu", "o ")):
                    values["makine_adi"] = result
                    break
        
        if values["makine_adi"] == "Bulunamadı":
            machine_field_patterns = extract_patterns.get('mmachine_field_patterns', [])
            for pattern in machine_field_patterns:
                match = re.search(pattern, text)
                if match:
                    result = match.group(1).strip()
                    cleanup_words = ["PROJESİ", "PROJECT", "SİSTEMİ", "SYSTEM", "EKİPMANI", "EQUIPMENT", "TALİMATI", "INSTRUCTION"]
                    for cleanup_word in cleanup_words:
                        result = re.sub(rf"\b{cleanup_word}\b", "", result, flags=re.IGNORECASE).strip()
                    result = re.sub(r'\s+', ' ', result).strip()
                    if 3 <= len(result) <= 50 and not result.lower().startswith(("bu", "şu", "o ")):
                        values["makine_adi"] = result
                        break
        
        # Model
        model_title_patterns = extract_patterns.get('makine_modeli', [])
        for pattern in model_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                if 2 <= len(result) <= 25:
                    values["makine_modeli"] = result
                    break
        
        if values["makine_modeli"] == "Bulunamadı":
            model_field_patterns = extract_patterns.get('model_field_patterns', [])
            for pattern in model_field_patterns:
                match = re.search(pattern, text)
                if match:
                    result = match.group(1).strip()
                    filter_words = ["konveyör", "tipi", "çeşit", "türü", "xxx", "tbd", "n/a"]
                    if not any(word in result.lower() for word in filter_words):
                        if 2 <= len(result) <= 25:
                            values["makine_modeli"] = result
                            break
        
        # Seri No
        serial_title_patterns = extract_patterns.get('seri_numarasi', [])
        placeholder_patterns = extract_patterns.get('placeholder_patterns', [])
        
        for pattern in serial_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                is_placeholder = any(re.match(p, result) for p in placeholder_patterns)
                if not is_placeholder and 3 <= len(result) <= 30:
                    values["seri_numarasi"] = result
                    break
        
        # Bakım Türü
        maintenance_title_patterns = extract_patterns.get('bakim_turu', [])
        for pattern in maintenance_title_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                result = match.group(1).strip()
                if len(result) <= 35:
                    values["bakim_turu"] = result
                    break
        
        if values["bakim_turu"] == "Bulunamadı":
            maintenance_word_patterns = extract_patterns.get('maintenance_word_patterns', [])
            for pattern in maintenance_word_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    result = match.group().strip()
                    if len(result) <= 35:
                        values["bakim_turu"] = result
                        break
        
        # Yetkili Personel
        personnel_title_patterns = extract_patterns.get('yetkili_personel', [])
        for pattern in personnel_title_patterns:
            match = re.search(pattern, text)
            if match:
                result = match.group(1).strip()
                filter_words = ["saklanması", "gereken", "yapılması", "bulunması", "olması", "edilmesi", "sağlanması"]
                if not any(word in result.lower() for word in filter_words):
                    if len(result.split()) >= 2 or (len(result.split()) == 1 and len(result) >= 4):
                        if 3 <= len(result) <= 35:
                            values["yetkili_personel"] = result
                            break
        
        return values

    def validate_maintenance_document(self, text: str) -> bool:
        """Validasyon - DB'den keywords"""
        maintenance_keywords = self.validation_keywords.get('critical_terms', [])
        
        if not maintenance_keywords:
            logger.warning("⚠️ Critical terms bulunamadı")
            return True
        
        critical_terms = []
        for item in maintenance_keywords:
            if isinstance(item, dict) and 'keywords' in item:
                critical_terms.append(item['keywords'])
            elif isinstance(item, list):
                critical_terms.append(item)
        
        found_keywords = 0
        found_words = []
        for category in critical_terms:
            for keyword in category:
                if re.search(rf"\b{keyword}\b", text, re.IGNORECASE):
                    found_keywords += 1
                    found_words.append(keyword)
                    break
        
        logger.info(f"Doküman validasyonu: {found_keywords} kategori bulundu")
        return found_keywords >= 2

    def generate_improvement_actions(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """İyileştirme önerileri - DB'den"""
        actions = []
        sorted_categories = sorted(scores["category_scores"].items(), key=lambda x: x[1]["percentage"])
        
        for category, score_data in sorted_categories[:5]:
            if score_data["percentage"] < 70:
                category_actions_list = self.category_actions.get(category, [])
                actions.extend(category_actions_list)
        
        if scores["percentage"] < 50:
            actions.insert(0, "ÖNCE: Doküman yapısını ve formatını yeniden gözden geçiriniz")
        
        return actions

    def generate_recommendations(self, analysis_results: Dict, scores: Dict) -> List[str]:
        recommendations = []
        total_percentage = scores["percentage"]
        
        if total_percentage >= 70:
            recommendations.append(f"✅ Bakım Talimatları GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ Bakım Talimatları GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            if category_score < 40:
                recommendations.append(f"🔴 {category} bölümü yetersiz (%{category_score:.1f})")
                missing_items = [name for name, result in results.items() if not result.found]
                if missing_items:
                    recommendations.append(f"   Eksik: {', '.join(missing_items[:3])}")
            elif category_score < 70:
                recommendations.append(f"🟡 {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"🟢 {category} bölümü yeterli (%{category_score:.1f})")
        
        return recommendations

    def analyze_maintenance_report(self, file_path: str) -> Dict[str, Any]:
        logger.info("Bakım Talimatları analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        if detected_lang != 'tr':
            logger.info(f"{detected_lang.upper()} dilinden Türkçe'ye çeviriliyor...")
            text = self.translate_to_turkish(text, detected_lang)
        
        if not self.validate_maintenance_document(text):
            return {
                "error": "YANLIŞ DOKÜMAN: Bu dosya bakım talimatları dökümanı değil!",
                "document_type": "UNKNOWN",
                "suggestion": "Lütfen geçerli bir bakım talimatları dökümanı yükleyiniz."
            }
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        recommendations = self.generate_recommendations(analysis_results, scores)
        improvement_actions = self.generate_improvement_actions(analysis_results, scores)
        
        final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
        
        return {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "iyilestirme_eylemleri": improvement_actions,
            "ozet": {
                "toplam_puan": scores["total_score"],
                "yuzde": scores["percentage"],
                "durum": final_status,
                "rapor_tipi": "BAKIM_TALIMATLARI"
            }
        }

# ============================================
# HELPER FUNCTIONS (3-STAGE VALIDATION)
# ============================================
def validate_document_server(text, validation_keywords):
    critical_terms_data = validation_keywords.get('critical_terms', [])
    
    critical_terms = []
    for item in critical_terms_data:
        if isinstance(item, dict) and 'keywords' in item:
            critical_terms.append(item['keywords'])
        elif isinstance(item, list):
            critical_terms.append(item)
    
    if not critical_terms:
        logger.warning("⚠️ Critical terms bulunamadı")
        return True
    
    category_found = [any(re.search(rf"\b{term}\b", text, re.IGNORECASE) for term in category) for category in critical_terms]
    valid = sum(category_found)
    logger.info(f"Validasyon: {valid}/{len(critical_terms)} kritik kategori")
    return valid >= 2

def check_strong_keywords_first_pages(filepath, validation_keywords):
    strong_keywords = validation_keywords.get('strong_keywords', [])
    
    if not strong_keywords:
        logger.warning("⚠️ Strong keywords bulunamadı")
        return True
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=300, first_page=1, last_page=1)
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        found = [kw for kw in strong_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa: {len(found)} özgü kelime")
        return len(found) >= 1
    except Exception as e:
        logger.warning(f"OCR hatası: {e}")
        return False

def check_excluded_keywords_first_pages(filepath, validation_keywords):
    excluded_keywords = validation_keywords.get('excluded_keywords', [])
    
    if not excluded_keywords:
        logger.warning("⚠️ Excluded keywords bulunamadı")
        return False
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=1)
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        found = [kw for kw in excluded_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        return len(found) >= 1
    except Exception as e:
        logger.warning(f"Excluded OCR hatası: {e}")
        return False

def get_conclusion_message(status, percentage):
    if status == "PASS":
        return f"Bakım talimatları yüksek kalitede ve standartlara uygun (%{percentage:.0f})"
    return f"Bakım talimatları yetersiz, kapsamlı revizyon gerekli (%{percentage:.0f})"

def get_main_issues(analysis_result):
    issues = []
    for category, score_data in analysis_result['puanlama']['category_scores'].items():
        if score_data['percentage'] < 50:
            issues.append(f"{category} kategorisinde ciddi eksiklikler")
    if not issues and analysis_result['puanlama']['total_score'] < 50:
        issues = ["Genel makine bilgileri eksik", "Güvenlik önlemleri yetersiz", "Bakım türleri tanımlanmamış", "Adım adım talimatlar eksik"]
    return issues[:4]

def map_language_code(lang_code):
    lang_mapping = {'tr': 'turkish', 'en': 'english', 'de': 'german', 'fr': 'french', 'es': 'spanish', 'it': 'italian'}
    return lang_mapping.get(lang_code, 'turkish')

# ============================================
# FLASK SERVICE
# ============================================
app = Flask(__name__)

app.config['SQLALCHEMY_DATABASE_URI'] = Config.SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = Config.SQLALCHEMY_TRACK_MODIFICATIONS
app.config['SQLALCHEMY_ECHO'] = Config.SQLALCHEMY_ECHO

UPLOAD_FOLDER = 'temp_uploads_bakim'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/bakimtalimatlari-report', methods=['POST'])
def analyze_maintenance_report():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '' or not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file'}), 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        analyzer = MaintenanceReportAnalyzer(app=app)
        file_ext = os.path.splitext(filepath)[1].lower()
        
        # 3 AŞAMALI KONTROL
        if file_ext == '.pdf':
            logger.info("Aşama 1: Bakım özgü kelime kontrolü...")
            if check_strong_keywords_first_pages(filepath, analyzer.validation_keywords):
                logger.info("✅ Aşama 1 geçti")
            else:
                logger.info("Aşama 2: Excluded kelime kontrolü...")
                if check_excluded_keywords_first_pages(filepath, analyzer.validation_keywords):
                    logger.info("❌ Excluded kelimeler bulundu")
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({'error': 'Invalid document type', 'message': 'Yüklediğiniz dosya bakım talimatları değil!'}), 400
                else:
                    logger.info("Aşama 3: Tam doküman kontrolü...")
                    try:
                        with open(filepath, 'rb') as f:
                            pdf_reader = PyPDF2.PdfReader(f)
                            text = "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
                        if not text or len(text.strip()) < 50 or not validate_document_server(text, analyzer.validation_keywords):
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({'error': 'Invalid document type', 'message': 'Yüklediğiniz dosya bakım talimatları değil!'}), 400
                    except Exception as e:
                        logger.error(f"Aşama 3 hatası: {e}")
                        try:
                            os.remove(filepath)
                        except:
                            pass
                        return jsonify({'error': 'Analysis failed'}), 500
                    
        elif file_ext in ['.docx', '.doc', '.txt']:
            logger.info(f"DOCX/TXT dosyası için tam doküman kontrolü: {file_ext}")
            text = ""
            if file_ext in ['.docx', '.doc']:
                text = analyzer.extract_text_from_docx(filepath)
            elif file_ext == '.txt':
                text = analyzer.extract_text_from_txt(filepath)
            
            if not text or len(text.strip()) == 0:
                try:
                    os.remove(filepath)
                except:
                    pass
                return jsonify({'error': 'Text extraction failed', 'message': 'Dosyadan yeterli metin çıkarılamadı'}), 400
            
            if not validate_document_server(text, analyzer.validation_keywords):
                try:
                    os.remove(filepath)
                except:
                    pass
                return jsonify({'error': 'Invalid document type', 'message': 'Yüklediğiniz dosya bakım talimatları değil!'}), 400

        logger.info(f"Bakım analizi yapılıyor: {filename}")
        analysis_result = analyzer.analyze_maintenance_report(filepath)
        
        try:
            os.remove(filepath)
        except:
            pass

        if 'error' in analysis_result:
            return jsonify({'error': 'Analysis failed', 'message': analysis_result['error']}), 400

        overall_percentage = analysis_result['puanlama']['percentage']
        status = "PASS" if overall_percentage >= 70 else "FAIL"
        
        response_data = {
            'analysis_date': analysis_result.get('analiz_tarihi'),
            'analysis_id': f"bakim_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            'category_scores': {},
            'extracted_values': analysis_result.get('cikarilan_degerler', {}),
            'file_type': 'BAKIM_TALIMATLARI',
            'filename': filename,
            'language_info': {'detected_language': map_language_code(analysis_result['dosya_bilgisi']['detected_language'])},
            'overall_score': {
                'percentage': round(overall_percentage, 2),
                'total_points': analysis_result['puanlama']['total_score'],
                'max_points': 100,
                'status': status,
                'status_tr': 'GEÇERLİ' if status == "PASS" else 'GEÇERSİZ'
            },
            'recommendations': analysis_result.get('oneriler', []),
            'improvement_actions': analysis_result.get('iyilestirme_eylemleri', []),
            'summary': {
                'is_valid': status == "PASS",
                'conclusion': get_conclusion_message(status, overall_percentage),
                'main_issues': [] if status == "PASS" else get_main_issues(analysis_result)
            }
        }
        
        for category, score_data in analysis_result['puanlama']['category_scores'].items():
            response_data['category_scores'][category] = {
                'score': score_data['normalized'],
                'max_score': score_data['max_weight'],
                'percentage': score_data['percentage'],
                'status': 'PASS' if score_data['percentage'] >= 70 else 'FAIL'
            }

        return jsonify({'success': True, 'message': 'Bakım Talimatları başarıyla analiz edildi', 'analysis_service': 'bakim_talimatları', 'data': response_data})

    except Exception as e:
        logger.error(f"API error: {str(e)}")
        return jsonify({'error': 'Server error', 'message': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy', 'service': 'Bakım Talimatları Analyzer API', 'version': '1.0.0', 'tesseract_available': tesseract_available, 'report_type': 'BAKIM_TALIMATLARI'})

@app.route('/', methods=['GET'])
def index():
    return jsonify({'service': 'Bakım Talimatları Analyzer API', 'version': '1.0.0', 'endpoints': {'POST /api/bakimtalimatlari-report': 'Bakım talimatları analizi', 'GET /api/health': 'Health check'}})

with app.app_context():
    db.init_app(app)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8014))
    logger.info(f"🚀 Bakım Talimatları Servisi - Port: {port}")
    app.run(host='0.0.0.0', port=port, debug=False)