# ============================================
# HRC KUVVET-BASINÇ ÖLÇÜM RAPORU ANALİZ SERVİSİ
# Standalone Service - Azure App Service Ready
# Port: 8013
# Database-Driven Configuration
# ============================================

# ============================================
# IMPORTS
# ============================================
import os
import re
from datetime import datetime
from typing import Dict, List, Any
import PyPDF2
from docx import Document
from dataclasses import dataclass
import logging
import pytesseract
import cv2
import numpy as np
from PIL import Image
import pdf2image

from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename

# ============================================
# DATABASE IMPORTS (YENİ)
# ============================================
from flask import current_app
from database import db, init_db
from db_loader import load_service_config

# ============================================
# CONFIG IMPORT (YENİ)
# ============================================
from config import Config

# ============================================
# LOGGING CONFIGURATION
# ============================================
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============================================
# LANGUAGE DETECTION
# ============================================
try:
    from langdetect import detect
    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    LANGUAGE_DETECTION_AVAILABLE = False
    logger.warning("langdetect modülü bulunamadı - dil tespiti devre dışı")

# ============================================
# TESSERACT CHECK
# ============================================
def check_tesseract_installation():
    """Check if Tesseract is properly installed"""
    try:
        version = pytesseract.get_tesseract_version()
        logger.info(f"Tesseract OCR kurulu - Sürüm: {version}")
        return True, f"Tesseract v{version}"
    except Exception as e:
        logger.error(f"Tesseract OCR kurulu değil: {e}")
        return False, str(e)

tesseract_available, tesseract_info = check_tesseract_installation()

# ============================================
# ANALİZ SINIFI - DATA CLASSES
# ============================================
@dataclass
class HRCAnalysisResult:
    """HRC analiz sonucu veri sınıfı"""
    criteria_name: str
    found: bool
    content: str
    score: int
    max_score: int
    details: Dict[str, Any]

# ============================================
# ANALİZ SINIFI - MAIN ANALYZER
# ============================================
class HRCReportAnalyzer:
    """HRC Kuvvet-Basınç Ölçüm Raporu analiz sınıfı"""
    
    def __init__(self, app=None):
        logger.info("HRC Kuvvet-Basınç Ölçüm Raporu analiz sistemi başlatılıyor...")
        logger.info("✅ OCR sistemi aktif - Tüm dosyalar OCR ile işlenmektedir")
        
        # Flask app context varsa DB'den yükle, yoksa boş başlat
        if app:
            with app.app_context():
                try:
                    config = load_service_config('hrc_report')
                    
                    # DB'den yüklenen veriler
                    self.criteria_weights = config.get('criteria_weights', {})
                    self.criteria_details = config.get('criteria_details', {})
                    self.pattern_definitions = config.get('pattern_definitions', {})
                    self.validation_keywords = config.get('validation_keywords', {})
                    self.category_actions = config.get('category_actions', {})
                    
                    logger.info(f"✅ Veritabanından yüklendi: {len(self.criteria_weights)} kategori")
                    
                except Exception as e:
                    logger.error(f"⚠️ Veritabanından yükleme başarısız: {e}")
                    logger.warning("⚠️ Fallback: Boş config kullanılıyor")
                    self.criteria_weights = {}
                    self.criteria_details = {}
                    self.pattern_definitions = {}
                    self.validation_keywords = {}
                    self.category_actions = {}
        else:
            # Flask app yoksa boş başlat (eski davranış)
            logger.warning("⚠️ Flask app context yok, boş config kullanılıyor")
            self.criteria_weights = {}
            self.criteria_details = {}
            self.pattern_definitions = {}
            self.validation_keywords = {}
            self.category_actions = {}
    
    def detect_language(self, text: str) -> str:
        """Metin dilini tespit et"""
        if not LANGUAGE_DETECTION_AVAILABLE:
            return 'tr'
        
        try:
            sample_text = text[:500].strip()
            if not sample_text:
                return 'tr'
                
            detected_lang = detect(sample_text)
            logger.info(f"Tespit edilen dil: {detected_lang}")
            return detected_lang
            
        except Exception as e:
            logger.warning(f"Dil tespiti başarısız: {e}")
            return 'tr'
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Dosya formatına göre OCR ile metin çıkarma"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_with_ocr(file_path)
        elif file_ext in ['.docx', '.doc']:
            try:
                text = self.extract_text_from_docx(file_path)
                if len(text.strip()) > 50:
                    return text
                else:
                    return self.extract_text_with_ocr(file_path)
            except:
                return self.extract_text_with_ocr(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.error(f"Desteklenmeyen dosya formatı: {file_ext}")
            return ""
    
    def extract_text_with_ocr(self, file_path: str) -> str:
        """OCR ile metin çıkarma"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                pages = pdf2image.convert_from_path(file_path, dpi=200)
                all_text = ""
                
                for i, page in enumerate(pages):
                    opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
                    gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
                    processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
                    
                    text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
                    all_text += text + "\n"
                    logger.info(f"PDF sayfa {i+1}/{len(pages)} OCR tamamlandı")
                
                return all_text.strip()
            else:
                return ""
                
        except Exception as e:
            logger.error(f"OCR hatası: {e}")
            return ""

    def extract_text_from_docx(self, docx_path: str) -> str:
        """DOCX'den metin çıkarma"""
        try:
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX metin çıkarma hatası: {e}")
            return ""
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """TXT dosyasından metin çıkarma"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except:
            try:
                with open(txt_path, 'r', encoding='cp1254') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"TXT metin çıkarma hatası: {e}")
                return ""
    
    def analyze_criteria(self, text: str, category: str) -> Dict[str, HRCAnalysisResult]:
        """Kriterleri analiz et"""
        results = {}
        criteria = self.criteria_details.get(category, {})
        
        for criterion_name, criterion_data in criteria.items():
            pattern = criterion_data["pattern"]
            weight = criterion_data["weight"]
            
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if matches:
                content = f"Bulunan: {str(matches[:3])}"
                found = True
                
                if weight <= 2:
                    score = min(weight, len(matches))
                else:
                    score = min(weight, len(matches) * (weight // 2))
                    score = max(score, weight // 2)
            else:
                content = "Bulunamadı"
                found = False
                score = 0
            
            results[criterion_name] = HRCAnalysisResult(
                criteria_name=criterion_name,
                found=found,
                content=content,
                score=score,
                max_score=weight,
                details={"pattern_used": pattern, "matches_count": len(matches) if matches else 0}
            )
        
        return results

    def calculate_scores(self, analysis_results: Dict[str, Dict[str, HRCAnalysisResult]]) -> Dict[str, Any]:
        """Puanları hesapla"""
        category_scores = {}
        total_score = 0
        
        for category, results in analysis_results.items():
            category_max = self.criteria_weights[category]
            category_earned = sum(result.score for result in results.values())
            category_possible = sum(result.max_score for result in results.values())
            
            if category_possible > 0:
                percentage = (category_earned / category_possible) * 100
                normalized_score = (percentage / 100) * category_max
            else:
                percentage = 0
                normalized_score = 0
            
            category_scores[category] = {
                "earned": category_earned,
                "possible": category_possible,
                "normalized": round(normalized_score, 2),
                "max_weight": category_max,
                "percentage": round(percentage, 2)
            }
            
            total_score += normalized_score
        
        return {
            "category_scores": category_scores,
            "total_score": round(total_score, 2),
            "percentage": round(total_score, 2)
        }

    def extract_specific_values(self, text: str) -> Dict[str, Any]:
        """HRC raporuna özgü spesifik değerleri çıkar"""
        values = {
            "robot_modeli": "Bulunamadı",
            "test_tarihi": "Bulunamadı",
            "olcum_cihazi": "Bulunamadı"
        }
        
        # DB'den pattern'leri al
        extract_values = self.pattern_definitions.get('extract_values', {})
        
        # ROBOT MODELİ
        robot_patterns = extract_values.get('robot_modeli', [])
        for pattern in robot_patterns:
            match = re.search(pattern, text)
            if match:
                values["robot_modeli"] = match.group(1).strip()
                break
        
        # TEST TARİHİ
        date_patterns = extract_values.get('test_tarihi', [])
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                values["test_tarihi"] = match.group(1).strip()
                break

        # ÖLÇÜM CİHAZI
        olcum_cihazi_patterns = extract_values.get('olcum_cihazi', [])
        for pattern in olcum_cihazi_patterns:
            match = re.search(pattern, text)
            if match:
                found_value = match.group(1).strip()
                if 10 <= len(found_value) <= 25:
                    values["olcum_cihazi"] = found_value
                    break
                
        return values

    def generate_recommendations(self, analysis_results: Dict, scores: Dict) -> List[str]:
        """Öneriler oluştur"""
        recommendations = []
        
        total_percentage = scores["percentage"]
        
        if total_percentage >= 70:
            recommendations.append(f"✅ HRC Kuvvet-Basınç Raporu GEÇERLİ (Toplam: %{total_percentage:.1f})")
        else:
            recommendations.append(f"❌ HRC Kuvvet-Basınç Raporu GEÇERSİZ (Toplam: %{total_percentage:.1f})")
        
        for category, results in analysis_results.items():
            category_score = scores["category_scores"][category]["percentage"]
            
            if category_score < 40:
                recommendations.append(f"🔴 {category} bölümü yetersiz (%{category_score:.1f})")
            elif category_score < 70:
                recommendations.append(f"🟡 {category} bölümü geliştirilmeli (%{category_score:.1f})")
            else:
                recommendations.append(f"🟢 {category} bölümü yeterli (%{category_score:.1f})")
        
        return recommendations

    def analyze_hrc_report(self, file_path: str) -> Dict[str, Any]:
        """Ana HRC analiz fonksiyonu"""
        logger.info("HRC Kuvvet-Basınç Ölçüm Raporu analizi başlatılıyor...")
        
        if not os.path.exists(file_path):
            return {"error": f"Dosya bulunamadı: {file_path}"}
        
        text = self.extract_text_from_file(file_path)
        if not text:
            return {"error": "Dosyadan metin çıkarılamadı"}
        
        detected_lang = self.detect_language(text)
        
        analysis_results = {}
        for category in self.criteria_weights.keys():
            analysis_results[category] = self.analyze_criteria(text, category)
        
        scores = self.calculate_scores(analysis_results)
        extracted_values = self.extract_specific_values(text)
        recommendations = self.generate_recommendations(analysis_results, scores)
        
        final_status = "PASS" if scores["percentage"] >= 70 else "FAIL"
        
        return {
            "analiz_tarihi": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "dosya_bilgisi": {
                "file_path": file_path,
                "file_type": os.path.splitext(file_path)[1].lower(),
                "detected_language": detected_lang
            },
            "cikarilan_degerler": extracted_values,
            "kategori_analizleri": analysis_results,
            "puanlama": scores,
            "oneriler": recommendations,
            "ozet": {
                "toplam_puan": scores["total_score"],
                "yuzde": scores["percentage"],
                "durum": final_status,
                "rapor_tipi": "HRC_KUVVET_BASINC_RAPORU"
            }
        }

# ============================================
# HELPER FUNCTIONS
# ============================================
def validate_document_server(text, validation_keywords):
    """Server kodunda HRC doküman validasyonu - DB'den gelen keywords ile"""
    
    # DB'den gelen critical_terms
    critical_terms_data = validation_keywords.get('critical_terms', [])
    
    # Liste formatına dönüştür (eski format uyumluluğu)
    critical_terms = []
    for item in critical_terms_data:
        if isinstance(item, dict) and 'keywords' in item:
            critical_terms.append(item['keywords'])
        elif isinstance(item, list):
            critical_terms.append(item)
    
    # Eğer DB'den veri gelmediyse, boş validasyon (hata verme)
    if not critical_terms:
        logger.warning("⚠️ Critical terms bulunamadı, validasyon atlanıyor")
        return True  # Varsayılan: geçerli kabul et
    
    category_found = []
    
    for i, category in enumerate(critical_terms):
        found_in_category = False
        for term in category:
            if re.search(rf"\b{term}\b", text, re.IGNORECASE):
                found_in_category = True
                logger.info(f"HRC Kategori {i+1} bulundu: '{term}'")
                break
        category_found.append(found_in_category)
    
    valid_categories = sum(category_found)
    logger.info(f"HRC doküman validasyonu: {valid_categories}/{len(critical_terms)} kritik kategori bulundu")
    
    return valid_categories >= len(critical_terms)


def check_strong_keywords_first_pages(filepath, validation_keywords):
    """İlk 1-2 sayfada HRC'ye özgü kelimeleri OCR ile ara - DB'den gelen keywords ile"""
    
    # DB'den strong keywords al
    strong_keywords = validation_keywords.get('strong_keywords', [])
    
    # Eğer DB'den veri gelmediyse, validasyon atla
    if not strong_keywords:
        logger.warning("⚠️ Strong keywords bulunamadı, validasyon atlanıyor")
        return True  # Varsayılan: geçerli kabul et
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=300, first_page=1, last_page=2)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 3 -l tur+eng')
            all_text += text.lower() + " "
        
        found_keywords = [kw for kw in strong_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa HRC kontrol: {len(found_keywords)} özgü kelime")
        return len(found_keywords) >= 2
        
    except Exception as e:
        logger.warning(f"İlk sayfa HRC kontrol hatası: {e}")
        return False


def check_excluded_keywords_first_pages(filepath, validation_keywords):
    """İlk 1-2 sayfada istenmeyen rapor türlerinin kelimelerini ara - DB'den"""
    
    # DB'den excluded keywords al
    excluded_keywords = validation_keywords.get('excluded_keywords', [])
    
    # Eğer DB'den veri gelmediyse, validasyon atla
    if not excluded_keywords:
        logger.warning("⚠️ Excluded keywords bulunamadı, validasyon atlanıyor")
        return False  # Varsayılan: excluded yok kabul et
    
    try:
        pages = pdf2image.convert_from_path(filepath, dpi=200, first_page=1, last_page=2)
        
        all_text = ""
        for page in pages:
            opencv_image = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
            
            text = pytesseract.image_to_string(processed, config='--oem 3 --psm 6 -l tur+eng')
            all_text += text.lower() + " "
        
        found_excluded = [kw for kw in excluded_keywords if re.search(rf"\b{kw.lower()}\b", all_text)]
        logger.info(f"İlk sayfa excluded kontrol: {len(found_excluded)} istenmeyen kelime")
        return len(found_excluded) >= 1
        
    except Exception as e:
        logger.warning(f"İlk sayfa excluded kontrol hatası: {e}")
        return False


def get_conclusion_message_hrc(status, percentage):
    """Sonuç mesajını döndür"""
    if status == "PASS":
        return f"HRC kuvvet-basınç raporu EN ISO 10218 ve ISO/TS 15066 standartlarına uygundur (%{percentage:.0f})"
    else:
        return f"HRC raporu standartlara uygun değil, kapsamlı revizyon gerekli (%{percentage:.0f})"


def get_main_issues_hrc(analysis_result):
    """Ana sorunları listele"""
    issues = []
    
    for category, score_data in analysis_result['puanlama']['category_scores'].items():
        if score_data['percentage'] < 50:
            issues.append(f"{category} kategorisinde ciddi eksiklikler")
    
    if not issues and analysis_result['puanlama']['total_score'] < 50:
        issues = [
            "Robot modeli ve seri numarası eksik",
            "Test koşulları ve senaryo tanımı eksik",
            "Kuvvet ve basınç ölçüm sonuçları eksik",
            "Risk değerlendirmesi yapılmamış"
        ]
    
    return issues[:4]


# ============================================
# FLASK SERVICE
# ============================================
app = Flask(__name__)

# Database configuration (YENİ)
app.config['SQLALCHEMY_DATABASE_URI'] = Config.SQLALCHEMY_DATABASE_URI
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = Config.SQLALCHEMY_TRACK_MODIFICATIONS
app.config['SQLALCHEMY_ECHO'] = Config.SQLALCHEMY_ECHO

UPLOAD_FOLDER = 'temp_uploads_hrc'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/api/hrc-report', methods=['POST'])
def analyze_hrc_report():
    """HRC Kuvvet-Basınç Ölçüm Raporu analiz endpoint"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        logger.info(f"HRC Kuvvet-Basınç Raporu kontrol ediliyor: {filename}")
        
        # Create analyzer instance
        analyzer = HRCReportAnalyzer(app=app)

        # ÜÇ AŞAMALI HRC KONTROLÜ
        logger.info(f"Üç aşamalı HRC kontrolü başlatılıyor: {filename}")
        file_ext = os.path.splitext(filepath)[1].lower()

        if file_ext == '.pdf':
            # PDF için üç aşamalı kontrol (OCR dahil)
            logger.info("Aşama 1: İlk sayfa HRC özgü kelime kontrolü...")
            if check_strong_keywords_first_pages(filepath, analyzer.validation_keywords):
                logger.info("✅ Aşama 1 geçti - HRC özgü kelimeler bulundu")
            else:
                logger.info("Aşama 2: İlk sayfa excluded kelime kontrolü...")
                if check_excluded_keywords_first_pages(filepath, analyzer.validation_keywords):
                    logger.info("❌ Aşama 2'de excluded kelimeler bulundu - HRC değil")
                    try:
                        os.remove(filepath)
                    except:
                        pass
                    return jsonify({
                        'error': 'Invalid document type',
                        'message': 'Bu dosya HRC raporu değil (farklı rapor türü tespit edildi). Lütfen HRC kuvvet-basınç ölçüm raporu yükleyiniz.',
                        'details': {
                            'filename': filename,
                            'document_type': 'OTHER_REPORT_TYPE',
                            'required_type': 'HRC_KUVVET_BASINC_RAPORU'
                        }
                    }), 400
                else:
                    # AŞAMA 3: PyPDF2 ile tam doküman kontrolü
                    logger.info("Aşama 3: Tam doküman critical terms kontrolü...")
                    try:
                        with open(filepath, 'rb') as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            text = ""
                            for page in pdf_reader.pages:
                                text += page.extract_text() + "\n"
                        
                        if not text or len(text.strip()) == 0:
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({
                                'error': 'Text extraction failed',
                                'message': 'Dosyadan yeterli metin çıkarılamadı'
                            }), 400
                        
                        if not validate_document_server(text, analyzer.validation_keywords):
                            try:
                                os.remove(filepath)
                            except:
                                pass
                            return jsonify({
                                'error': 'Invalid document type',
                                'message': 'Yüklediğiniz dosya HRC kuvvet-basınç ölçüm raporu değil! Lütfen geçerli bir HRC raporu yükleyiniz.',
                                'details': {
                                    'filename': filename,
                                    'document_type': 'NOT_HRC_REPORT',
                                    'required_type': 'HRC_KUVVET_BASINC_RAPORU'
                                }
                            }), 400
                            
                    except Exception as e:
                        logger.error(f"Aşama 3 hatası: {e}")
                        try:
                            os.remove(filepath)
                        except:
                            pass
                        return jsonify({
                            'error': 'Analysis failed',
                            'message': 'Dosya analizi sırasında hata oluştu'
                        }), 500

        elif file_ext in ['.docx', '.doc', '.txt']:
            # DOCX/TXT için sadece tam doküman kontrolü
            logger.info(f"DOCX/TXT dosyası için tam doküman kontrolü: {file_ext}")
            text = ""
            if file_ext in ['.docx', '.doc']:
                text = analyzer.extract_text_from_docx(filepath)
            elif file_ext == '.txt':
                text = analyzer.extract_text_from_txt(filepath)
            
            if not text or len(text.strip()) == 0:
                try:
                    os.remove(filepath)
                except:
                    pass
                return jsonify({
                    'error': 'Text extraction failed',
                    'message': 'Dosyadan yeterli metin çıkarılamadı'
                }), 400
            
            if not validate_document_server(text, analyzer.validation_keywords):
                try:
                    os.remove(filepath)
                except:
                    pass
                return jsonify({
                    'error': 'Invalid document type',
                    'message': 'Yüklediğiniz dosya HRC kuvvet-basınç ölçüm raporu değil! Lütfen geçerli bir HRC raporu yükleyiniz.',
                    'details': {
                        'filename': filename,
                        'document_type': 'NOT_HRC_REPORT',
                        'required_type': 'HRC_KUVVET_BASINC_RAPORU'
                    }
                }), 400

        logger.info(f"HRC raporu doğrulandı, analiz başlatılıyor: {filename}")
        analysis_result = analyzer.analyze_hrc_report(filepath)
        
        try:
            os.remove(filepath)
        except:
            pass

        if 'error' in analysis_result:
            return jsonify({'error': 'Analysis failed', 'message': analysis_result['error']}), 400

        overall_percentage = analysis_result['puanlama']['percentage']
        status = "PASS" if overall_percentage >= 70 else "FAIL"
        
        response_data = {
            'analysis_date': analysis_result.get('analiz_tarihi'),
            'analysis_id': f"hrc_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            'category_scores': {},
            'extracted_values': analysis_result.get('cikarilan_degerler', {}),
            'file_type': 'HRC_KUVVET_BASINC_RAPORU',
            'filename': filename,
            'overall_score': {
                'percentage': round(overall_percentage, 2),
                'total_points': analysis_result['puanlama']['total_score'],
                'max_points': 100,
                'status': status,
                'status_tr': 'GEÇERLİ' if status == "PASS" else 'GEÇERSİZ'
            },
            'recommendations': analysis_result.get('oneriler', []),
            'summary': {
                'is_valid': status == "PASS",
                'conclusion': get_conclusion_message_hrc(status, overall_percentage),
                'main_issues': [] if status == "PASS" else get_main_issues_hrc(analysis_result)
            }
        }
        
        for category, score_data in analysis_result['puanlama']['category_scores'].items():
            response_data['category_scores'][category] = {
                'score': score_data['normalized'],
                'max_score': score_data['max_weight'],
                'percentage': score_data['percentage'],
                'status': 'PASS' if score_data['percentage'] >= 70 else 'FAIL'
            }

        return jsonify({
            'success': True,
            'message': 'HRC Kuvvet-Basınç Raporu başarıyla analiz edildi',
            'analysis_service': 'hrc',
            'data': response_data
        })

    except Exception as e:
        logger.error(f"API error: {str(e)}")
        return jsonify({'error': 'Server error', 'message': str(e)}), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check"""
    return jsonify({
        'status': 'healthy',
        'service': 'HRC Force-Pressure Report Analyzer API',
        'version': '1.0.0',
        'tesseract_available': tesseract_available,
        'report_type': 'HRC_KUVVET_BASINC_RAPORU'
    })


@app.route('/', methods=['GET'])
def index():
    """Index page"""
    return jsonify({
        'service': 'HRC Force-Pressure Report Analyzer API',
        'version': '1.0.0',
        'endpoints': {
            'POST /api/hrc-report': 'HRC kuvvet-basınç raporu analizi',
            'GET /api/health': 'Health check',
            'GET /': 'API info'
        }
    })


# ============================================
# DATABASE INITIALIZATION
# ============================================
with app.app_context():
    db.init_app(app)


# ============================================
# APPLICATION ENTRY POINT (Azure-Friendly)
# ============================================
if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("HRC Kuvvet-Basınç Ölçüm Raporu Analiz Servisi")
    logger.info("=" * 60)
    
    port = int(os.environ.get('PORT', 8013))
    logger.info(f"🚀 Servis başlatılıyor - Port: {port}")
    
    app.run(host='0.0.0.0', port=port, debug=False)